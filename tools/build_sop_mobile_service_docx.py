from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("SOP_Mobile_Service_Setup_and_Pack_Down_Template.docx")

FONT = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
RISK_FILL = "FBEAEA"
BORDER = "C9D3DF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, before=0, after=6, line=1.25, keep_next=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next is not None:
        fmt.keep_with_next = keep_next


def style_para(p, size=11, color=BODY, bold=False, italic=False):
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_para(
    doc,
    text="",
    *,
    style=None,
    size=11,
    color=BODY,
    bold=False,
    italic=False,
    align=None,
    before=0,
    after=6,
    line=1.25,
    keep_next=None,
):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=before, after=after, line=line, keep_next=keep_next)
    if align is not None:
        p.alignment = align
    return p


def add_mixed_para(doc, parts, *, before=0, after=6, line=1.25, align=None, keep_next=None):
    p = doc.add_paragraph()
    for text, opts in parts:
        run = p.add_run(text)
        set_run_font(
            run,
            size=opts.get("size", 11),
            color=opts.get("color", BODY),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )
    set_paragraph_format(p, before=before, after=after, line=line, keep_next=keep_next)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    style_para(
        p,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_cell_text(cell, text, *, bold=False, size=9.5, color=BODY, align=None, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    if text is not None:
        for idx, part in enumerate(str(text).split("\n")):
            if idx:
                p.add_run().add_break()
            run = p.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold)
    set_paragraph_format(p, before=0, after=0, line=1.15)
    if align is not None:
        p.alignment = align


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **margins):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ["top", "start", "bottom", "end"]:
        if m not in margins:
            continue
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[m]))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, sz="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell, **CELL_MARGIN)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            size=font_size,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=header_fill,
        )
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for idx, text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, len(row_data) - 1) and len(str(text)) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], text, size=font_size, align=align)
    add_para(doc, "", after=3)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D7DBE2", sz="6")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.text = ""
    r = p.add_run(title.upper() + "\n")
    set_run_font(r, size=10, color=DARK_BLUE, bold=True)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=BODY)
    set_paragraph_format(p, before=0, after=0, line=1.2)
    add_para(doc, "", after=2)
    return table


def add_step_table(doc, steps):
    rows = [(str(idx), step) for idx, step in enumerate(steps, start=1)]
    return add_table(
        doc,
        ["Step", "Action"],
        rows,
        [900, 8460],
        header_fill=LIGHT_BLUE,
        font_size=9,
    )


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abs_base = max(existing_abs + [0]) + 1
    num_base = max(existing_num + [0]) + 1

    def make_abs(abs_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def make_num(num_id, abs_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abs_id))
        num.append(abstract_ref)
        numbering.append(num)

    make_abs(abs_base, "bullet", "\uf0b7", "Symbol")
    make_num(num_base, abs_base)
    make_abs(abs_base + 1, "decimal", "%1.")
    make_num(num_base + 1, abs_base + 1)
    return {"bullet": num_base, "decimal": num_base + 1, "decimal_abs": abs_base + 1}


def add_numbering_instance(doc, abstract_num_id):
    numbering = doc.part.numbering_part.element
    existing_num = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    num_id = max(existing_num + [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(p, num_id):
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.375)
    fmt.first_line_indent = Inches(-0.188)
    fmt.space_after = Pt(4)
    fmt.line_spacing = 1.25


def add_list_item(doc, text, num_id, *, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, color=BODY)
    apply_num(p, num_id)
    return p


def add_bullet_item(doc, text, *, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.text = ""
    r = p.add_run(text)
    set_run_font(r, size=size, color=BODY)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.375)
    fmt.first_line_indent = Inches(-0.188)
    fmt.space_after = Pt(4)
    fmt.line_spacing = 1.25
    return p


def build_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True


def setup_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    for cell in table.rows[0].cells:
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            pass
    left, right = table.rows[0].cells
    set_cell_text(left, "SOP | Mobile Service Setup and Pack Down", size=8.5, color=MUTED)
    set_cell_text(right, "[Business Name]", size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    footer = section.footer
    footer.paragraphs[0].text = ""
    ftable = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ftable, [6500, 2860], indent_dxa=0)
    lcell, rcell = ftable.rows[0].cells
    set_cell_text(lcell, "Controlled document - verify latest version before use", size=8.5, color=MUTED)
    p = rcell.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" of ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "NUMPAGES")
    set_paragraph_format(p, before=0, after=0, line=1.0)
    for cell in ftable.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)


def build_document():
    doc = Document()
    build_styles(doc)
    nums = add_numbering(doc)
    setup_header_footer(doc)

    props = doc.core_properties
    props.title = "SOP Mobile Service Setup and Pack Down Template"
    props.subject = "Operational SOP template for a South African cleaning services business"
    props.author = "OpenAI Codex"
    props.keywords = "SOP, cleaning services, mobile service, setup, pack down, South Africa"

    # Cover
    add_para(doc, "", after=80)
    add_para(
        doc,
        "Standard Operating Procedure Template",
        size=10.5,
        bold=True,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=16,
    )
    add_para(
        doc,
        "Mobile Service Setup and Pack Down",
        size=28,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.08,
    )
    add_para(
        doc,
        "For a South African Cleaning Services Business",
        size=15,
        color=DARK_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=36,
    )
    add_para(
        doc,
        "Prepared for: [Business Name] | Template date: June 2026",
        size=11.5,
        bold=True,
        color=BODY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=6,
    )
    add_para(
        doc,
        "Use this controlled SOP to standardise crew mobilisation, on-site setup, service readiness, pack down, return-to-base checks, and field records.",
        size=10.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=36,
    )
    add_table(
        doc,
        ["Field", "Template Entry"],
        [
            ("Document ID", "[SOP-OPS-###]"),
            ("Version", "[0.1 / Draft / Approved]"),
            ("Effective date", "[dd/mm/yyyy]"),
            ("Review date", "[dd/mm/yyyy]"),
            ("Document owner", "[Operations Manager / SHEQ Officer]"),
            ("Approved by", "[Owner / Director / Accountable Manager]"),
            ("Applies to", "Mobile teams, drivers, supervisors, subcontractors, and temporary staff"),
            ("Service area", "[Province / Municipality / Client sites]"),
        ],
        [2500, 6860],
        font_size=9.5,
    )
    add_callout(
        doc,
        "Template control",
        "Replace bracketed fields before issue. Keep one controlled master copy and record any local changes in the revision history.",
        fill=CAUTION_FILL,
    )
    add_page_break(doc)

    # Document control
    add_heading(doc, "1. Document Control", 1)
    add_table(
        doc,
        ["Version", "Date", "Changed by", "Summary of change", "Approved by"],
        [
            ("0.1", "[dd/mm/yyyy]", "[Name]", "Initial template issue", "[Name]"),
            ("", "", "", "", ""),
            ("", "", "", "", ""),
        ],
        [1050, 1400, 1700, 3610, 1600],
        header_fill=LIGHT_GREY,
        font_size=9,
    )
    add_table(
        doc,
        ["Role / Area", "Controlled copy holder", "Issue method"],
        [
            ("Operations", "[Name / Role]", "[Shared drive / printed controlled copy]"),
            ("Mobile crew supervisors", "[Name / Role]", "[Crew pack / mobile tablet]"),
            ("SHEQ / HR / Training", "[Name / Role]", "[Training file / onboarding pack]"),
            ("Client-specific file", "[Account / Site]", "[Service agreement appendix]"),
        ],
        [2200, 3000, 4160],
        header_fill=LIGHT_GREY,
        font_size=9.2,
    )
    add_callout(
        doc,
        "Issue rule",
        "A printed copy is uncontrolled unless it shows the current version, effective date, and approval signature.",
    )

    # Purpose and scope
    add_heading(doc, "2. Purpose and Scope", 1)
    add_para(
        doc,
        "This SOP sets the minimum operating standard for preparing, setting up, completing, packing down, and resetting a mobile cleaning service team at a client site. It is written as a business template and must be adapted to the company risk profile, service menu, client contracts, and local municipal requirements.",
    )
    add_heading(doc, "2.1 Purpose", 2)
    add_table(
        doc,
        ["Purpose area", "Requirement"],
        [
            ("Readiness", "Ensure every mobile crew arrives with the right people, equipment, chemicals, PPE, records, and site information."),
            ("Safety and control", "Prevent unsafe work, uncontrolled chemical use, property damage, waste mismanagement, and incomplete client handover."),
            ("Hold point", "Give supervisors a practical checklist before cleaning starts and before the team leaves site."),
            ("Evidence", "Create consistent records for job completion, incidents, client sign-off, inventory control, and continuous improvement."),
        ],
        [2100, 7260],
        header_fill=LIGHT_GREY,
        font_size=9,
    )
    add_heading(doc, "2.2 Scope", 2)
    add_heading(doc, "Included services", 3)
    add_para(
        doc,
        "Commercial, residential, hospitality, retail, office, post-construction, deep-clean, carpet/upholstery, washroom, hygiene, and periodic specialist cleaning tasks performed by mobile teams.",
        after=4,
    )
    add_heading(doc, "Excluded unless separately authorised", 3)
    add_para(
        doc,
        "Pest control, medical/biohazard decontamination, high-access rope work, confined-space entry, asbestos work, hazardous waste transport, or any work requiring a specialist licence, permit, method statement, or client-specific approval.",
        after=4,
    )
    add_heading(doc, "2.3 Stop-Work Conditions", 2)
    add_table(
        doc,
        ["Condition", "Required response"],
        [
            ("Changed site hazards", "Client site hazards differ materially from the job card or risk assessment."),
            ("Missing controls", "Required PPE, SDS, signage, barriers, ventilation, water, electricity, or safe access is not available."),
            ("Chemical control failure", "Chemicals are unlabelled, leaking, expired, incompatible, or mixed without an approved dilution instruction."),
            ("Uncontrolled exposure", "A staff member, client, visitor, or member of the public could be exposed to an uncontrolled hazard."),
            ("Out-of-scope work", "The requested work is outside the approved scope or requires a permit, specialist training, or additional insurance approval."),
        ],
        [2600, 6760],
        header_fill=LIGHT_GREY,
        font_size=9,
    )

    # Compliance
    add_heading(doc, "3. South African Compliance Reference Points", 1)
    add_para(
        doc,
        "This section is a practical reminder for SOP users. It does not replace legal advice, client specifications, or a competent safety risk assessment.",
        italic=True,
        color=MUTED,
        size=10.5,
    )
    add_table(
        doc,
        ["Reference point", "Operational meaning for this SOP"],
        [
            (
                "Occupational Health and Safety Act 85 of 1993, including Hazardous Chemical Agents duties where applicable",
                "Keep risk assessments, training, PPE, chemical labels, and safety data sheets current. Do not start work where chemical, electrical, slip, trip, manual handling, or public-interface risks are uncontrolled.",
            ),
            (
                "Compensation for Occupational Injuries and Diseases Act / Compensation Fund obligations",
                "Maintain the company incident reporting process and ensure supervisors know how to escalate injuries, occupational exposures, and near misses.",
            ),
            (
                "Protection of Personal Information Act 4 of 2013",
                "Protect client contact details, access codes, photographs, incident records, and any personal information collected on job cards or mobile devices.",
            ),
            (
                "National Environmental Management: Waste Act 59 of 2008 and local municipal by-laws",
                "Separate general, recyclable, contaminated, and special waste streams where required. Do not discharge chemicals, slurry, or wash water into drains unless authorised by the client and local rules.",
            ),
            (
                "Client site rules, service agreement, and manufacturer instructions",
                "Follow site induction, permit, access, PPE, surface-care, chemical dilution, equipment, warranty, and sign-off requirements specific to each job.",
            ),
        ],
        [2800, 6560],
        header_fill=LIGHT_GREY,
        font_size=8.8,
    )
    add_para(
        doc,
        "Useful official sources for company verification: labour.gov.za, gov.za, justice.gov.za, and the relevant municipality or client site rules.",
        size=9.5,
        color=MUTED,
        italic=True,
        after=3,
    )

    # Roles
    add_heading(doc, "4. Roles and Responsibilities", 1)
    add_table(
        doc,
        ["Role", "Setup responsibilities", "Pack-down responsibilities"],
        [
            (
                "Operations Manager",
                "Approves SOP, confirms resources, assigns competent supervisors, checks service plans for high-risk or unusual work.",
                "Reviews completion records, incidents, stock losses, client feedback, and corrective actions.",
            ),
            (
                "Mobile Supervisor / Team Leader",
                "Leads toolbox talk, performs arrival assessment, confirms site boundaries, authorises service readiness.",
                "Controls final inspection, waste handling, client sign-off, inventory reconciliation, and departure clearance.",
            ),
            (
                "Driver / Mobile Unit Custodian",
                "Completes vehicle pre-use check, load restraint, fuel/route check, and safe parking plan.",
                "Checks load security, vehicle cleanliness, waste segregation, refuelling, mileage, and return-to-base reset.",
            ),
            (
                "Cleaning Technician",
                "Uses assigned PPE, prepares tools, follows chemical dilution and method statements, reports hazards.",
                "Cleans tools, returns chemicals, bags waste, clears signs only when area is safe, reports defects or damage.",
            ),
            (
                "Stores / Dispatcher",
                "Issues equipment, chemicals, SDS, consumables, registers, and job packs.",
                "Receives returns, updates stock, tags defective equipment, quarantines unknown chemicals or contaminated items.",
            ),
            (
                "SHEQ / Training Coordinator",
                "Maintains training matrix, risk templates, SDS library, induction records, and audit programme.",
                "Reviews incidents, near misses, toolbox talks, and corrective action evidence.",
            ),
        ],
        [2100, 3630, 3630],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )

    # Requirements
    add_heading(doc, "5. Required Equipment, PPE, Chemicals, and Records", 1)
    add_table(
        doc,
        ["Category", "Minimum field pack", "Pre-departure check"],
        [
            (
                "PPE",
                "Safety shoes, nitrile/chemical gloves, eye protection, reflective vest where needed, respiratory protection if required by SDS or risk assessment.",
                "Correct sizes, clean, undamaged, and available for each crew member.",
            ),
            (
                "Chemical control",
                "Approved chemicals, labelled containers, SDS access, dilution chart, measuring tools, secondary containment, spill kit.",
                "No unlabelled containers. No incompatible chemicals packed together. Dilutions match job method.",
            ),
            (
                "Cleaning equipment",
                "Mops, buckets, microfibre cloths, vacuum/extractor, floor tools, extension leads, hoses, machine accessories, batteries/chargers where applicable.",
                "Clean, inspected, fit for task, tagged if required, and loaded to prevent leaks or damage.",
            ),
            (
                "Site protection",
                "Wet-floor signs, cones, barrier tape, drop sheets, corner guards, door stops, masking/protection where authorised.",
                "Enough signs and barriers for public-facing or shared areas.",
            ),
            (
                "Admin records",
                "Job card, client contacts, risk assessment, toolbox talk sheet, incident form, client sign-off, photos approval note if used.",
                "Current version, legible, and available offline if mobile signal is unreliable.",
            ),
            (
                "Vehicle controls",
                "First aid kit, fire extinguisher where applicable, straps, load plan, emergency contacts, fuel/toll card, route plan.",
                "Vehicle pre-use check completed and defects escalated before departure.",
            ),
            (
                "Waste controls",
                "Bags, labels, ties, containers for wet/contaminated waste if authorised, client waste instructions.",
                "Waste route agreed. Special or hazardous waste is not handled unless trained and authorised.",
            ),
        ],
        [1800, 4860, 2700],
        header_fill=LIGHT_GREY,
        font_size=8.7,
    )

    # Pre-departure
    add_heading(doc, "6. Pre-Departure Procedure", 1)
    pre_steps = [
        "Review the job card, service level agreement, client access rules, site contact, arrival window, service scope, and special instructions.",
        "Confirm the crew is trained, medically fit for the task if required, properly briefed, and allocated according to the work plan.",
        "Check the weather, route, travel time, parking restrictions, loading-zone rules, and client access requirements.",
        "Pack PPE, signage, first aid, spill kit, method statements, SDS access, and the current SOP/job pack.",
        "Inspect and test cleaning equipment, power leads, chargers, hoses, attachments, and consumables before loading.",
        "Prepare chemicals using approved labels, compatible containers, dilution instructions, and secondary containment.",
        "Load the vehicle so chemicals, equipment, waste containers, and client property controls are segregated and restrained.",
        "Confirm client arrival details and any permit, induction, or security requirements before leaving base.",
        "Complete the mobile unit issue register, vehicle pre-use check, and supervisor release signature.",
        "Depart only after the supervisor confirms that the crew, vehicle, records, and equipment are ready for safe service.",
    ]
    add_step_table(doc, pre_steps)
    add_callout(
        doc,
        "Pre-departure hold point",
        "If a required control is missing, the supervisor must delay departure, substitute approved equipment, revise the job plan, or escalate to Operations.",
        fill=CAUTION_FILL,
    )

    # Setup
    add_heading(doc, "7. On-Site Setup Procedure", 1)
    setup_steps = [
        "Park safely in the agreed loading or service area without blocking emergency access, pedestrian routes, client operations, or public traffic.",
        "Report to the client site contact, confirm identity/access requirements, and obtain any site induction or permit instructions.",
        "Complete a site walk-through with the client or authorised representative before unloading where practical.",
        "Confirm the work area, boundaries, priority rooms/areas, excluded zones, access keys/cards, alarm rules, and client property controls.",
        "Identify changes from the job card, including occupancy, fragile items, electrical points, water sources, surfaces, slip risks, pets, public access, waste rules, or security restrictions.",
        "Position the staging area so chemicals, equipment, cables, hoses, waste bags, and clean/dirty tools are separated and protected from clients and the public.",
        "Place signs, cones, barriers, drop sheets, and floor protection before unpacking wet or powered equipment.",
        "Connect water and electricity only with client approval and using safe routing that prevents trip, shock, splash, or overload hazards.",
        "Prepare chemicals at the designated chemical station using the approved dilution chart and SDS controls.",
        "Conduct a toolbox talk covering hazards, PPE, work zones, communications, emergency arrangements, stop-work rules, and task sequence.",
        "Perform a surface test patch or equipment test where the service method, chemical, or surface condition requires it.",
        "Start cleaning only after the supervisor completes the Service Readiness Checklist and resolves any hold items.",
    ]
    add_step_table(doc, setup_steps)

    add_heading(doc, "7.1 Service Readiness Checklist", 2)
    add_table(
        doc,
        ["Readiness item", "Pass criteria", "Initial"],
        [
            ("Client check-in", "Site contact, work area, access, and exclusions confirmed.", "[ ]"),
            ("Risk changes", "New hazards recorded and controls agreed before work starts.", "[ ]"),
            ("Staging area", "Clean/dirty zones, chemical point, equipment zone, and waste point are separated.", "[ ]"),
            ("Public/client protection", "Signs, cones, barriers, drop sheets, and cable/hose routing are in place.", "[ ]"),
            ("Chemical control", "Labels, SDS access, dilution, PPE, and ventilation controls are confirmed.", "[ ]"),
            ("Equipment control", "Tools are inspected, safe to use, and matched to the service method.", "[ ]"),
            ("Waste route", "Waste collection, temporary storage, and disposal route agreed.", "[ ]"),
            ("Emergency arrangements", "First aid, emergency contacts, assembly point, and escalation route confirmed.", "[ ]"),
            ("Supervisor release", "Supervisor authorises cleaning to begin.", "[ ]"),
        ],
        [2800, 5360, 1200],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )

    # During service
    add_heading(doc, "8. During-Service Control", 1)
    add_table(
        doc,
        ["Control area", "Minimum standard"],
        [
            ("Communication", "The supervisor remains contactable and gives the client updates on delays, access limits, scope changes, incidents, or completion time changes."),
            ("Chemical use", "Only approved chemicals and dilutions may be used. Never mix chemicals unless the method statement specifically authorises it."),
            ("Electrical and water safety", "Keep leads and hoses routed away from walkways, standing water, sharp edges, heat, and unauthorised plugs or adapters."),
            ("Public interface", "Maintain signs and barriers. Do not leave wet floors, cables, chemicals, open buckets, or powered equipment unattended."),
            ("Property protection", "Move or protect client property only with authorisation. Record pre-existing damage before service begins where practical."),
            ("Quality control", "Supervisor checks high-risk or high-visibility areas during the job, not only at the end."),
            ("Scope control", "Additional work requires supervisor approval and client authorisation before the crew starts it."),
        ],
        [2200, 7160],
        header_fill=LIGHT_GREY,
        font_size=9,
    )

    # Pack down
    add_heading(doc, "9. Pack-Down Procedure", 1)
    pack_steps = [
        "Stop active cleaning in a controlled sequence so wet areas, chemicals, cables, and equipment are not abandoned.",
        "Complete a supervisor quality check against the job card, service areas, priority items, and client instructions.",
        "Collect loose tools, cloths, buckets, chemical containers, nozzles, accessories, batteries, leads, hoses, and signage from the work area.",
        "Close, wipe, and secure chemical containers. Return unused chemicals to secondary containment and record any spills, losses, or unusual use.",
        "Dispose of or bag waste according to the agreed route. Do not pour chemicals, slurry, or contaminated wash water into drains unless authorised.",
        "Clean equipment enough for safe transport and prevent leaking, odour, residue build-up, and cross-contamination between sites.",
        "Remove floor protection, barriers, and signs only when floors are dry, safe, and clear for handover.",
        "Restore furniture, keys, access cards, switches, doors, windows, alarms, and client property as agreed.",
        "Complete the final walk-through with the client representative where available.",
        "Record photographs only where authorised, avoid unnecessary personal information, and file images in the approved company location.",
        "Obtain client sign-off, record any snags or rework, and agree the follow-up owner and deadline.",
        "Load the vehicle safely, reconcile inventory, confirm no team member or item is left behind, and report departure to the site contact.",
    ]
    add_step_table(doc, pack_steps)
    add_callout(
        doc,
        "Pack-down hold point",
        "The team may not leave site until the supervisor confirms the area is safe, the client has been offered handover, waste is controlled, and the vehicle load is secure.",
        fill=CAUTION_FILL,
    )

    # Return to base
    add_heading(doc, "10. Return-to-Base and Vehicle Reset", 1)
    add_table(
        doc,
        ["Reset item", "Required action", "Done"],
        [
            ("Vehicle", "Refuel/recharge if required, remove waste, clean load area, lock vehicle, report defects.", "[ ]"),
            ("Equipment", "Clean, dry, inspect, charge, tag defective items, and return to assigned storage.", "[ ]"),
            ("Chemicals", "Return stock, segregate incompatible products, quarantine unknown or damaged containers.", "[ ]"),
            ("PPE and laundry", "Bag soiled PPE/cloths, send laundry, replace damaged PPE, record shortages.", "[ ]"),
            ("Records", "Submit job card, checklist, client sign-off, photos, incident forms, and waste evidence.", "[ ]"),
            ("Stock", "Update consumable use, breakages, losses, and re-order triggers.", "[ ]"),
            ("Debrief", "Record client feedback, near misses, rework, training needs, and improvement actions.", "[ ]"),
        ],
        [1900, 6260, 1200],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )

    # Incidents
    add_heading(doc, "11. Incidents, Spills, Damage, and Non-Conformance", 1)
    add_para(
        doc,
        "The first response is to make the area safe, protect people, preserve important evidence where appropriate, and notify the right person. The supervisor must decide whether work can continue, must pause, or must be escalated.",
    )
    add_table(
        doc,
        ["Trigger", "Immediate response", "Record / escalation"],
        [
            ("Injury, exposure, or near miss", "Stop work, give first aid, isolate hazard, call emergency support if needed.", "Incident report; notify Operations/SHEQ; follow company COIDA process."),
            ("Chemical spill or leak", "Keep people away, use spill kit and SDS controls, ventilate where safe.", "Spill report; photograph only if safe and authorised; quarantine container."),
            ("Client property damage", "Stop work in the affected area, protect item, notify client contact and Operations.", "Damage report with description, photos if authorised, witness names, and corrective action."),
            ("Lost key, access card, or code", "Notify supervisor immediately and follow client security instruction.", "Security incident report; restrict access to personal information."),
            ("Missing or defective equipment", "Remove from service, tag if defective, substitute approved equipment only.", "Equipment defect/loss report and stores update."),
            ("Waste disposal concern", "Do not dispose. Secure waste and obtain instruction from client/Operations.", "Waste exception report and receipt/evidence where applicable."),
            ("Scope change or unsafe request", "Pause task and obtain supervisor approval before starting extra work.", "Variation note, revised risk assessment, client approval."),
        ],
        [2000, 4200, 3160],
        header_fill=LIGHT_GREY,
        font_size=8.5,
    )
    add_heading(doc, "11.1 Emergency Contacts", 2)
    add_table(
        doc,
        ["Emergency / escalation contact", "Name", "Number", "After-hours instruction"],
        [
            ("Operations Manager", "[Name]", "[Number]", "[Instruction]"),
            ("SHEQ / First Aid Lead", "[Name]", "[Number]", "[Instruction]"),
            ("Client site contact", "[Name]", "[Number]", "[Instruction]"),
            ("Security / building manager", "[Name]", "[Number]", "[Instruction]"),
            ("Emergency services", "[Local emergency number]", "[Number]", "[Instruction]"),
        ],
        [2600, 2100, 1900, 2760],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )

    # Records
    add_heading(doc, "12. Records and Retention", 1)
    add_table(
        doc,
        ["Record", "Owner", "Where filed", "Minimum retention"],
        [
            ("Approved SOP and revision history", "Operations / SHEQ", "[Controlled document folder]", "[Company schedule]"),
            ("Job card and client sign-off", "Mobile Supervisor", "[Client/job file]", "[Company schedule]"),
            ("Vehicle and equipment checks", "Driver / Stores", "[Fleet/equipment register]", "[Company schedule]"),
            ("Toolbox talk and risk assessment", "Mobile Supervisor", "[SHEQ/job file]", "[Company schedule]"),
            ("Chemical issue and dilution logs", "Stores / Supervisor", "[Chemical control file]", "[Company schedule]"),
            ("Incident, near-miss, and damage reports", "SHEQ / Operations", "[Incident register]", "[Company/legal schedule]"),
            ("Waste receipts or disposal evidence", "Supervisor / Admin", "[Waste/client file]", "[Company/client schedule]"),
            ("Training and competency records", "HR / Training", "[Training file]", "[Company/legal schedule]"),
        ],
        [2600, 2000, 2760, 2000],
        header_fill=LIGHT_GREY,
        font_size=8.5,
    )

    # Appendices
    add_page_break(doc)
    add_heading(doc, "Appendix A: Mobile Service Setup Checklist", 1)
    add_table(
        doc,
        ["Item", "Requirement", "Done", "Initial"],
        [
            ("Job card", "Scope, site, arrival time, contact, exclusions, and special instructions confirmed.", "[ ]", ""),
            ("Crew", "Names, roles, competency, PPE sizes, and attendance confirmed.", "[ ]", ""),
            ("Vehicle", "Pre-use check, fuel, load area, emergency kit, and route completed.", "[ ]", ""),
            ("PPE", "All required PPE issued and inspected.", "[ ]", ""),
            ("Chemicals", "Labels, SDS access, dilution chart, secondary containment, and spill kit confirmed.", "[ ]", ""),
            ("Equipment", "Machines, tools, leads, hoses, accessories, batteries, and consumables checked.", "[ ]", ""),
            ("Records", "Risk assessment, toolbox talk, client sign-off, incident form, and registers packed.", "[ ]", ""),
            ("Arrival", "Client/site contact, access, induction, parking, and work boundaries confirmed.", "[ ]", ""),
            ("Staging", "Clean/dirty zones, chemical point, waste point, and equipment zone established.", "[ ]", ""),
            ("Protection", "Signs, barriers, drop sheets, cable/hose routing, and floor protection in place.", "[ ]", ""),
            ("Toolbox talk", "Hazards, PPE, emergency arrangements, work sequence, and stop-work rules briefed.", "[ ]", ""),
            ("Release", "Supervisor authorises service to start.", "[ ]", ""),
        ],
        [1400, 6040, 900, 1020],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )

    add_heading(doc, "Appendix B: Pack-Down Checklist", 1)
    add_table(
        doc,
        ["Item", "Requirement", "Done", "Initial"],
        [
            ("Quality check", "Supervisor checked job areas against service scope and client priorities.", "[ ]", ""),
            ("Snags", "Rework, exclusions, and follow-ups recorded with owner and due date.", "[ ]", ""),
            ("Chemicals", "Containers closed, wiped, labelled, and secured for transport.", "[ ]", ""),
            ("Waste", "Waste bagged, labelled where needed, and routed as agreed.", "[ ]", ""),
            ("Equipment", "Tools cleaned, inspected, counted, and loaded safely.", "[ ]", ""),
            ("Site restored", "Furniture, doors, windows, keys/cards, lights, alarms, and client property restored.", "[ ]", ""),
            ("Area safe", "Floors dry/safe, signs removed last, and no trip/slip hazards remain.", "[ ]", ""),
            ("Client sign-off", "Client offered final walk-through and sign-off captured where available.", "[ ]", ""),
            ("Vehicle load", "Load restrained, no leaks, no blocked exits, crew accounted for.", "[ ]", ""),
            ("Departure", "Site contact notified and departure time recorded.", "[ ]", ""),
            ("Return to base", "Records submitted, stock updated, defects tagged, and vehicle reset.", "[ ]", ""),
        ],
        [1400, 6040, 900, 1020],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix C: Mobile Unit Inventory Register", 1)
    add_table(
        doc,
        ["Date", "Item", "Qty out", "Condition out", "Qty returned", "Condition returned", "Defect/loss action"],
        [
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
        ],
        [1150, 2000, 900, 1400, 1050, 1500, 1360],
        header_fill=LIGHT_GREY,
        font_size=8.2,
    )

    add_heading(doc, "Appendix D: Chemical Control and Dilution Log", 1)
    add_table(
        doc,
        ["Date", "Product", "Task / area", "Dilution", "SDS checked", "Issued by", "Returned / used", "Notes"],
        [
            ("", "", "", "", "[ ]", "", "", ""),
            ("", "", "", "", "[ ]", "", "", ""),
            ("", "", "", "", "[ ]", "", "", ""),
            ("", "", "", "", "[ ]", "", "", ""),
            ("", "", "", "", "[ ]", "", "", ""),
            ("", "", "", "", "[ ]", "", "", ""),
            ("", "", "", "", "[ ]", "", "", ""),
            ("", "", "", "", "[ ]", "", "", ""),
        ],
        [900, 1500, 1400, 900, 900, 1100, 1200, 1460],
        header_fill=LIGHT_GREY,
        font_size=8,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix E: Site Risk Assessment and Toolbox Talk", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Client / site", "[Client name and address]"),
            ("Job number / date", "[Job number] / [dd/mm/yyyy]"),
            ("Supervisor", "[Name]"),
            ("Crew members", "[Names]"),
            ("Service scope", "[Scope summary]"),
            ("Main hazards found on arrival", "[Slips/trips, chemicals, electrical, public access, fragile surfaces, waste, security, weather, manual handling]"),
            ("Controls agreed", "[PPE, barriers, signs, ventilation, isolation, sequencing, client support, permits]"),
            ("Emergency arrangements", "[First aid, assembly point, site contact, emergency number, spill response]"),
            ("Stop-work triggers", "[Site-specific triggers]"),
            ("Crew sign-off", "[Names/signatures]"),
        ],
        [2500, 6860],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )
    add_table(
        doc,
        ["Hazard", "Who may be harmed", "Controls", "Residual risk", "Action owner"],
        [
            ("", "", "", "[Low / Med / High]", ""),
            ("", "", "", "[Low / Med / High]", ""),
            ("", "", "", "[Low / Med / High]", ""),
            ("", "", "", "[Low / Med / High]", ""),
            ("", "", "", "[Low / Med / High]", ""),
        ],
        [1700, 1900, 3260, 1200, 1300],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix F: Client Service Completion Sign-Off", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Client / site", "[Client name and address]"),
            ("Job number", "[Job number]"),
            ("Service date and time", "[Start] to [Finish]"),
            ("Areas completed", "[Areas / rooms / assets completed]"),
            ("Exclusions / access limits", "[Areas not completed and reason]"),
            ("Snags or rework", "[Details, owner, due date]"),
            ("Waste handled", "[General / recyclable / special / client retained / contractor removed]"),
            ("Photos authorised", "[Yes / No / Not applicable]"),
            ("Client representative", "[Name, role, signature, date]"),
            ("Supervisor", "[Name, signature, date]"),
        ],
        [2500, 6860],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )

    add_heading(doc, "Appendix G: Incident and Non-Conformance Report", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Report type", "[Injury / near miss / chemical spill / property damage / complaint / non-conformance / other]"),
            ("Date, time, site", "[Details]"),
            ("Reported by", "[Name and role]"),
            ("People involved", "[Names / roles / contact details where appropriate]"),
            ("Description", "[What happened, where, and what was affected]"),
            ("Immediate action taken", "[Make safe, first aid, isolation, client notice, emergency call, spill control]"),
            ("Photos or evidence", "[Authorised? Where filed?]"),
            ("Root cause / contributing factors", "[Initial assessment]"),
            ("Corrective action", "[Action, owner, due date]"),
            ("Closed by", "[Name, role, date]"),
        ],
        [2500, 6860],
        header_fill=LIGHT_GREY,
        font_size=8.8,
    )

    # Preset audit marker in document properties/comment-like final note.
    add_para(
        doc,
        "End of template. Adapt this SOP after consultation with the business owner, supervisor, SHEQ representative, client contract, insurer, and competent compliance adviser.",
        italic=True,
        color=MUTED,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=8,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT.resolve())
