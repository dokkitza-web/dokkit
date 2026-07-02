from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Risk_Assessment_and_Incident_Report_Template.docx")

FONT = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
RISK_RED = RGBColor(0x9B, 0x1C, 0x1C)
CAUTION_GOLD = RGBColor(0x7A, 0x5A, 0x00)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
RISK_FILL = "FBEAEA"
BORDER = "C9D3DF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn("w:ascii"), name)
    r_pr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, before=0, after=6, line=1.25, keep_next=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next is not None:
        fmt.keep_with_next = keep_next


def style_para(p, size=11, color=BODY, bold=False, italic=False):
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_para(
    doc,
    text="",
    *,
    style=None,
    size=11,
    color=BODY,
    bold=False,
    italic=False,
    align=None,
    before=0,
    after=6,
    line=1.25,
    keep_next=None,
):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=before, after=after, line=line, keep_next=keep_next)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    style_para(
        p,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **margins):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ["top", "start", "bottom", "end"]:
        if m not in margins:
            continue
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[m]))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, size=9.5, color=BODY, align=None, fill=None, italic=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    if text is not None:
        for idx, part in enumerate(str(text).split("\n")):
            if idx:
                p.add_run().add_break()
            run = p.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=0, after=0, line=1.15)
    if align is not None:
        p.alignment = align


def set_table_borders(table, color=BORDER, sz="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell, **CELL_MARGIN)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            size=font_size,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=header_fill,
        )
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, text in enumerate(row_data):
            short = len(str(text)) < 18
            align = WD_ALIGN_PARAGRAPH.CENTER if short and idx in (0, len(row_data) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], text, size=font_size, align=align)
    add_para(doc, "", after=3)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT_FILL, title_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D7DBE2", sz="6")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.text = ""
    title_run = p.add_run(title.upper() + "\n")
    set_run_font(title_run, size=10, color=title_color, bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=10.5, color=BODY)
    set_paragraph_format(p, before=0, after=0, line=1.2)
    add_para(doc, "", after=2)
    return table


def add_step_table(doc, steps, *, font_size=8.8):
    rows = [(str(idx), step) for idx, step in enumerate(steps, start=1)]
    return add_table(
        doc,
        ["Step", "Action"],
        rows,
        [850, 8510],
        header_fill=LIGHT_BLUE,
        font_size=font_size,
    )


def build_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True


def setup_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = ""
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    left, right = table.rows[0].cells
    set_cell_text(left, "SHEQ | Risk Assessment and Incident Report", size=8.5, color=MUTED)
    set_cell_text(right, "[Business Name]", size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    footer = section.footer
    footer.paragraphs[0].text = ""
    ftable = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ftable, [6500, 2860], indent_dxa=0)
    lcell, rcell = ftable.rows[0].cells
    set_cell_text(lcell, "Controlled document - verify latest version before use", size=8.5, color=MUTED)
    p = rcell.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" of ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "NUMPAGES")
    set_paragraph_format(p, before=0, after=0, line=1.0)
    for cell in ftable.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)


def add_cover(doc):
    add_para(doc, "", after=48)
    add_para(
        doc,
        "Business Template",
        size=10.5,
        bold=True,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
    )
    add_para(
        doc,
        "Risk Assessment and Incident Report",
        size=29,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.08,
    )
    add_para(
        doc,
        "For a South African Cleaning Services Business",
        size=15,
        color=DARK_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=24,
    )
    add_para(
        doc,
        "Prepared for: [Business Name] | Template date: June 2026",
        size=11.5,
        bold=True,
        color=BODY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=5,
    )
    add_para(
        doc,
        "Use this controlled template to identify cleaning-service hazards, rate and control risk, report incidents, capture evidence, assign corrective actions, and keep a reliable SHEQ record trail.",
        size=10.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=28,
    )
    add_table(
        doc,
        ["Field", "Template Entry"],
        [
            ("Document ID", "[SHEQ-RAIR-###]"),
            ("Version", "[0.1 / Draft / Approved]"),
            ("Effective date", "[dd/mm/yyyy]"),
            ("Review date", "[dd/mm/yyyy]"),
            ("Document owner", "[SHEQ Officer / Operations Manager]"),
            ("Approved by", "[Owner / Director / Accountable Manager]"),
            ("Applies to", "Cleaners, supervisors, drivers, subcontractors, temporary staff, and managers"),
            ("Service area", "[Province / Municipality / Client sites]"),
        ],
        [2500, 6860],
        font_size=9.5,
    )
    add_callout(
        doc,
        "Template control",
        "Replace bracketed fields before issue. Adapt the controls to the actual cleaning methods, client sites, products, vehicles, insurance conditions, and competent legal/SHEQ advice.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )
    add_page_break(doc)


def add_document_control(doc):
    add_heading(doc, "1. Document Control", 1)
    add_table(
        doc,
        ["Version", "Date", "Changed by", "Summary of change", "Approved by"],
        [
            ("0.1", "[dd/mm/yyyy]", "[Name]", "Initial template issue", "[Name]"),
            ("", "", "", "", ""),
            ("", "", "", "", ""),
        ],
        [1050, 1400, 1700, 3610, 1600],
        header_fill=LIGHT_GREY,
        font_size=9,
    )
    add_table(
        doc,
        ["Controlled copy holder", "Role / area", "Issue method", "Review trigger"],
        [
            ("[Name]", "Operations", "[Shared drive / printed controlled copy]", "New service, new chemical, incident trend, legal change"),
            ("[Name]", "SHEQ / HR", "[Training file / onboarding pack]", "Training gap, injury, exposure, near miss"),
            ("[Name]", "Supervisors", "[Crew pack / mobile app]", "Client-site change or method change"),
            ("[Name]", "Client service desk", "[Admin folder]", "Complaint, damage claim, privacy concern"),
        ],
        [2300, 1800, 2600, 2660],
        header_fill=LIGHT_GREY,
        font_size=8.8,
    )


def add_purpose_scope(doc):
    add_heading(doc, "2. Purpose, Scope, and Use", 1)
    add_para(
        doc,
        "This template provides a practical risk-assessment and incident-reporting system for cleaning services. It is designed for residential, commercial, mobile, and once-off cleaning work, including work at client premises where site conditions can change without notice.",
    )
    add_table(
        doc,
        ["Use case", "How the template should be used"],
        [
            ("Pre-job planning", "Identify predictable hazards before accepting or scheduling work, especially chemicals, access, electricity, manual handling, public interface, and waste."),
            ("On-arrival assessment", "Check the actual site before unloading or starting work. Record new hazards and controls on the task risk assessment."),
            ("Toolbox talk", "Brief the team on the job method, hazards, PPE, emergency arrangements, stop-work triggers, and supervisor release point."),
            ("Incident response", "Make the area safe, give first aid or emergency support, preserve evidence where appropriate, and report through the correct route."),
            ("Investigation and learning", "Capture facts, root cause, corrective actions, owners, due dates, close-out evidence, and training updates."),
        ],
        [2200, 7160],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )
    add_callout(
        doc,
        "Stop-work rule",
        "Any employee may pause work where there is an uncontrolled risk to people, client property, the public, the environment, company equipment, or personal information. The supervisor must reassess before work continues.",
        fill=RISK_FILL,
        title_color=RISK_RED,
    )


def add_compliance(doc):
    add_heading(doc, "3. South African Compliance Reference Points", 1)
    add_para(
        doc,
        "This section is an operational reference, not legal advice. The business should verify reporting duties with a competent adviser, the Department of Employment and Labour, the Compensation Fund, the client contract, insurer, and local municipal requirements.",
        italic=True,
        color=MUTED,
        size=10.5,
    )
    add_table(
        doc,
        ["Reference point", "Operational meaning for cleaning work"],
        [
            (
                "Occupational Health and Safety Act 85 of 1993",
                "Provide and maintain, as far as reasonably practicable, a work environment that is safe and without risk to health. Assess cleaning hazards before work starts and control risks arising from the work.",
            ),
            (
                "OHS Act incident reporting requirements",
                "Escalate serious incidents immediately for determination of inspector/client/insurer reporting duties, including death, major injury, unconsciousness, permanent physical defect risk, extended inability to work, major incident, or dangerous substance release.",
            ),
            (
                "Compensation for Occupational Injuries and Diseases Act / Compensation Fund",
                "Report injury-on-duty claims within the required timeframes. Department guidance states injury claims within 7 days of notice and occupational disease claims within 14 days of notice.",
            ),
            (
                "Regulations for Hazardous Chemical Agents, 2021",
                "Keep chemical labels, GHS-aligned safety data sheets, training, exposure controls, PPE, ventilation, and emergency controls aligned to the products used.",
            ),
            (
                "Protection of Personal Information Act 4 of 2013",
                "Limit incident photos, ID numbers, medical details, addresses, client access codes, witness contact details, and staff records to what is necessary and keep them secure.",
            ),
            (
                "National Environmental Management: Waste Act 59 of 2008 and municipal by-laws",
                "Prevent pollution, control cleaning waste, chemical residues, contaminated absorbents, sharps, broken glass, and unusual waste through approved routes.",
            ),
        ],
        [3000, 6360],
        header_fill=LIGHT_GREY,
        font_size=8.2,
    )
    add_table(
        doc,
        ["Official verification source", "Link / note"],
        [
            ("Department of Employment and Labour", "https://www.labour.gov.za/"),
            ("OHS Act on gov.za", "https://www.gov.za/documents/occupational-health-and-safety-act"),
            ("Compensation Fund employer obligations", "https://www.labour.gov.za/compensation-fund-obligations-of-the-employer"),
            ("POPIA on gov.za", "https://www.gov.za/documents/protection-personal-information-act"),
            ("Waste Act on gov.za", "https://www.gov.za/documents/national-environmental-management-waste-act"),
        ],
        [3000, 6360],
        header_fill=LIGHT_GREY,
        font_size=8.1,
    )


def add_roles(doc):
    add_heading(doc, "4. Roles and Responsibilities", 1)
    add_table(
        doc,
        ["Role", "Risk-assessment responsibility", "Incident-reporting responsibility"],
        [
            ("Owner / Director", "Approve the system, appoint competent persons, resource PPE/training, and review high-risk trends.", "Ensure statutory/client/insurer reporting routes are followed and serious events are closed at management level."),
            ("SHEQ Officer / Competent Person", "Maintain risk matrix, task assessments, chemical controls, training records, and audit schedule.", "Classify incidents, advise on reportability, lead investigations, track corrective actions, and retain records."),
            ("Operations Manager", "Confirm job scope, staffing, equipment, vehicle, client rules, and escalation contacts before work is released.", "Notify client/insurer/management where required and ensure operational corrective actions are completed."),
            ("Supervisor / Team Leader", "Complete site/task risk assessment, toolbox talk, PPE check, stop-work decisions, and release to start.", "Make area safe, arrange first aid, protect evidence, take authorised photos, collect witness details, and submit report."),
            ("Cleaner / Driver / Crew Member", "Follow controls, use PPE, report hazards, avoid unauthorised chemicals/tasks, and stop unsafe work.", "Report injuries, exposure, near misses, spills, damage, complaints, lost keys, or security concerns immediately."),
            ("Admin / HR", "File controlled records and training evidence according to company schedule.", "Support COIDA claim documents, medical certificates, personal-information safeguards, and close-out filing."),
        ],
        [1900, 3850, 3610],
        header_fill=LIGHT_BLUE,
        font_size=8.2,
    )


def add_risk_method(doc):
    add_heading(doc, "5. Risk-Assessment Method", 1)
    add_step_table(
        doc,
        [
            "Define the task, location, client area, cleaning method, team, equipment, chemicals, timing, and public/client interface.",
            "Identify hazards before arrival using the job card, client information, product SDS, past incidents, and supervisor experience.",
            "Recheck hazards on arrival before unloading, mixing chemicals, connecting equipment, moving furniture, or entering restricted areas.",
            "List who may be harmed, including staff, client employees, residents, children, vulnerable persons, pets, contractors, and members of the public.",
            "Rate initial risk by combining likelihood and severity before controls are applied.",
            "Select controls using elimination, substitution, isolation/barriers, engineering controls, administrative controls, training, signage, and PPE.",
            "Rate residual risk after controls. Do not start high residual-risk work without manager/SHEQ approval and additional controls.",
            "Brief the team, record names/signatures, assign open actions, and obtain supervisor release to start.",
            "Monitor controls during work, especially wet floors, chemical use, cables, access routes, equipment defects, and changing client activity.",
            "Review the assessment after any incident, near miss, client complaint, chemical change, equipment change, or repeated non-conformance.",
        ],
        font_size=8.6,
    )
    add_heading(doc, "5.1 Risk Rating Matrix", 2)
    add_table(
        doc,
        ["Likelihood", "Description", "Score"],
        [
            ("Rare", "Not expected during normal cleaning work but possible in unusual conditions.", "1"),
            ("Unlikely", "Could happen, but controls and experience make it uncommon.", "2"),
            ("Possible", "Might occur during the job or has happened before in similar work.", "3"),
            ("Likely", "Expected to occur if controls are weak or conditions change.", "4"),
            ("Almost certain", "Expected repeatedly or already happening on the job.", "5"),
        ],
        [1700, 6360, 1300],
        header_fill=LIGHT_GREY,
        font_size=8.5,
    )
    add_table(
        doc,
        ["Severity", "Description", "Score"],
        [
            ("Minor", "First-aid only, minor spill, no lasting harm, no service disruption.", "1"),
            ("Moderate", "Medical treatment, minor property damage, complaint, or short work interruption.", "2"),
            ("Serious", "Lost-time injury risk, significant chemical exposure, major damage, or client escalation.", "3"),
            ("Major", "Permanent harm risk, hospitalisation, major fire/flood/chemical event, or major client loss.", "4"),
            ("Catastrophic", "Fatality, multiple serious injuries, major environmental harm, or severe legal exposure.", "5"),
        ],
        [1700, 6360, 1300],
        header_fill=LIGHT_GREY,
        font_size=8.5,
    )
    add_table(
        doc,
        ["Risk score", "Risk level", "Required decision"],
        [
            ("1-4", "Low", "Proceed with normal controls and supervisor monitoring."),
            ("5-9", "Medium", "Proceed only after controls are confirmed and briefed."),
            ("10-16", "High", "Supervisor and Operations/SHEQ approval required before work starts or continues."),
            ("17-25", "Critical", "Do not start or continue. Stop work and escalate for redesign, specialist support, or refusal of unsafe scope."),
        ],
        [1500, 1800, 6060],
        header_fill=LIGHT_BLUE,
        font_size=8.6,
    )


def add_hazard_register(doc):
    add_heading(doc, "6. Cleaning-Service Hazard Register", 1)
    add_table(
        doc,
        ["Hazard group", "Typical cleaning examples", "Minimum controls"],
        [
            ("Slips, trips, falls", "Wet floors, loose mats, hoses, leads, stairs, poor lighting, cluttered homes/sites.", "Signs, barriers, dry-route planning, cable/hose routing, suitable footwear, good housekeeping, no unattended wet stairs."),
            ("Chemicals and exposure", "Disinfectants, degreasers, acids/alkalis, aerosols, odours, incorrect dilution, product mixing.", "Approved products, labels, SDS access, dilution chart, ventilation, gloves/eye protection, no unauthorised mixing, spill kit."),
            ("Electrical", "Vacuum leads, extension cords, wet areas, damaged plugs, overloaded sockets, client appliances.", "Pre-use inspection, dry connections, RCD where required, safe routing, no damaged equipment, client permission for power points."),
            ("Manual handling", "Moving furniture, lifting machines, carrying water buckets, loading vehicles, stairs.", "Team lift, trolley, weight limits, route check, avoid twisting, training, supervisor approval for heavy items."),
            ("Biological / hygiene", "Bodily fluids, sharps, pest waste, mould, spoiled food, toilets, contaminated cloths.", "Stop-work for excluded biohazards, gloves, segregation, colour-coded cloths, hand hygiene, sealed waste, specialist scope if needed."),
            ("Client property", "Fragile surfaces, valuables, electronics, antiques, keys, access cards, alarms.", "Pre-start walk-through, photos if authorised, protect/move only with permission, secure access records, damage reporting."),
            ("Vehicles and travel", "Route risk, loading, reversing, parking, public access, fuel/charging, fatigue.", "Driver checks, load restraint, parking plan, journey plan, no blocked exits, defect reporting."),
            ("Public/security", "Clients at home, office staff, children, pets, aggressive behaviour, restricted areas.", "Work-zone control, client communication, no childcare/pet handling, security escalation, team check-in/out."),
            ("Waste and environment", "Chemical residue, contaminated absorbents, broken glass, sharps, wash water, client waste.", "Segregate waste, use approved disposal route, do not pour unknown chemicals, label and escalate unusual waste."),
        ],
        [1800, 3660, 3900],
        header_fill=LIGHT_GREY,
        font_size=7.7,
    )
    add_callout(
        doc,
        "Chemical mixing warning",
        "Do not mix cleaning chemicals unless the approved method statement specifically permits it. Treat bleach, ammonia, acids, degreasers, disinfectants, and unknown client products as incompatible until verified.",
        fill=RISK_FILL,
        title_color=RISK_RED,
    )


def add_controls(doc):
    add_heading(doc, "7. Required Control Checks", 1)
    add_table(
        doc,
        ["Control area", "Minimum check before work starts", "Done"],
        [
            ("Scope", "Job card matches the client request and excludes specialist, high-access, biohazard, pest, mould, or hazardous waste work unless approved.", "[ ]"),
            ("Competency", "Crew is trained for the service method, equipment, chemicals, PPE, and incident process.", "[ ]"),
            ("PPE", "Gloves, eye protection, footwear, masks/respirators where required, aprons, and high-visibility items are suitable and available.", "[ ]"),
            ("SDS and labels", "Every decanted product is labelled and SDS/product safety information is accessible to the supervisor.", "[ ]"),
            ("Equipment", "Machines, leads, hoses, buckets, ladders/steps, signs, barriers, and spill kit are inspected and safe.", "[ ]"),
            ("Client/site rules", "Access, alarms, keys, parking, waste, toilets, water, electricity, photos, and privacy limits are confirmed.", "[ ]"),
            ("Emergency", "First aid, emergency contacts, assembly point, nearest clinic/hospital route, and supervisor escalation are known.", "[ ]"),
            ("Release", "Supervisor confirms controls are in place and open high-risk items are escalated before work starts.", "[ ]"),
        ],
        [1900, 6260, 1200],
        header_fill=LIGHT_BLUE,
        font_size=8.3,
    )
    add_heading(doc, "7.1 PPE and Equipment Quick Guide", 2)
    add_table(
        doc,
        ["Task / exposure", "Minimum PPE / equipment", "Extra approval trigger"],
        [
            ("General domestic/commercial cleaning", "Closed non-slip shoes, suitable gloves, company uniform, cleaning signs for wet areas.", "Unusual client chemicals, aggressive clients, biohazards, or restricted areas."),
            ("Chemical dilution or decanting", "Chemical-resistant gloves, eye protection, labelled containers, ventilation, spill kit.", "Unknown product, damaged container, missing SDS, strong fumes, or product mixing risk."),
            ("Bathroom/toilet cleaning", "Gloves, eye protection if splash risk, colour-coded cloths, hand hygiene controls.", "Bodily fluids, sharps, blocked drains, sewage, or heavy contamination."),
            ("Machine/floor cleaning", "Signs/barriers, safe leads, footwear, hearing protection if needed, equipment pre-check.", "Damaged plugs/leads, wet electrical points, public access, or unfamiliar machine."),
            ("Waste handling", "Gloves, closed shoes, bags/containers, sharps caution, hand hygiene.", "Sharps, chemical residue, broken glass, medical waste, dead pests, or unknown liquid."),
        ],
        [2300, 4360, 2700],
        header_fill=LIGHT_GREY,
        font_size=8.0,
    )


def add_incident_process(doc):
    add_heading(doc, "8. Incident Response and Escalation", 1)
    add_step_table(
        doc,
        [
            "Stop the task and make the area safe if this can be done without creating further risk.",
            "Check people first. Provide first aid within training limits and call emergency services where needed.",
            "Isolate the hazard: switch off equipment, close containers, cordon wet floors, ventilate, or keep people away.",
            "Notify the supervisor immediately. The supervisor notifies Operations/SHEQ and the client contact where appropriate.",
            "Do not admit liability or argue with the client. Record facts, actions taken, and who was notified.",
            "Preserve evidence where appropriate, including product labels, equipment, photos if authorised, CCTV details, witness names, and damaged items.",
            "Decide whether work may continue, must be paused, or must be stopped pending investigation.",
            "Complete the incident report before end of shift where possible and submit supporting documents.",
            "SHEQ/Management classifies reportability, including OHS Act, COIDA, client, insurer, and municipal/environmental requirements.",
            "Assign corrective actions, verify close-out evidence, update risk assessments/training, and brief affected teams.",
        ],
        font_size=8.6,
    )
    add_table(
        doc,
        ["Incident type", "Examples", "Immediate escalation"],
        [
            ("Injury / illness / exposure", "Cut, slip, fall, back strain, eye splash, inhalation, skin burn, suspected occupational disease.", "First aid/emergency support, supervisor, SHEQ, Operations, HR/COIDA owner."),
            ("Near miss", "Almost slipped on wet floor, chemical almost mixed, lead almost caused trip, object almost fell.", "Supervisor and SHEQ for risk review."),
            ("Chemical spill / release", "Leaking container, product splash, strong fumes, unknown reaction, spill into drain.", "Keep people away, SDS controls, SHEQ/Operations, client contact where on site."),
            ("Client property damage", "Broken item, stained surface, scratched floor, water damage, lost key/access card.", "Supervisor, Operations, client contact, insurer/admin where required."),
            ("Security / behaviour", "Threatening client, theft allegation, lost access code, unauthorised entry, harassment.", "Supervisor, Operations, client management/security, HR where staff affected."),
            ("Environmental / waste", "Improper disposal, contaminated absorbent, wash-water concern, unusual waste.", "Supervisor, SHEQ, client contact, municipal/specialist route if required."),
            ("Complaint / non-conformance", "Missed scope, poor quality, late arrival, privacy complaint, unauthorised photo.", "Supervisor, Operations, client service desk."),
        ],
        [1800, 4660, 2900],
        header_fill=LIGHT_GREY,
        font_size=8.0,
    )


def add_investigation(doc):
    add_heading(doc, "9. Investigation, Corrective Action, and Close-Out", 1)
    add_para(
        doc,
        "Investigations should establish facts and prevent recurrence. They should not be used to blame workers for reporting hazards or incidents in good faith.",
    )
    add_table(
        doc,
        ["Investigation stage", "Required output"],
        [
            ("Fact capture", "Date, time, location, people involved, task, equipment, chemicals, weather/site conditions, photos if authorised, witness details, and immediate actions."),
            ("Classification", "Injury, near miss, exposure, property damage, spill, complaint, non-conformance, privacy/security, environmental, or other."),
            ("Root cause review", "Consider unsafe condition, unsafe act, missing control, training gap, supervision, equipment, product, client condition, schedule pressure, or procedure gap."),
            ("Corrective action", "Action must have an owner, due date, risk priority, evidence required, and person responsible for verification."),
            ("Communication", "Brief affected teams and update toolbox talk content, risk register, job card, client instructions, or chemical register where needed."),
            ("Close-out", "SHEQ/Operations confirms action completion, files evidence, updates trend register, and records final status."),
        ],
        [2300, 7060],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )
    add_table(
        doc,
        ["Close-out quality check", "Yes / No", "Notes"],
        [
            ("Immediate cause and root/contributing causes recorded.", "[ ]", ""),
            ("Reportability checked against COIDA/OHS/client/insurer requirements.", "[ ]", ""),
            ("Corrective actions are realistic, assigned, dated, and verified.", "[ ]", ""),
            ("Affected risk assessment, SOP, SDS, training, or client instruction updated.", "[ ]", ""),
            ("Personal information and photos stored only in approved locations.", "[ ]", ""),
            ("Learning shared with supervisors/crews without exposing unnecessary personal details.", "[ ]", ""),
        ],
        [5960, 1200, 2200],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )


def add_records(doc):
    add_heading(doc, "10. Records and Retention", 1)
    add_table(
        doc,
        ["Record", "Owner", "Where filed", "Minimum retention"],
        [
            ("Approved template and revision history", "SHEQ / Operations", "[Controlled document folder]", "[Company schedule]"),
            ("Task/site risk assessments and toolbox talks", "Supervisor", "[SHEQ/job file]", "[Company/client schedule]"),
            ("Chemical register, SDS, dilution records", "SHEQ / Stores", "[Chemical control file]", "[Company/legal schedule]"),
            ("PPE and equipment inspections", "Supervisor / Stores", "[Equipment register]", "[Company schedule]"),
            ("Incident reports and evidence", "SHEQ / Operations", "[Incident register]", "[Company/legal schedule]"),
            ("COIDA claim documents and medical certificates", "HR / COIDA owner", "[Confidential employee file]", "[Legal/company schedule]"),
            ("Corrective action tracker and close-out evidence", "SHEQ / Operations", "[SHEQ action tracker]", "[Company schedule]"),
            ("Client notifications, complaints, and damage claims", "Operations / Admin", "[Client/job file]", "[Company/client/insurer schedule]"),
        ],
        [2700, 1900, 2760, 2000],
        header_fill=LIGHT_GREY,
        font_size=8.2,
    )
    add_callout(
        doc,
        "Privacy and evidence",
        "Incident evidence may include sensitive personal information. Limit access, avoid unnecessary medical or identity details, store photos securely, and do not share incident images on personal devices or social media.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )


def add_appendices(doc):
    add_page_break(doc)
    add_heading(doc, "Appendix A: Site / Task Risk Assessment Form", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Client / site / address", "[Client name, site, address, unit/floor/area]"),
            ("Job number / date / time", "[Job number] / [dd/mm/yyyy] / [start-finish]"),
            ("Service type", "[Residential / commercial / deep clean / floor care / move-out / other]"),
            ("Supervisor / competent person", "[Name and contact]"),
            ("Crew members", "[Names and roles]"),
            ("Chemicals / equipment", "[Products, machines, leads, ladders/steps, vehicles, special tools]"),
            ("Emergency arrangements", "[First aider, emergency services, nearest clinic, assembly point, site contact]"),
            ("Client/site rules", "[Access, photos, waste, parking, alarms, security, occupants, pets, public areas]"),
            ("Stop-work triggers", "[Site-specific stop-work triggers]"),
            ("Supervisor release", "[Name, signature, time]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )
    add_table(
        doc,
        ["Hazard", "Who may be harmed", "Initial risk", "Controls required", "Residual risk", "Action owner"],
        [
            ("", "", "[L x S = ]", "", "[Low / Med / High]", ""),
            ("", "", "[L x S = ]", "", "[Low / Med / High]", ""),
            ("", "", "[L x S = ]", "", "[Low / Med / High]", ""),
            ("", "", "[L x S = ]", "", "[Low / Med / High]", ""),
            ("", "", "[L x S = ]", "", "[Low / Med / High]", ""),
            ("", "", "[L x S = ]", "", "[Low / Med / High]", ""),
        ],
        [1500, 1900, 1300, 2960, 1300, 1400],
        header_fill=LIGHT_GREY,
        font_size=7.9,
    )

    add_heading(doc, "Appendix B: Toolbox Talk and Attendance", 1)
    add_table(
        doc,
        ["Briefing item", "Covered", "Notes"],
        [
            ("Scope, exclusions, client/site rules, and work sequence.", "[ ]", ""),
            ("Key hazards and controls from the task risk assessment.", "[ ]", ""),
            ("PPE, chemical labels, SDS access, and product mixing restrictions.", "[ ]", ""),
            ("Wet floor, cable/hose, public access, and property protection controls.", "[ ]", ""),
            ("Emergency contacts, first aid, assembly point, and incident reporting route.", "[ ]", ""),
            ("Stop-work triggers and supervisor release to start.", "[ ]", ""),
        ],
        [6100, 1000, 2260],
        header_fill=LIGHT_BLUE,
        font_size=8.3,
    )
    add_table(
        doc,
        ["Name", "Role", "Signature", "Time"],
        [
            ("", "", "", ""),
            ("", "", "", ""),
            ("", "", "", ""),
            ("", "", "", ""),
        ],
        [2600, 2000, 2860, 1900],
        header_fill=LIGHT_GREY,
        font_size=8.5,
    )

    add_heading(doc, "Appendix C: Chemical and PPE Risk Register", 1)
    add_table(
        doc,
        ["Product / task", "Hazards", "Required controls / PPE", "SDS checked", "Approved by"],
        [
            ("", "", "", "[ ]", ""),
            ("", "", "", "[ ]", ""),
            ("", "", "", "[ ]", ""),
            ("", "", "", "[ ]", ""),
            ("", "", "", "[ ]", ""),
            ("", "", "", "[ ]", ""),
            ("", "", "", "[ ]", ""),
            ("", "", "", "[ ]", ""),
        ],
        [1900, 2160, 3200, 1100, 1000],
        header_fill=LIGHT_GREY,
        font_size=8.0,
    )
    add_table(
        doc,
        ["PPE item", "Issued to", "Condition", "Replacement needed", "Initial"],
        [
            ("Gloves", "", "[Good / Damaged]", "[ ]", ""),
            ("Eye protection", "", "[Good / Damaged]", "[ ]", ""),
            ("Mask / respirator", "", "[Good / Damaged / N/A]", "[ ]", ""),
            ("Apron / overalls", "", "[Good / Damaged]", "[ ]", ""),
            ("Footwear", "", "[Good / Damaged]", "[ ]", ""),
            ("Other", "", "", "[ ]", ""),
        ],
        [2000, 2400, 2100, 1700, 1160],
        header_fill=LIGHT_GREY,
        font_size=8.3,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix D: Incident Report Form", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Report type", "[Injury / near miss / exposure / spill / property damage / security / complaint / environmental / other]"),
            ("Date, time, location", "[Details]"),
            ("Client / site / job number", "[Details]"),
            ("Reported by", "[Name, role, contact]"),
            ("Supervisor", "[Name, role, contact]"),
            ("Person(s) involved", "[Names, roles, contact details where appropriate]"),
            ("Description of incident", "[What happened, task in progress, area, equipment/chemical involved, sequence of events]"),
            ("Immediate action taken", "[First aid, area made safe, spill controlled, equipment isolated, client notified]"),
            ("Injury / exposure details", "[Body part, symptoms, treatment, medical referral, product/exposure details]"),
            ("Damage / loss details", "[Item, area, owner, condition, value estimate if known]"),
            ("Photos / evidence", "[Authorised? File location? Labels/products retained? CCTV/witnesses?]"),
            ("Work status", "[Continued / paused / stopped / rebooked / area isolated]"),
            ("Notifications made", "[Supervisor, Operations, SHEQ, client, HR/COIDA, insurer, emergency services]"),
            ("Reportability check", "[COIDA / OHS / client / insurer / municipal / not reportable - reason]"),
            ("Completed by", "[Name, signature, date]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_BLUE,
        font_size=8.1,
    )

    add_heading(doc, "Appendix E: Witness Statement", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Witness name and contact", "[Name, role, phone/email]"),
            ("Relationship to incident", "[Employee / client / public / contractor / other]"),
            ("Where were you?", "[Position and activity at the time]"),
            ("What did you see or hear?", "[Witness account in own words]"),
            ("What happened immediately after?", "[Actions, notifications, first aid, area control]"),
            ("Other witnesses / evidence", "[Names, CCTV, photos, documents, products, equipment]"),
            ("Statement given voluntarily", "[Yes / No]"),
            ("Signature and date", "[Signature / date / time]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix F: Injury / Exposure and First Aid Log", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Employee / person affected", "[Name, role, employee number if applicable]"),
            ("Injury / exposure", "[Cut / strain / slip / splash / inhalation / burn / other]"),
            ("Body part / symptoms", "[Details]"),
            ("Chemical or equipment involved", "[Product/equipment name, label, SDS reference, batch if relevant]"),
            ("First aid given", "[Action, time, by whom]"),
            ("Medical referral", "[Clinic/hospital/doctor, transport, time, person accompanying]"),
            ("Returned to work?", "[Yes / No / restricted duties / sent home]"),
            ("COIDA owner notified", "[Name, time, method]"),
            ("Follow-up required", "[Medical certificate, statement, claim number, return-to-work review]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_BLUE,
        font_size=8.4,
    )
    add_table(
        doc,
        ["First-aid item used", "Qty", "Replacement needed", "Initial"],
        [
            ("", "", "[ ]", ""),
            ("", "", "[ ]", ""),
            ("", "", "[ ]", ""),
            ("", "", "[ ]", ""),
        ],
        [4400, 1200, 2200, 1560],
        header_fill=LIGHT_GREY,
        font_size=8.5,
    )

    add_heading(doc, "Appendix G: Property Damage / Client Notification", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Client / site", "[Name and address]"),
            ("Affected item / area", "[Description]"),
            ("Condition before work", "[Known pre-existing damage? Photos? Client notes?]"),
            ("What happened", "[Facts only]"),
            ("Immediate protective action", "[Area isolated, item protected, water stopped, product removed]"),
            ("Client notified", "[Name, time, method, summary]"),
            ("Photos / evidence", "[Authorised? File location?]"),
            ("Follow-up owner", "[Name, role, due date]"),
            ("Resolution", "[Repair / replacement / insurer / client credit / no action - reason]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix H: Corrective Action Tracker", 1)
    add_table(
        doc,
        ["Ref", "Finding / root cause", "Corrective action", "Owner", "Due date", "Status", "Verified by"],
        [
            ("", "", "", "", "", "[Open]", ""),
            ("", "", "", "", "", "[Open]", ""),
            ("", "", "", "", "", "[Open]", ""),
            ("", "", "", "", "", "[Open]", ""),
            ("", "", "", "", "", "[Open]", ""),
            ("", "", "", "", "", "[Open]", ""),
            ("", "", "", "", "", "[Open]", ""),
            ("", "", "", "", "", "[Open]", ""),
        ],
        [650, 2100, 2750, 1100, 950, 900, 910],
        header_fill=LIGHT_BLUE,
        font_size=7.8,
    )

    add_heading(doc, "Appendix I: Emergency and Escalation Contacts", 1)
    add_table(
        doc,
        ["Emergency / escalation contact", "Name", "Number", "After-hours instruction"],
        [
            ("Emergency services", "[Local emergency number]", "[Number]", "[Instruction]"),
            ("Operations Manager", "[Name]", "[Number]", "[Instruction]"),
            ("SHEQ Officer / Competent Person", "[Name]", "[Number]", "[Instruction]"),
            ("First Aider", "[Name]", "[Number]", "[Instruction]"),
            ("HR / COIDA Owner", "[Name]", "[Number]", "[Instruction]"),
            ("Client site contact", "[Name]", "[Number]", "[Instruction]"),
            ("Insurer / broker", "[Name]", "[Number]", "[Instruction]"),
            ("Waste / spill contractor", "[Name]", "[Number]", "[Instruction]"),
        ],
        [2800, 2100, 1800, 2660],
        header_fill=LIGHT_GREY,
        font_size=8.3,
    )
    add_para(
        doc,
        "End of template. Adapt this document with the business owner, supervisors, SHEQ representative, HR/COIDA owner, insurer, client contract owner, and competent compliance adviser.",
        italic=True,
        color=MUTED,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=8,
    )


def build_document():
    doc = Document()
    build_styles(doc)
    setup_header_footer(doc)

    props = doc.core_properties
    props.title = "Risk Assessment and Incident Report Template"
    props.subject = "Risk assessment and incident reporting business template for a South African cleaning services business"
    props.author = "OpenAI Codex"
    props.keywords = "risk assessment, incident report, cleaning services, SHEQ, South Africa, OHS, COIDA"

    add_cover(doc)
    add_document_control(doc)
    add_purpose_scope(doc)
    add_compliance(doc)
    add_roles(doc)
    add_risk_method(doc)
    add_hazard_register(doc)
    add_controls(doc)
    add_incident_process(doc)
    add_investigation(doc)
    add_records(doc)
    add_appendices(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT.resolve())
