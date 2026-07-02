from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Month_End_Business_Review.docx")

FONT = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
CAUTION_GOLD = RGBColor(0x7A, 0x5A, 0x00)
RISK_RED = RGBColor(0x9B, 0x1C, 0x1C)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
RISK_FILL = "FBEAEA"
BORDER = "C9D3DF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn("w:ascii"), name)
    r_pr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, before=0, after=6, line=1.25, keep_next=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next is not None:
        fmt.keep_with_next = keep_next


def style_para(p, size=11, color=BODY, bold=False, italic=False):
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_para(
    doc,
    text="",
    *,
    style=None,
    size=11,
    color=BODY,
    bold=False,
    italic=False,
    align=None,
    before=0,
    after=6,
    line=1.25,
    keep_next=None,
):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=before, after=after, line=line, keep_next=keep_next)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    style_para(
        p,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **margins):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ["top", "start", "bottom", "end"]:
        if margin not in margins:
            continue
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[margin]))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, size=8.7, color=BODY, align=None, fill=None, italic=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    if text is not None:
        for idx, part in enumerate(str(text).split("\n")):
            if idx:
                p.add_run().add_break()
            run = p.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=0, after=0, line=1.13)
    if align is not None:
        p.alignment = align


def set_table_borders(table, color=BORDER, sz="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell, **CELL_MARGIN)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            size=font_size,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=header_fill,
        )
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            align = None
            if str(value).strip() in {"[ ]", "[R]", "[%]", "[Score]", "[No.]", "[Yes / No / N/A]"}:
                align = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], value, size=font_size, align=align)
        prevent_row_split(table.rows[-1])
    add_para(doc, "", after=3)
    return table


def add_label_table(doc, rows, label_width=2700, font_size=9.2):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [label_width, CONTENT_WIDTH_DXA - label_width])
    set_table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=NAVY, fill=LIGHT_GREY, size=font_size)
        set_cell_text(cells[1], value, size=font_size)
        prevent_row_split(table.rows[-1])
    add_para(doc, "", after=3)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT_FILL, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D5DCE5", sz="4")
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(title)
    set_run_font(run, size=10.2, color=title_color, bold=True)
    p.add_run().add_break()
    run = p.add_run(body)
    set_run_font(run, size=9.4, color=BODY)
    set_paragraph_format(p, before=0, after=0, line=1.18)
    add_para(doc, "", after=5)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.2, color=BODY)
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_headers_and_footers(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    left = hp.add_run("MBR | Month End Business Review")
    set_run_font(left, size=9, color=MUTED)
    right = hp.add_run("    [Business Name]")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    left = fp.add_run("Controlled document - verify latest version before use")
    set_run_font(left, size=9, color=MUTED)


def add_cover(doc):
    add_para(doc, "Management Review Template", size=10, color=CAUTION_GOLD, bold=True, after=0)
    add_para(
        doc,
        "Month End Business Review",
        size=27,
        color=NAVY,
        bold=True,
        after=6,
        line=1.1,
    )
    add_para(
        doc,
        "For a South African cleaning services business",
        size=13,
        color=MUTED,
        after=18,
    )
    add_label_table(
        doc,
        [
            ("Business name", "[Business Name]"),
            ("Review month", "[Month YYYY]"),
            ("Prepared by", "[Name and role]"),
            ("Review meeting date", "[YYYY-MM-DD]"),
            ("Attendees", "[Owner / operations / finance / HR / sales / supervisors]"),
            ("Business unit / region", "[All / branch / region / service line]"),
            ("Document status", "[Draft / reviewed / approved / archived]"),
            ("Record location", "[Shared drive / CRM / accounting folder / board pack]"),
        ],
    )
    add_callout(
        doc,
        "Purpose",
        "Use this pack to review the month just closed, agree the next month's priorities, and keep a practical record of decisions. The template covers revenue, cash, operations, quality, client experience, people, safety, compliance, sales pipeline, risks, and actions.",
    )


def add_document_control(doc):
    add_heading(doc, "1. Document Control", 1)
    add_table(
        doc,
        ["Control field", "Template entry", "Notes"],
        [
            ("Document number", "MBR-[YYYY-MM]", "Use one pack per month or per branch."),
            ("Version", "[1.0]", "Update when KPI definitions or review rules change."),
            ("Owner", "[Managing Director / Operations Manager]", "Assign pack preparation and final sign-off."),
            ("Review cycle", "Monthly, within [x] working days after month end.", "Align with accounting close, payroll, and client invoicing."),
            ("Evidence folder", "[Folder/link]", "Store source reports, minutes, action proof, and signed review pack."),
            ("Retention", "[Company retention schedule]", "Keep only necessary personal information and store securely."),
        ],
        [2100, 4100, 3160],
    )
    add_heading(doc, "2. How to Use This Review Pack", 1)
    add_bullet(doc, "Populate the dashboard from source reports before the meeting: accounting, payroll, CRM, job management, audits, incident logs, and stock records.")
    add_bullet(doc, "Discuss exceptions first: missed targets, client complaints, safety events, cash constraints, overdue debtors, capacity gaps, and repeated defects.")
    add_bullet(doc, "Record decisions in plain language with an owner and due date. Carry unresolved actions into next month's review.")
    add_bullet(doc, "Avoid adding unnecessary personal information. Use employee or client names only where required for a legitimate business record.")
    add_bullet(doc, "Treat compliance prompts as a management reminder, not legal advice. Verify current laws, SARS rules, sector requirements, and client contracts.")
    add_heading(doc, "2.1 Meeting Cadence", 2)
    add_table(
        doc,
        ["Step", "Owner", "Timing", "Output"],
        [
            ("Close source data", "[Finance / Ops / HR]", "Day 1-3 after month end", "Invoices, bank receipts, payroll, job records, stock, audits, incidents."),
            ("Prepare pack", "[Pack owner]", "Day 3-5", "Draft dashboard, exceptions, decisions needed, and proposed actions."),
            ("Review meeting", "[Leadership team]", "Day 5-7", "Decisions, accountable owners, target changes, and escalations."),
            ("Follow-up", "[Action owners]", "Weekly until closed", "Action evidence and status update for next month."),
        ],
        [2100, 2000, 2300, 2960],
        font_size=8.7,
    )


def add_executive_dashboard(doc):
    add_page_break(doc)
    add_heading(doc, "3. Executive Summary Dashboard", 1)
    add_callout(
        doc,
        "Month-end decision focus",
        "Start the meeting with the dashboard trend, then agree only the decisions that change next month's cash, service quality, client retention, safety, people capacity, or sales conversion.",
    )
    add_table(
        doc,
        ["KPI", "Target", "Actual", "Trend", "Comment / action"],
        [
            ("Revenue invoiced", "R [target]", "R [actual]", "[Up / flat / down]", "[Drivers, lost work, new work]"),
            ("Gross margin", "[%]", "[%]", "[Trend]", "[Labour, chemicals, transport, rework impact]"),
            ("Net profit / EBITDA", "R [%]", "R [%]", "[Trend]", "[Main variance]"),
            ("Cash collected", "R [target]", "R [actual]", "[Trend]", "[Debtors collected / shortfall]"),
            ("Debtor days / overdue value", "[Days / R]", "[Actual]", "[Trend]", "[Top overdue accounts]"),
            ("Jobs completed", "[No.]", "[No.]", "[Trend]", "[Recurring / once-off / deep clean mix]"),
            ("On-time arrival / completion", "[%]", "[%]", "[Trend]", "[Route, staff, traffic, access issues]"),
            ("Quality audit average", "[%]", "[%]", "[Trend]", "[Repeat defects]"),
            ("Client complaints", "[No. max]", "[No.]", "[Trend]", "[Open / closed / escalated]"),
            ("Incidents / near misses", "[0 target]", "[No.]", "[Trend]", "[Safety theme and control]"),
            ("Staff availability", "[%]", "[%]", "[Trend]", "[Absence, turnover, overtime, vacancies]"),
            ("Sales pipeline value", "R [target]", "R [actual]", "[Trend]", "[Quote conversion / next wins]"),
        ],
        [2200, 1450, 1450, 1250, 3010],
        font_size=8.1,
    )
    add_heading(doc, "3.1 Decisions and Escalations", 2)
    add_table(
        doc,
        ["No.", "Decision / escalation", "Reason", "Owner", "Due date", "Status"],
        [
            ("1", "[Decision required]", "[Impact on margin, client, cash, safety, people, or compliance]", "[Name]", "[Date]", "[Open]"),
            ("2", "[Decision required]", "[Reason]", "[Name]", "[Date]", "[Open]"),
            ("3", "[Decision required]", "[Reason]", "[Name]", "[Date]", "[Open]"),
            ("4", "[Decision required]", "[Reason]", "[Name]", "[Date]", "[Open]"),
        ],
        [650, 2550, 2900, 1250, 1050, 960],
        font_size=8.3,
    )


def add_financial_review(doc):
    add_page_break(doc)
    add_heading(doc, "4. Financial Performance Review", 1)
    add_heading(doc, "4.1 Revenue by Service Line", 2)
    add_table(
        doc,
        ["Service line", "Budget", "Actual", "Variance", "Gross margin", "Comment"],
        [
            ("Recurring residential cleaning", "R [ ]", "R [ ]", "R / %", "[%]", "[New/cancelled recurring clients]"),
            ("Once-off / deep cleaning", "R [ ]", "R [ ]", "R / %", "[%]", "[Promotion, seasonal demand, capacity]"),
            ("Move-in / move-out cleaning", "R [ ]", "R [ ]", "R / %", "[%]", "[Estate agents, property managers]"),
            ("Commercial / office cleaning", "R [ ]", "R [ ]", "R / %", "[%]", "[Contract changes or missed shifts]"),
            ("Specialist work", "R [ ]", "R [ ]", "R / %", "[%]", "[Carpet, windows, post-construction, other]"),
            ("Total", "R [ ]", "R [ ]", "R / %", "[%]", "[Explain total variance]"),
        ],
        [2100, 1200, 1200, 1200, 1300, 2360],
        font_size=8.3,
    )
    add_heading(doc, "4.2 Profit and Loss Snapshot", 2)
    add_table(
        doc,
        ["Line item", "Budget", "Actual", "Variance", "Management note"],
        [
            ("Revenue invoiced", "R [ ]", "R [ ]", "R [ ]", "[Billing completeness and timing]"),
            ("Direct labour", "R [ ]", "R [ ]", "R [ ]", "[Normal hours, overtime, subcontractors]"),
            ("Cleaning chemicals and consumables", "R [ ]", "R [ ]", "R [ ]", "[Usage, waste, price increases]"),
            ("Transport / fuel / vehicle costs", "R [ ]", "R [ ]", "R [ ]", "[Route efficiency and maintenance]"),
            ("Equipment repairs / hire", "R [ ]", "R [ ]", "R [ ]", "[Vacuum, machine, ladder, tool issues]"),
            ("Admin, rent, insurance, marketing", "R [ ]", "R [ ]", "R [ ]", "[Overheads]"),
            ("Net profit / EBITDA", "R [ ]", "R [ ]", "R [ ]", "[Corrective management action]"),
        ],
        [2450, 1350, 1350, 1350, 2860],
        font_size=8.3,
    )
    add_heading(doc, "4.3 Cash, Debtors, and Credit Control", 2)
    add_table(
        doc,
        ["Cash / debtor item", "Amount / count", "Risk", "Owner / action"],
        [
            ("Opening bank balance", "R [ ]", "[Low / normal / high]", "[Note]"),
            ("Cash received in month", "R [ ]", "[Trend]", "[Note]"),
            ("Invoices issued but unpaid", "R [ ]", "[Debtor days]", "[Top accounts and follow-up]"),
            ("Overdue 30+ days", "R [ ]", "[Client/service risk]", "[Collection plan]"),
            ("Credit notes / write-offs", "R [ ]", "[Reason]", "[Approval and prevention]"),
            ("Forecast cash pressure next month", "R [ ]", "[Payroll, suppliers, tax, vehicles]", "[Mitigation]"),
        ],
        [2600, 1850, 2300, 2610],
        font_size=8.4,
    )


def add_month_end_admin(doc):
    add_heading(doc, "5. Month-End Admin, Tax, and Payroll Checklist", 1)
    add_table(
        doc,
        ["Checklist item", "Done", "Evidence / exception", "Owner"],
        [
            ("All completed jobs invoiced and matched to job cards / client approvals.", "[ ]", "[Invoice list / exception]", "[Finance]"),
            ("Supplier invoices captured, especially chemicals, consumables, fuel, PPE, repairs, and subcontractors.", "[ ]", "[Supplier folder]", "[Finance]"),
            ("Bank receipts allocated and debtor follow-ups issued.", "[ ]", "[Bank / age analysis]", "[Finance]"),
            ("VAT status checked and VAT201 prepared if the business is VAT-registered.", "[ ]", "[VAT201 / accountant note]", "[Finance]"),
            ("EMP201 / PAYE, SDL, and UIF reviewed where applicable.", "[ ]", "[Payroll report / SARS proof]", "[Payroll]"),
            ("Payroll hours, overtime, leave, deductions, loans, and payslips checked.", "[ ]", "[Payroll sign-off]", "[Payroll / HR]"),
            ("National Minimum Wage and any sector/client requirements verified for pay period.", "[ ]", "[Current rate/source checked]", "[HR / Payroll]"),
            ("COIDA / Compensation Fund obligations and Letter of Good Standing status checked where needed.", "[ ]", "[LOGS / ROE note]", "[HR / Finance]"),
            ("CIPC annual return and beneficial ownership timing checked if month is near anniversary date.", "[ ]", "[CIPC status]", "[Company secretary / owner]"),
        ],
        [3950, 750, 3350, 1310],
        font_size=8.2,
    )
    add_callout(
        doc,
        "Finance caution",
        "This checklist does not replace advice from a registered tax practitioner, accountant, payroll provider, or labour adviser. Use it to make sure month-end questions are asked early and evidence is stored.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )


def add_operations_review(doc):
    add_page_break(doc)
    add_heading(doc, "6. Operations Performance Review", 1)
    add_heading(doc, "6.1 Job Delivery and Capacity", 2)
    add_table(
        doc,
        ["Operational metric", "Target", "Actual", "Variance / cause", "Next action"],
        [
            ("Recurring visits completed", "[No.]", "[No.]", "[Missed / extra / rescheduled]", "[Action]"),
            ("Once-off jobs completed", "[No.]", "[No.]", "[Demand / capacity]", "[Action]"),
            ("Cancelled visits", "[No. max]", "[No.]", "[Client / staff / access / weather]", "[Prevention]"),
            ("Rework visits", "[No. max]", "[No.]", "[Quality / scope / staff]", "[Training or process action]"),
            ("Average job duration", "[Hours]", "[Hours]", "[Scope creep / inefficiency]", "[Action]"),
            ("Route productivity", "[Jobs per team/day]", "[Actual]", "[Travel / loading / traffic]", "[Route change]"),
            ("Stock-outs or emergency purchases", "[0]", "[No.]", "[Cause]", "[Min/max stock action]"),
            ("Equipment downtime", "[Hours]", "[Hours]", "[Machine/tool issue]", "[Repair / replace]"),
        ],
        [2500, 1250, 1250, 2700, 1660],
        font_size=8.3,
    )
    add_heading(doc, "6.2 Mobile Service Setup and Pack Down Review", 2)
    add_table(
        doc,
        ["Control area", "Review question", "Result", "Action / owner"],
        [
            ("Vehicle loading", "Were chemicals, equipment, PPE, warning signs, and job packs loaded correctly?", "[Good / issue]", "[Action]"),
            ("Chemical transport", "Were products closed, labelled, upright, and segregated where needed?", "[Good / issue]", "[Action]"),
            ("Arrival setup", "Were team briefings, site risks, access rules, and wet-floor controls applied?", "[Good / issue]", "[Action]"),
            ("Pack down", "Were products capped, equipment cleaned, waste handled, keys returned, and job cards updated?", "[Good / issue]", "[Action]"),
            ("Route learning", "What route, parking, access, stock, or staffing lesson must be built into next month?", "[Learning]", "[Action]"),
        ],
        [2100, 4300, 1250, 1710],
        font_size=8.4,
    )


def add_quality_client_review(doc):
    add_page_break(doc)
    add_heading(doc, "7. Quality and Client Experience Review", 1)
    add_heading(doc, "7.1 Quality Score and Defect Themes", 2)
    add_table(
        doc,
        ["Quality area", "Score / count", "Top finding", "Root cause", "Corrective action"],
        [
            ("Audit average", "[%]", "[Best / worst site]", "[Cause]", "[Action]"),
            ("Bathroom / hygiene defects", "[No.]", "[Repeat item]", "[Training / time / product]", "[Action]"),
            ("Kitchen / food-area defects", "[No.]", "[Repeat item]", "[Cause]", "[Action]"),
            ("Floors, carpets, and edges", "[No.]", "[Repeat item]", "[Cause]", "[Action]"),
            ("High-touch points", "[No.]", "[Repeat item]", "[Cause]", "[Action]"),
            ("Presentation / final walk-through", "[No.]", "[Repeat item]", "[Cause]", "[Action]"),
            ("Client property / damage", "[No.]", "[Incident theme]", "[Cause]", "[Action]"),
        ],
        [2100, 1300, 2500, 1800, 1660],
        font_size=8.4,
    )
    add_heading(doc, "7.2 Complaints, Compliments, and Retention", 2)
    add_table(
        doc,
        ["Client / segment", "Feedback type", "Status", "Value at risk", "Owner / next step"],
        [
            ("[Client / residential segment]", "[Complaint / compliment / request]", "[Open / closed]", "R [ ]", "[Action]"),
            ("[Client / segment]", "[Feedback]", "[Status]", "R [ ]", "[Action]"),
            ("[Client / segment]", "[Feedback]", "[Status]", "R [ ]", "[Action]"),
            ("[Client / segment]", "[Feedback]", "[Status]", "R [ ]", "[Action]"),
            ("[Client / segment]", "[Feedback]", "[Status]", "R [ ]", "[Action]"),
        ],
        [2050, 2300, 1300, 1350, 2360],
        font_size=8.4,
    )
    add_heading(doc, "7.3 Client Renewal and Churn Watchlist", 2)
    add_table(
        doc,
        ["Client", "Renewal / risk date", "Reason for risk", "Retention action", "Owner"],
        [
            ("[Client]", "[Date]", "[Price / quality / relationship / competitor]", "[Call, visit, service recovery, quote]", "[Name]"),
            ("[Client]", "[Date]", "[Reason]", "[Action]", "[Name]"),
            ("[Client]", "[Date]", "[Reason]", "[Action]", "[Name]"),
        ],
        [1900, 1650, 2500, 2210, 1100],
        font_size=8.4,
    )


def add_people_review(doc):
    add_page_break(doc)
    add_heading(doc, "8. People, HR, and Training Review", 1)
    add_heading(doc, "8.1 Workforce Snapshot", 2)
    add_table(
        doc,
        ["People metric", "Target / plan", "Actual", "Issue / trend", "Action"],
        [
            ("Headcount by role", "[Cleaners / leaders / drivers / supervisors]", "[Actual]", "[Vacancies / surplus]", "[Recruit / redeploy]"),
            ("Attendance / absenteeism", "[%]", "[%]", "[Pattern / reason]", "[Counselling / support / action]"),
            ("Overtime hours", "[Hours / cost]", "[Actual]", "[Demand / inefficiency]", "[Plan]"),
            ("Leave liability / leave taken", "[Plan]", "[Actual]", "[Operational impact]", "[Roster plan]"),
            ("Turnover / resignations", "[No. max]", "[No.]", "[Reason]", "[Retention action]"),
            ("Disciplinary / grievance matters", "[Open / closed]", "[Actual]", "[Theme]", "[HR action]"),
            ("Uniform / PPE issue", "[Complete]", "[Gaps]", "[Size / stock / damage]", "[Issue / replace]"),
        ],
        [2450, 2200, 1200, 2250, 1260],
        font_size=8.4,
    )
    add_heading(doc, "8.2 Training and Competence Review", 2)
    add_table(
        doc,
        ["Training area", "Due / required", "Completed", "Gap", "Owner / due"],
        [
            ("OHS induction / refresher", "[No.]", "[No.]", "[Employees / teams]", "[Name / date]"),
            ("Chemical and SDS training", "[No.]", "[No.]", "[Products / staff]", "[Name / date]"),
            ("PPE use and replacement", "[No.]", "[No.]", "[Gap]", "[Name / date]"),
            ("Equipment authorisation", "[No.]", "[No.]", "[Vacuum / machine / ladder / vehicle]", "[Name / date]"),
            ("Residential cleaning process", "[No.]", "[No.]", "[Area-specific weakness]", "[Name / date]"),
            ("Client privacy and confidentiality", "[No.]", "[No.]", "[Refresher due]", "[Name / date]"),
            ("Supervisor coaching completed", "[No.]", "[No.]", "[Team / topic]", "[Name / date]"),
        ],
        [2500, 1550, 1350, 2600, 1360],
        font_size=8.4,
    )
    add_callout(
        doc,
        "People decision prompt",
        "Confirm whether next month needs recruitment, roster changes, supervisor coaching, disciplinary follow-up, wage-rate verification, leave planning, or extra training before new client work is accepted.",
    )


def add_safety_compliance_review(doc):
    add_page_break(doc)
    add_heading(doc, "9. Safety, Compliance, and Risk Controls", 1)
    add_table(
        doc,
        ["Control area", "Month result", "Exception / evidence", "Action owner"],
        [
            ("Incidents, injuries, and near misses", "[No. / severity]", "[Incident register / investigation]", "[Owner]"),
            ("COIDA / Compensation Fund reporting", "[Checked]", "[Injury report / LOGS / ROE status]", "[Owner]"),
            ("OHS risk assessments", "[Current / needs update]", "[New sites, tasks, chemicals, equipment]", "[Owner]"),
            ("Hazardous chemical agents", "[Compliant / issue]", "[SDS, labels, dilution, storage, disposal, exposure controls]", "[Owner]"),
            ("PPE availability and use", "[Good / issue]", "[Gloves, eye protection, footwear, masks, aprons]", "[Owner]"),
            ("Vehicle and driving risk", "[Good / issue]", "[Licences, incidents, maintenance, route risk]", "[Owner]"),
            ("Client privacy / POPIA", "[Good / issue]", "[Photos, access codes, keys, client records, staff data]", "[Owner]"),
            ("Waste / environmental controls", "[Good / issue]", "[Chemical containers, spills, client waste rules]", "[Owner]"),
            ("Insurance and contracts", "[Current / review]", "[Claims, policy renewals, client SLA changes]", "[Owner]"),
        ],
        [2650, 1750, 3650, 1310],
        font_size=8.2,
    )
    add_heading(doc, "9.1 Compliance Calendar Watchlist", 2)
    add_table(
        doc,
        ["Item", "Frequency / trigger", "Next due", "Evidence location"],
        [
            ("SARS tax and employer submissions", "Monthly / bi-monthly / provisional periods as applicable", "[Date]", "[eFiling / accountant folder]"),
            ("Payroll and minimum wage review", "Each pay period and when rates change", "[Date]", "[Payroll folder]"),
            ("COIDA Return of Earnings / Letter of Good Standing", "Annual and as required for clients/tenders", "[Date]", "[Compensation Fund portal]"),
            ("CIPC annual return / beneficial ownership", "Annual / when company details change", "[Date]", "[CIPC profile]"),
            ("Insurance renewals", "Annual and after material business change", "[Date]", "[Policy folder]"),
            ("SDS / product approvals", "When products change or SDS updates", "[Date]", "[Chemical register]"),
            ("POPIA/privacy review", "At least annually or after incident/change", "[Date]", "[Privacy folder]"),
        ],
        [2600, 3100, 1400, 2260],
        font_size=8.3,
    )


def add_sales_review(doc):
    add_page_break(doc)
    add_heading(doc, "10. Sales, Marketing, and Pipeline Review", 1)
    add_table(
        doc,
        ["Pipeline metric", "Target", "Actual", "Conversion / quality note", "Next action"],
        [
            ("Leads received", "[No.]", "[No.]", "[Source quality]", "[Action]"),
            ("Site assessments completed", "[No.]", "[No.]", "[No-shows / fit]", "[Action]"),
            ("Quotes issued", "R [ ]", "R [ ]", "[Average value / margin]", "[Action]"),
            ("Quotes won", "R [ ]", "R [ ]", "[Conversion %]", "[Action]"),
            ("Quotes lost", "R [ ]", "R [ ]", "[Price / timing / competitor / fit]", "[Action]"),
            ("Recurring contracts added", "[No.]", "[No.]", "[Monthly value]", "[Onboarding action]"),
            ("Reviews / referrals requested", "[No.]", "[No.]", "[Google / WhatsApp / client referral]", "[Action]"),
        ],
        [2300, 1250, 1250, 2800, 1760],
        font_size=8.4,
    )
    add_heading(doc, "10.1 Top Opportunities and Renewals", 2)
    add_table(
        doc,
        ["Opportunity / client", "Service", "Value", "Stage", "Decision date", "Owner / next step"],
        [
            ("[Prospect/client]", "[Residential / commercial / deep clean]", "R [ ]", "[Lead / quote / negotiation]", "[Date]", "[Next step]"),
            ("[Prospect/client]", "[Service]", "R [ ]", "[Stage]", "[Date]", "[Next step]"),
            ("[Prospect/client]", "[Service]", "R [ ]", "[Stage]", "[Date]", "[Next step]"),
            ("[Prospect/client]", "[Service]", "R [ ]", "[Stage]", "[Date]", "[Next step]"),
            ("[Prospect/client]", "[Service]", "R [ ]", "[Stage]", "[Date]", "[Next step]"),
        ],
        [1850, 2000, 1100, 1700, 1250, 1460],
        font_size=8.2,
    )
    add_heading(doc, "10.2 Pricing and Margin Review", 2)
    add_table(
        doc,
        ["Pricing item", "Finding", "Impact", "Decision required"],
        [
            ("Hourly / per-job rate adequacy", "[Rate covers labour, travel, chemicals, overhead, margin?]", "[R / %]", "[Increase / hold / reprice]"),
            ("Chemical and consumable cost trend", "[Price increase / usage issue]", "[R / %]", "[Supplier / stock / client price action]"),
            ("Travel and route cost", "[Fuel/time/parking/tolls]", "[R / %]", "[Zone fee / route change]"),
            ("Discounts and promotions", "[Used this month]", "[Margin impact]", "[Continue / stop / tighten approval]"),
        ],
        [2300, 3300, 1400, 2360],
        font_size=8.4,
    )


def add_risk_action_review(doc):
    add_page_break(doc)
    add_heading(doc, "11. Risk, Opportunity, and Action Plan", 1)
    add_heading(doc, "11.1 Risk and Opportunity Register", 2)
    add_table(
        doc,
        ["Type", "Risk / opportunity", "Likelihood", "Impact", "Control / action", "Owner"],
        [
            ("Risk", "[Cash shortfall from overdue debtors]", "[L/M/H]", "[L/M/H]", "[Collection plan / credit terms]", "[Name]"),
            ("Risk", "[Quality defects causing churn]", "[L/M/H]", "[L/M/H]", "[Audit / coaching / service recovery]", "[Name]"),
            ("Risk", "[Staff absence or turnover]", "[L/M/H]", "[L/M/H]", "[Recruitment / roster / training]", "[Name]"),
            ("Risk", "[Chemical incident or unsafe work]", "[L/M/H]", "[L/M/H]", "[SDS / PPE / supervision]", "[Name]"),
            ("Risk", "[Vehicle or equipment failure]", "[L/M/H]", "[L/M/H]", "[Maintenance / backup plan]", "[Name]"),
            ("Opportunity", "[New route density / recurring client cluster]", "[L/M/H]", "[L/M/H]", "[Sales and route plan]", "[Name]"),
            ("Opportunity", "[Upsell deep clean / carpets / windows]", "[L/M/H]", "[L/M/H]", "[Campaign / client call list]", "[Name]"),
        ],
        [1050, 2600, 1250, 1050, 2700, 710],
        font_size=8.1,
    )
    add_heading(doc, "11.2 Next Month Targets", 2)
    add_table(
        doc,
        ["Target area", "Next month target", "Key driver", "Owner"],
        [
            ("Revenue", "R [target]", "[New jobs, recurring retention, upsells]", "[Name]"),
            ("Gross margin", "[%]", "[Labour, route density, materials, rework]", "[Name]"),
            ("Cash collected", "R [target]", "[Debtor follow-up and payment terms]", "[Name]"),
            ("Quality audit average", "[%]", "[Defect reduction and coaching]", "[Name]"),
            ("Client complaints", "[No. max]", "[Service recovery and prevention]", "[Name]"),
            ("Incidents / near misses", "[0 target / reporting target]", "[Risk controls and toolbox talks]", "[Name]"),
            ("Sales pipeline", "R [target]", "[Lead generation, quotes, renewals]", "[Name]"),
        ],
        [2200, 2050, 3800, 1310],
        font_size=8.4,
    )
    add_heading(doc, "11.3 Action Plan", 2)
    add_table(
        doc,
        ["Action", "Expected outcome", "Owner", "Due date", "Status / evidence"],
        [
            ("[Action 1]", "[Outcome]", "[Name]", "[Date]", "[Open / proof]"),
            ("[Action 2]", "[Outcome]", "[Name]", "[Date]", "[Open / proof]"),
            ("[Action 3]", "[Outcome]", "[Name]", "[Date]", "[Open / proof]"),
            ("[Action 4]", "[Outcome]", "[Name]", "[Date]", "[Open / proof]"),
            ("[Action 5]", "[Outcome]", "[Name]", "[Date]", "[Open / proof]"),
            ("[Action 6]", "[Outcome]", "[Name]", "[Date]", "[Open / proof]"),
        ],
        [2700, 2600, 1100, 1100, 1860],
        font_size=8.4,
    )


def add_appendices(doc):
    add_page_break(doc)
    add_heading(doc, "Appendix A: KPI Definitions", 1)
    add_table(
        doc,
        ["KPI", "Suggested definition", "Source"],
        [
            ("Revenue invoiced", "Total invoices issued for the month, excluding VAT if management accounts are prepared that way.", "[Accounting system]"),
            ("Gross margin", "(Revenue minus direct labour, consumables, transport, direct subcontractors, and job-specific costs) / revenue.", "[Management accounts]"),
            ("Net profit / EBITDA", "Profit measure agreed with accountant or management team; keep definition consistent month to month.", "[Management accounts]"),
            ("Debtor days", "Accounts receivable / average daily revenue, or the method used by the accounting system.", "[Age analysis]"),
            ("On-time arrival", "Visits where team arrived within agreed client window / total visits.", "[Job system / supervisor logs]"),
            ("Quality audit average", "Total achieved points / total applicable audit points.", "[QC audit forms]"),
            ("Staff availability", "Available staff hours / planned staff hours, excluding approved leave if agreed.", "[Roster / payroll]"),
            ("Client retention", "Active recurring clients retained / active recurring clients at start of period.", "[CRM / contract list]"),
        ],
        [2200, 5200, 1960],
        font_size=8.2,
    )
    add_heading(doc, "Appendix B: Meeting Attendance and Minutes", 1)
    add_label_table(
        doc,
        [
            ("Meeting date and time", "[YYYY-MM-DD / start-finish]"),
            ("Chair", "[Name]"),
            ("Minute taker", "[Name]"),
            ("Attendees", "[Names and roles]"),
            ("Apologies", "[Names and roles]"),
            ("Previous actions reviewed", "[Yes / no / exceptions]"),
        ],
        label_width=2500,
    )
    add_table(
        doc,
        ["Discussion item", "Key point / decision", "Owner", "Due date"],
        [
            ("Finance", "[Summary]", "[Name]", "[Date]"),
            ("Operations", "[Summary]", "[Name]", "[Date]"),
            ("Quality / clients", "[Summary]", "[Name]", "[Date]"),
            ("People / HR", "[Summary]", "[Name]", "[Date]"),
            ("Safety / compliance", "[Summary]", "[Name]", "[Date]"),
            ("Sales / pipeline", "[Summary]", "[Name]", "[Date]"),
        ],
        [2000, 4400, 1500, 1460],
        font_size=8.5,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix C: Source and Review List", 1)
    add_table(
        doc,
        ["Source", "URL / review note"],
        [
            ("SARS Businesses and Employers", "https://www.sars.gov.za/businesses-and-employers/"),
            ("SARS small business taxpayer guidance", "https://www.sars.gov.za/businesses-and-employers/small-businesses-taxpayers/"),
            ("SARS Pay As You Earn guidance", "https://www.sars.gov.za/types-of-tax/pay-as-you-earn/"),
            ("SARS payment rules", "https://www.sars.gov.za/guide-to-sars-payment-rules/"),
            ("Occupational Health and Safety Act 85 of 1993", "https://www.gov.za/documents/occupational-health-and-safety-act"),
            ("Regulations for Hazardous Chemical Agents, 2021", "https://www.labour.gov.za/DocumentCenter/Publications/Occupational%20Health%20and%20Safety/Regulations%20for%20Hazardous%20Chemical%20Agents%202021.pdf"),
            ("Basic Conditions of Employment Act 75 of 1997", "https://www.gov.za/documents/basic-conditions-employment-act"),
            ("National Minimum Wage Amendment 2026", "https://www.gov.za/sites/default/files/gcis_document/202602/54075rg11941gon7083.pdf"),
            ("Compensation for Occupational Injuries and Diseases Act 130 of 1993", "https://www.gov.za/documents/compensation-occupational-injuries-and-diseases-act"),
            ("Compensation Fund ROE online information", "https://www.labour.gov.za/Online-Tools/Pages/ROE-Online-%28cfonline-labour-gov-za%29.aspx"),
            ("Protection of Personal Information Act 4 of 2013", "https://www.gov.za/documents/protection-personal-information-act"),
            ("CIPC annual returns", "https://annualreturns.cipc.co.za/"),
        ],
        [3600, 5760],
        font_size=7.8,
    )
    add_callout(
        doc,
        "Review note",
        "Before issuing a completed month-end pack, confirm that source links, rates, deadlines, and client contract obligations are still current for the month under review.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )


def build_doc():
    doc = Document()
    configure_styles(doc)
    set_headers_and_footers(doc)

    add_cover(doc)
    add_document_control(doc)
    add_executive_dashboard(doc)
    add_financial_review(doc)
    add_page_break(doc)
    add_month_end_admin(doc)
    add_operations_review(doc)
    add_quality_client_review(doc)
    add_people_review(doc)
    add_safety_compliance_review(doc)
    add_sales_review(doc)
    add_risk_action_review(doc)
    add_appendices(doc)

    doc.core_properties.title = "Month End Business Review"
    doc.core_properties.subject = "South African cleaning services month-end business review template"
    doc.core_properties.keywords = "cleaning services, South Africa, month end, business review, operations, finance, OHS, POPIA, SARS, COIDA"
    doc.core_properties.comments = "Template generated for business use; verify current tax, payroll, labour, safety, privacy, and client requirements before implementation."
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUT.resolve()}")
