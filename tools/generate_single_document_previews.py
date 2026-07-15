from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from PIL import Image, ImageDraw, ImageFont


SOURCE_ROOT = Path(
    r"C:\Elzano Cox\Projects\dokkit-template-production\04_APPROVED_READY_SINGLE DOCUMENTS"
)
OUTPUT_ROOT = Path(
    r"C:\Elzano Cox\Projects\dokkit\public\images\previews\single-documents"
)
HR_SOURCE_ROOT = Path(
    r"C:\Elzano Cox\Projects\dokkit-template-production\03_PACKAGED_PRODUCTS\03_SINGLE_TEMPLATES"
)

DOCUMENTS = [
    {
        "slug": "permanent-employment-agreement-template",
        "title": "South African Permanent Employment Agreement Template",
        "source": HR_SOURCE_ROOT / "02_Permanent_Employment_Agreement.docx",
        "format": "Word",
    },
    {
        "slug": "business-financial-income-statement-template",
        "title": "Business Financial Income Statement Template",
        "source": SOURCE_ROOT
        / "Excel"
        / "DokKit_Business_Financial_Income_Statement_Template_v1.0.xlsx",
        "format": "Excel",
    },
    {
        "slug": "crm-tracker",
        "title": "CRM Tracker",
        "source": SOURCE_ROOT / "Excel" / "DokKit_CRM_Tracker_v1.xlsx",
        "format": "Excel",
    },
    {
        "slug": "income-and-expense-tracker",
        "title": "Income and Expense Tracker",
        "source": SOURCE_ROOT / "Excel" / "DokKit_Income_and_Expense_Tracker_v1.xlsx",
        "format": "Excel",
    },
    {
        "slug": "invoice-workbook-template",
        "title": "Invoice Workbook Template",
        "source": SOURCE_ROOT / "Excel" / "DokKit_Invoice_Workbook_Template_v1.xlsx",
        "format": "Excel",
    },
    {
        "slug": "fixed-term-employment-contract-template",
        "title": "South African Fixed-Term Employment Agreement Template",
        "source": HR_SOURCE_ROOT / "03_Fixed_Term_Employment_Agreement.docx",
        "format": "Word",
    },
    {
        "slug": "job-description-template",
        "title": "Job Description Template for South African Small Businesses",
        "source": HR_SOURCE_ROOT / "06_Job_Description_Template.docx",
        "format": "Word",
    },
    {
        "slug": "employee-onboarding-checklist-template",
        "title": "Employee Onboarding Checklist Template",
        "source": HR_SOURCE_ROOT / "11_Employee_Onboarding_Checklist.docx",
        "format": "Word",
    },
    {
        "slug": "employee-timesheet-template",
        "title": "Employee Timesheet Template",
        "source": HR_SOURCE_ROOT / "18_Timesheet_Template.docx",
        "format": "Word",
    },
    {
        "slug": "leave-application-form-template",
        "title": "Leave Application Form Template",
        "source": HR_SOURCE_ROOT / "20_Leave_Application_Form.docx",
        "format": "Word",
    },
    {
        "slug": "disciplinary-code-and-procedure-template",
        "title": "Disciplinary Code and Procedure Template",
        "source": HR_SOURCE_ROOT / "26_Disciplinary_Code_and_Procedure.docx",
        "format": "Word",
    },
    {
        "slug": "general-service-agreement-template",
        "title": "General Service Agreement Template",
        "source": SOURCE_ROOT
        / "Word"
        / "Dokkit_General_Service_Agreement_Template_v1.docx",
        "format": "Word",
    },
    {
        "slug": "joint-venture-structure-agreement-template",
        "title": "Joint Venture Structure Agreement Template",
        "source": SOURCE_ROOT
        / "Word"
        / "DokKit_Joint_Venture_Structure_Agreement_Template_v1.docx",
        "format": "Word",
    },
    {
        "slug": "master-quotation-template",
        "title": "Master Quotation Template",
        "source": SOURCE_ROOT / "Word" / "DokKit_Master_Quotation_Template_v1.docx",
        "format": "Word",
    },
    {
        "slug": "non-disclosure-agreement-template",
        "title": "Non-Disclosure Agreement Template",
        "source": SOURCE_ROOT
        / "Word"
        / "DokKit_Non_Disclosure_Agreement_Template_v1.docx",
        "format": "Word",
    },
    {
        "slug": "popia-privacy-policy-statement-template",
        "title": "POPIA Privacy Policy Statement Template",
        "source": SOURCE_ROOT
        / "Word"
        / "DokKit_POPIA_Privacy_Policy_Statement_Template_v1.docx",
        "format": "Word",
    },
    {
        "slug": "terms-and-conditions-template",
        "title": "Terms and Conditions Template",
        "source": SOURCE_ROOT
        / "Word"
        / "DokKit_Terms_and_Conditions_Template_v1.docx",
        "format": "Word",
    },
    {
        "slug": "vat-compliant-invoice-template",
        "title": "VAT-Compliant Invoice Template",
        "source": SOURCE_ROOT
        / "Word"
        / "DokKit_VAT_Compliant_Invoice_Template_v1.docx",
        "format": "Word",
    },
    {
        "slug": "vat-ready-purchase-order-template",
        "title": "VAT-Ready Purchase Order Template",
        "source": SOURCE_ROOT
        / "Word"
        / "DokKit_VAT_Ready_Purchase_Order_Template_v1.docx",
        "format": "Word",
    },
]

DOCX_SECTION_OVERRIDES = {
    "master-quotation-template": [
        (
            "Quoted to and references",
            "Client, contact, billing and quotation reference fields.",
        ),
        (
            "Goods or services",
            "Description, quantity, unit, price, discount, VAT and line totals.",
        ),
        (
            "Commercial notes",
            "Price basis, validity, scheduling and service or delivery details.",
        ),
        (
            "Acceptance",
            "Client name, company, signature and date fields.",
        ),
    ],
    "vat-compliant-invoice-template": [
        (
            "Billed to",
            "Customer, billing address and invoice reference fields.",
        ),
        (
            "Invoice items",
            "Descriptions, quantities, unit prices, VAT rates and line totals.",
        ),
        (
            "Amount payable",
            "Subtotal, VAT and invoice total sections.",
        ),
        (
            "Payment details",
            "Banking, remittance and payment-reference fields.",
        ),
    ],
    "vat-ready-purchase-order-template": [
        (
            "Supplier details",
            "Supplier identity, VAT status and quotation references.",
        ),
        (
            "Ordered items",
            "Descriptions, quantities, pricing, VAT and line totals.",
        ),
        (
            "Delivery details",
            "Delivery address, required date and purchasing notes.",
        ),
        (
            "Authorisation",
            "Prepared, checked and approved-by fields.",
        ),
    ],
}

DOCX_INTRO_OVERRIDES = {
    "master-quotation-template": (
        "A structured quotation layout for client details, pricing, VAT, "
        "commercial notes and acceptance."
    ),
    "vat-compliant-invoice-template": (
        "A VAT-ready tax invoice layout with customer, line-item, total and "
        "payment sections."
    ),
    "vat-ready-purchase-order-template": (
        "A VAT-ready purchase order layout with supplier, line-item, delivery "
        "and approval sections."
    ),
}

PAGE_W = 1100
PAGE_H = 1420
MARGIN = 76
CONTENT_W = PAGE_W - (MARGIN * 2)
FOOTER_H = 74
CONTENT_BOTTOM = PAGE_H - FOOTER_H - 24

BLACK = (17, 17, 17)
INK = (38, 38, 40)
MUTED = (95, 95, 102)
ORANGE = (255, 106, 0)
ORANGE_DARK = (217, 84, 0)
ORANGE_SOFT = (255, 244, 235)
WARM = (246, 244, 241)
LINE = (229, 224, 218)
WHITE = (255, 255, 255)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


FONT_BRAND = font(19, True)
FONT_HEADER_LABEL = font(17, True)
FONT_TITLE = font(34, True)
FONT_TITLE_SMALL = font(28, True)
FONT_SECTION = font(20, True)
FONT_BODY = font(19)
FONT_BODY_BOLD = font(19, True)
FONT_CELL = font(17)
FONT_CELL_BOLD = font(17, True)
FONT_SMALL = font(15)
FONT_SMALL_BOLD = font(15, True)
FONT_WATERMARK = font(72, True)


def clean_text(value: object) -> str:
    text = "" if value is None else str(value)
    return re.sub(r"\s+", " ", text).strip()


def display_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d %b %Y")
    if isinstance(value, date):
        return value.strftime("%d %b %Y")
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return clean_text(value)


def wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font_obj: ImageFont.ImageFont,
    max_width: int,
) -> list[str]:
    words = clean_text(text).split()
    if not words:
        return []

    lines: list[str] = []
    line = words[0]
    for word in words[1:]:
        candidate = f"{line} {word}"
        if draw.textbbox((0, 0), candidate, font=font_obj)[2] <= max_width:
            line = candidate
        else:
            lines.append(line)
            line = word
    lines.append(line)
    return lines


def clipped_lines(
    draw: ImageDraw.ImageDraw,
    text: str,
    font_obj: ImageFont.ImageFont,
    max_width: int,
    max_lines: int,
) -> list[str]:
    lines = wrap_text(draw, text, font_obj, max_width)
    if len(lines) <= max_lines:
        return lines

    clipped = lines[:max_lines]
    last = clipped[-1]
    while last and draw.textbbox((0, 0), f"{last}...", font=font_obj)[2] > max_width:
        last = last[:-1].rstrip()
    clipped[-1] = f"{last}..."
    return clipped


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font_obj: ImageFont.ImageFont,
    fill: tuple[int, int, int],
) -> None:
    bounds = draw.textbbox((0, 0), text, font=font_obj)
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    x = box[0] + ((box[2] - box[0] - width) // 2)
    y = box[1] + ((box[3] - box[1] - height) // 2) - bounds[1]
    draw.text((x, y), text, fill=fill, font=font_obj)


def draw_header(draw: ImageDraw.ImageDraw, title: str, file_format: str) -> int:
    header_h = 154
    draw.rectangle((0, 0, PAGE_W, header_h), fill=BLACK)

    logo_box = (MARGIN, 35, MARGIN + 66, 101)
    draw.rounded_rectangle(logo_box, radius=12, fill=WHITE)
    draw_centered_text(draw, logo_box, "DK", FONT_BRAND, BLACK)

    title_x = MARGIN + 88
    draw.text(
        (title_x, 28),
        "DokKit Preview",
        fill=(255, 176, 111),
        font=FONT_HEADER_LABEL,
    )
    format_width = draw.textbbox((0, 0), file_format, font=FONT_HEADER_LABEL)[2]
    draw.text(
        (PAGE_W - MARGIN - format_width, 28),
        file_format,
        fill=(255, 176, 111),
        font=FONT_HEADER_LABEL,
    )

    max_title_width = PAGE_W - MARGIN - title_x
    title_font = FONT_TITLE
    title_lines = wrap_text(draw, title, title_font, max_title_width)
    if len(title_lines) > 2:
        title_font = FONT_TITLE_SMALL
        title_lines = wrap_text(draw, title, title_font, max_title_width)
    title_lines = title_lines[:2]

    title_y = 56
    for line in title_lines:
        draw.text((title_x, title_y), line, fill=WHITE, font=title_font)
        title_y += 37

    return header_h + 28


def draw_watermark(image: Image.Image) -> None:
    overlay = Image.new("RGBA", image.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)
    text = "PREVIEW"
    text_box = draw.textbbox((0, 0), text, font=FONT_WATERMARK)
    text_w = text_box[2] - text_box[0]
    text_h = text_box[3] - text_box[1]
    watermark = Image.new("RGBA", (text_w + 80, text_h + 80), (255, 255, 255, 0))
    watermark_draw = ImageDraw.Draw(watermark)
    watermark_draw.text((40, 40), text, font=FONT_WATERMARK, fill=(255, 106, 0, 42))
    watermark = watermark.rotate(-28, expand=True)
    overlay.alpha_composite(
        watermark,
        ((PAGE_W - watermark.width) // 2, (PAGE_H - watermark.height) // 2),
    )
    image.alpha_composite(overlay)


def draw_footer(draw: ImageDraw.ImageDraw) -> None:
    draw.rectangle((0, PAGE_H - FOOTER_H, PAGE_W, PAGE_H), fill=WARM)
    draw.text(
        (MARGIN, PAGE_H - 47),
        "Preview only. Editable Word and Excel files unlock after verified payment.",
        fill=MUTED,
        font=FONT_SMALL,
    )


def draw_section_band(draw: ImageDraw.ImageDraw, y: int, label: str) -> int:
    draw.rounded_rectangle(
        (MARGIN, y, PAGE_W - MARGIN, y + 44),
        radius=8,
        fill=ORANGE_SOFT,
    )
    draw.text((MARGIN + 18, y + 12), label, fill=ORANGE_DARK, font=FONT_SMALL_BOLD)
    return y + 62


def draw_intro_panel(draw: ImageDraw.ImageDraw, y: int, text: str) -> int:
    lines = clipped_lines(draw, text, FONT_BODY, CONTENT_W - 52, 3)
    panel_h = max(94, 38 + (len(lines) * 25))
    draw.rounded_rectangle(
        (MARGIN, y, PAGE_W - MARGIN, y + panel_h),
        radius=8,
        fill=WARM,
        outline=LINE,
        width=1,
    )
    draw.rectangle((MARGIN, y, MARGIN + 7, y + panel_h), fill=ORANGE)
    text_y = y + ((panel_h - (len(lines) * 25)) // 2)
    for line in lines:
        draw.text((MARGIN + 26, text_y), line, fill=INK, font=FONT_BODY)
        text_y += 25
    return y + panel_h


def is_heading_style(style_name: str) -> bool:
    return style_name.casefold().startswith("heading")


def extract_table_labels(document: Document) -> list[str]:
    labels: list[str] = []
    seen: set[str] = set()
    excluded = {
        "N/A",
        "EA",
        "QTY",
        "UNIT",
        "DATE",
        "NAME",
        "SIGNATURE",
        "AMOUNT",
        "DESCRIPTION",
    }

    def add_label(value: str) -> None:
        label = clean_text(value).strip(" |:")
        key = label.casefold()
        if (
            not label
            or key in seen
            or label in excluded
            or "[" in label
            or "]" in label
            or len(label) > 48
            or len(label.split()) > 7
        ):
            return
        seen.add(key)
        labels.append(label)

    for table in document.tables[:12]:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            nonempty = [value for value in cells if value]
            if len(nonempty) >= 2 and len(nonempty[0]) <= 40:
                add_label(nonempty[0])
            for value in nonempty:
                letters = [char for char in value if char.isalpha()]
                uppercase_ratio = (
                    sum(char.isupper() for char in letters) / len(letters)
                    if letters
                    else 0
                )
                if uppercase_ratio >= 0.85:
                    add_label(value)
            if len(labels) >= 10:
                return labels
    return labels


def extract_docx_content(
    source: Path,
    slug: str,
) -> tuple[str, list[tuple[str, str]], list[str]]:
    document = Document(source)
    paragraphs = [
        (clean_text(paragraph.text), paragraph.style.name)
        for paragraph in document.paragraphs
        if clean_text(paragraph.text)
    ]

    intro = DOCX_INTRO_OVERRIDES.get(slug, "")
    if not intro:
        for text, style_name in paragraphs:
            if is_heading_style(style_name) or text.isupper():
                continue
            if 25 <= len(text) <= 230:
                intro = text
                break
    if not intro:
        intro = "Selected editable fields and document sections are shown below."

    sections = list(DOCX_SECTION_OVERRIDES.get(slug, []))
    if not sections:
        for index, (heading, style_name) in enumerate(paragraphs):
            if not is_heading_style(style_name):
                continue
            body_parts: list[str] = []
            for body, body_style in paragraphs[index + 1 :]:
                if is_heading_style(body_style):
                    break
                if body:
                    body_parts.append(body)
                if sum(len(part) for part in body_parts) >= 430:
                    break
            body_text = " ".join(body_parts)
            if not body_text:
                body_text = "Editable particulars and completion fields for this section."
            sections.append((heading, body_text))
            if len(sections) >= 4:
                break

    if not sections:
        for text, _ in paragraphs:
            if ":" not in text:
                continue
            heading, body = text.split(":", 1)
            if 3 <= len(heading) <= 38 and body.strip():
                sections.append((heading, body.strip()))
            if len(sections) >= 4:
                break

    table_labels = extract_table_labels(document)
    while len(sections) < 4 and table_labels:
        label = table_labels.pop(0)
        sections.append((label.title(), "Editable fields are provided in this section."))

    return intro, sections[:4], table_labels[:8]


def draw_section_card(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    number: int,
    title: str,
    body: str,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=8, fill=WHITE, outline=LINE, width=1)
    draw.rounded_rectangle((x1, y1, x1 + 7, y2), radius=4, fill=ORANGE)
    draw.text(
        (x1 + 24, y1 + 18),
        f"SECTION {number:02d}",
        fill=ORANGE_DARK,
        font=FONT_SMALL_BOLD,
    )

    title_lines = clipped_lines(draw, title, FONT_SECTION, (x2 - x1) - 48, 2)
    title_y = y1 + 47
    for line in title_lines:
        draw.text((x1 + 24, title_y), line, fill=BLACK, font=FONT_SECTION)
        title_y += 25

    body_y = title_y + 9
    body_lines = clipped_lines(
        draw,
        body,
        FONT_CELL,
        (x2 - x1) - 48,
        max(2, (y2 - body_y - 18) // 22),
    )
    for line in body_lines:
        draw.text((x1 + 24, body_y), line, fill=MUTED, font=FONT_CELL)
        body_y += 22


def draw_chip_grid(
    draw: ImageDraw.ImageDraw,
    y: int,
    label: str,
    items: list[str],
    columns: int = 2,
) -> int:
    if not items:
        return y
    y = draw_section_band(draw, y, label)
    gap = 14
    chip_w = (CONTENT_W - (gap * (columns - 1))) // columns
    chip_h = 50
    for index, item in enumerate(items):
        row = index // columns
        column = index % columns
        x = MARGIN + (column * (chip_w + gap))
        chip_y = y + (row * (chip_h + 12))
        draw.rounded_rectangle(
            (x, chip_y, x + chip_w, chip_y + chip_h),
            radius=8,
            fill=WHITE,
            outline=LINE,
            width=1,
        )
        draw.ellipse((x + 16, chip_y + 19, x + 28, chip_y + 31), fill=ORANGE)
        item_lines = clipped_lines(draw, item, FONT_CELL_BOLD, chip_w - 58, 1)
        if item_lines:
            draw.text(
                (x + 40, chip_y + 15),
                item_lines[0],
                fill=INK,
                font=FONT_CELL_BOLD,
            )
    rows = (len(items) + columns - 1) // columns
    return y + (rows * (chip_h + 12))


def render_docx_preview(item: dict[str, object]) -> None:
    image = create_canvas()
    draw = ImageDraw.Draw(image)
    y = draw_header(draw, str(item["title"]), "Word")
    intro, sections, table_labels = extract_docx_content(
        Path(str(item["source"])),
        str(item["slug"]),
    )

    y = draw_section_band(draw, y, "Inside this Word template")
    y = draw_intro_panel(draw, y, intro) + 24

    card_gap = 18
    card_w = (CONTENT_W - card_gap) // 2
    card_h = 218
    for index, (title, body) in enumerate(sections):
        row = index // 2
        column = index % 2
        x = MARGIN + (column * (card_w + card_gap))
        card_y = y + (row * (card_h + card_gap))
        draw_section_card(
            draw,
            (x, card_y, x + card_w, card_y + card_h),
            index + 1,
            title,
            body,
        )
    y += (2 * card_h) + card_gap + 24

    labels = table_labels[:6]
    if not labels:
        labels = [title for title, _ in sections[:6]]
    draw_chip_grid(draw, y, "Editable template sections", labels)

    draw_watermark(image)
    draw_footer(draw)
    image.convert("RGB").save(
        OUTPUT_ROOT / f"{item['slug']}.png",
        optimize=True,
    )


def row_values(worksheet, row_number: int, max_col: int = 14) -> list[str]:
    values = [
        display_value(worksheet.cell(row=row_number, column=column).value)
        for column in range(1, min(worksheet.max_column, max_col) + 1)
    ]
    return [value for value in values if value]


def preceding_section_title(worksheet, row_number: int, fallback: str) -> str:
    for candidate_row in range(row_number - 1, max(0, row_number - 8), -1):
        values = row_values(worksheet, candidate_row)
        if not values:
            continue
        value = values[0]
        if len(value) <= 48:
            return value.title() if value.isupper() else value
    return fallback


def extract_xlsx_content(source: Path) -> dict[str, object]:
    workbook = load_workbook(source, data_only=True, read_only=False)
    worksheet = workbook[workbook.sheetnames[0]]
    sheet_names = workbook.sheetnames[:6]

    all_rows = [
        (row_number, row_values(worksheet, row_number))
        for row_number in range(1, min(worksheet.max_row, 45) + 1)
    ]
    nonempty_rows = [(row_number, values) for row_number, values in all_rows if values]
    title = nonempty_rows[0][1][0] if nonempty_rows else source.stem

    subtitle = ""
    for _, values in nonempty_rows[1:]:
        candidate = " ".join(values)
        if 45 <= len(candidate) <= 280:
            subtitle = candidate
            break
    if not subtitle:
        subtitle = "Selected workbook sections and editable fields are shown below."

    metrics: list[tuple[str, str]] = []
    for row_number, values in nonempty_rows:
        next_values = row_values(worksheet, row_number + 1)
        if (
            3 <= len(values) <= 4
            and len(next_values) == len(values)
            and all(len(value) <= 32 for value in values)
        ):
            metrics = list(zip(values, next_values))
            break

    section_title = "Workbook extract"
    table_rows: list[list[str]] = []
    list_items: list[str] = []
    has_header = False

    for row_number, values in nonempty_rows:
        if values[0].casefold() != "month" or len(values) < 3:
            continue
        section_title = "Monthly performance"
        table_rows.append(values[:4])
        for data_row in range(row_number + 1, min(row_number + 9, worksheet.max_row + 1)):
            data_values = row_values(worksheet, data_row)
            if len(data_values) < 2:
                break
            table_rows.append(data_values[:4])
        has_header = True
        break

    if not table_rows:
        for row_number, values in nonempty_rows:
            if not re.fullmatch(r"Step \d+", values[0], flags=re.IGNORECASE):
                continue
            section_title = preceding_section_title(worksheet, row_number, "Quick setup")
            for data_row in range(row_number, min(row_number + 8, worksheet.max_row + 1)):
                data_values = row_values(worksheet, data_row)
                if not data_values or not re.fullmatch(
                    r"Step \d+",
                    data_values[0],
                    flags=re.IGNORECASE,
                ):
                    break
                table_rows.append(data_values[:2])
            break

    if not table_rows:
        setup_row = next(
            (
                row_number
                for row_number, values in nonempty_rows
                if values[0].casefold() == "setup sequence"
            ),
            None,
        )
        if setup_row is not None:
            section_title = "Setup sequence"
            for data_row in range(setup_row + 1, min(setup_row + 12, worksheet.max_row + 1)):
                data_values = row_values(worksheet, data_row)
                if not data_values:
                    continue
                if re.match(r"^\d+\.", data_values[0]):
                    list_items.append(data_values[0])
                elif list_items:
                    break

    if not table_rows and not list_items:
        for index, (row_number, values) in enumerate(nonempty_rows):
            if len(values) < 2:
                continue
            candidate_rows = [values[:4]]
            for _, next_values in nonempty_rows[index + 1 : index + 7]:
                if len(next_values) < 2:
                    break
                candidate_rows.append(next_values[:4])
            if len(candidate_rows) >= 3:
                section_title = preceding_section_title(
                    worksheet,
                    row_number,
                    "Workbook extract",
                )
                table_rows = candidate_rows
                break

    workbook.close()
    return {
        "title": title,
        "subtitle": subtitle,
        "metrics": metrics,
        "section_title": section_title,
        "table_rows": table_rows,
        "list_items": list_items,
        "has_header": has_header,
        "sheet_names": sheet_names,
    }


def calculate_column_widths(
    draw: ImageDraw.ImageDraw,
    rows: list[list[str]],
    total_width: int,
) -> list[int]:
    column_count = max(len(row) for row in rows)
    if column_count == 1:
        return [total_width]

    measured: list[int] = []
    for column in range(column_count):
        widths = []
        for row in rows:
            value = row[column] if column < len(row) else ""
            sample = value[:42]
            widths.append(draw.textbbox((0, 0), sample, font=FONT_CELL)[2] + 30)
        measured.append(max(80, min(max(widths, default=80), 360)))

    if column_count == 2:
        first_width = max(130, min(measured[0], 210))
        return [first_width, total_width - first_width]

    minimum = 120 if column_count <= 4 else 92
    base_total = minimum * column_count
    if base_total >= total_width:
        widths = [total_width // column_count] * column_count
        widths[-1] += total_width - sum(widths)
        return widths

    extra_space = total_width - base_total
    weights = [max(1, width - minimum) for width in measured]
    weight_total = sum(weights)
    widths = [
        minimum + int(extra_space * (weight / weight_total))
        for weight in weights
    ]
    widths[-1] += total_width - sum(widths)
    return widths


def draw_table(
    draw: ImageDraw.ImageDraw,
    y: int,
    rows: list[list[str]],
    has_header: bool,
    bottom_limit: int,
) -> int:
    if not rows:
        return y
    widths = calculate_column_widths(draw, rows, CONTENT_W)
    column_count = len(widths)

    for row_index, row in enumerate(rows):
        is_header = has_header and row_index == 0
        fonts = [FONT_CELL_BOLD if is_header or column == 0 else FONT_CELL for column in range(column_count)]
        line_sets = [
            clipped_lines(
                draw,
                row[column] if column < len(row) else "",
                fonts[column],
                widths[column] - 24,
                3,
            )
            for column in range(column_count)
        ]
        row_h = max(50, 22 + (max((len(lines) for lines in line_sets), default=1) * 21))
        if y + row_h > bottom_limit:
            break

        x = MARGIN
        for column, width in enumerate(widths):
            if is_header:
                fill = BLACK
                text_fill = WHITE
            elif column == 0:
                fill = ORANGE_SOFT
                text_fill = ORANGE_DARK
            else:
                fill = WHITE if row_index % 2 else WARM
                text_fill = INK

            draw.rectangle(
                (x, y, x + width, y + row_h),
                fill=fill,
                outline=LINE,
                width=1,
            )
            lines = line_sets[column]
            text_y = y + ((row_h - (len(lines) * 21)) // 2)
            for line in lines:
                draw.text(
                    (x + 12, text_y),
                    line,
                    fill=text_fill,
                    font=fonts[column],
                )
                text_y += 21
            x += width
        y += row_h
    return y


def draw_metric_cards(
    draw: ImageDraw.ImageDraw,
    y: int,
    metrics: list[tuple[str, str]],
) -> int:
    metrics = metrics[:4]
    if not metrics:
        return y
    gap = 12
    card_w = (CONTENT_W - (gap * 3)) // 4
    card_h = 104
    for index, (label, value) in enumerate(metrics):
        x = MARGIN + (index * (card_w + gap))
        draw.rounded_rectangle(
            (x, y, x + card_w, y + card_h),
            radius=8,
            fill=BLACK,
        )
        label_lines = clipped_lines(draw, label, FONT_SMALL_BOLD, card_w - 24, 2)
        label_y = y + 15
        for line in label_lines:
            draw.text((x + 12, label_y), line, fill=(255, 176, 111), font=FONT_SMALL_BOLD)
            label_y += 18
        draw.text(
            (x + 12, y + 70),
            value,
            fill=WHITE,
            font=FONT_BODY_BOLD,
        )
    return y + card_h


def draw_setup_list(
    draw: ImageDraw.ImageDraw,
    y: int,
    items: list[str],
    bottom_limit: int,
) -> int:
    for index, item in enumerate(items[:8], 1):
        lines = clipped_lines(draw, re.sub(r"^\d+\.\s*", "", item), FONT_CELL, CONTENT_W - 92, 2)
        row_h = max(54, 22 + (len(lines) * 21))
        if y + row_h > bottom_limit:
            break
        fill = WHITE if index % 2 else WARM
        draw.rounded_rectangle(
            (MARGIN, y, PAGE_W - MARGIN, y + row_h),
            radius=7,
            fill=fill,
            outline=LINE,
            width=1,
        )
        number_box = (MARGIN + 14, y + 10, MARGIN + 50, y + 46)
        draw.rounded_rectangle(number_box, radius=7, fill=ORANGE)
        draw_centered_text(draw, number_box, str(index), FONT_SMALL_BOLD, WHITE)
        text_y = y + ((row_h - (len(lines) * 21)) // 2)
        for line in lines:
            draw.text((MARGIN + 68, text_y), line, fill=INK, font=FONT_CELL)
            text_y += 21
        y += row_h + 8
    return y


def render_xlsx_preview(item: dict[str, object]) -> None:
    image = create_canvas()
    draw = ImageDraw.Draw(image)
    y = draw_header(draw, str(item["title"]), "Excel")
    content = extract_xlsx_content(Path(str(item["source"])))

    y = draw_section_band(draw, y, "Inside this Excel workbook")
    y = draw_intro_panel(draw, y, str(content["subtitle"])) + 22

    metrics = list(content["metrics"])
    if metrics:
        y = draw_metric_cards(draw, y, metrics) + 22

    y = draw_section_band(draw, y, str(content["section_title"]))
    bottom_limit = CONTENT_BOTTOM - 210
    table_rows = list(content["table_rows"])
    list_items = list(content["list_items"])
    if table_rows:
        y = draw_table(
            draw,
            y,
            table_rows,
            bool(content["has_header"]),
            bottom_limit,
        )
    else:
        y = draw_setup_list(draw, y, list_items, bottom_limit)

    y = min(y + 24, CONTENT_BOTTOM - 180)
    draw_chip_grid(
        draw,
        y,
        "Workbook sections",
        list(content["sheet_names"]),
        columns=3,
    )

    draw_watermark(image)
    draw_footer(draw)
    image.convert("RGB").save(
        OUTPUT_ROOT / f"{item['slug']}.png",
        optimize=True,
    )


def create_canvas() -> Image.Image:
    return Image.new("RGBA", (PAGE_W, PAGE_H), (*WHITE, 255))


def main() -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

    for item in DOCUMENTS:
        source = Path(str(item["source"]))
        if not source.exists():
            raise FileNotFoundError(source)
        if item["format"] == "Word":
            render_docx_preview(item)
        else:
            render_xlsx_preview(item)
        print(f"Created {OUTPUT_ROOT / f'{item['slug']}.png'}")


if __name__ == "__main__":
    main()
