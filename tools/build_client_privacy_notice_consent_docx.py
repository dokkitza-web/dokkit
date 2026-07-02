from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Client_Privacy_Notice_and_Consent_Template.docx")

FONT = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
CAUTION_GOLD = RGBColor(0x7A, 0x5A, 0x00)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
BORDER = "C9D3DF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn("w:ascii"), name)
    r_pr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, before=0, after=6, line=1.25, keep_next=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next is not None:
        fmt.keep_with_next = keep_next


def style_para(p, size=11, color=BODY, bold=False, italic=False):
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_para(
    doc,
    text="",
    *,
    style=None,
    size=11,
    color=BODY,
    bold=False,
    italic=False,
    align=None,
    before=0,
    after=6,
    line=1.25,
    keep_next=None,
):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=before, after=after, line=line, keep_next=keep_next)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    style_para(
        p,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **margins):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ["top", "start", "bottom", "end"]:
        if margin not in margins:
            continue
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[margin]))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, size=9.5, color=BODY, align=None, fill=None, italic=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    if text is not None:
        for idx, part in enumerate(str(text).split("\n")):
            if idx:
                p.add_run().add_break()
            run = p.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=0, after=0, line=1.15)
    if align is not None:
        p.alignment = align


def set_table_borders(table, color=BORDER, sz="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell, **CELL_MARGIN)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            size=font_size,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=header_fill,
        )
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, text in enumerate(row_data):
            short = len(str(text)) < 18
            align = WD_ALIGN_PARAGRAPH.CENTER if short and idx in (0, len(row_data) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], text, size=font_size, align=align)
    add_para(doc, "", after=3)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT_FILL, title_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D7DBE2", sz="6")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.text = ""
    title_run = p.add_run(title.upper() + "\n")
    set_run_font(title_run, size=10, color=title_color, bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=10.3, color=BODY)
    set_paragraph_format(p, before=0, after=0, line=1.2)
    add_para(doc, "", after=2)
    return table


def add_step_table(doc, steps, *, font_size=8.7):
    rows = [(str(idx), step) for idx, step in enumerate(steps, start=1)]
    return add_table(
        doc,
        ["Step", "Action"],
        rows,
        [850, 8510],
        header_fill=LIGHT_BLUE,
        font_size=font_size,
    )


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def build_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True


def setup_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = ""
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    left, right = table.rows[0].cells
    set_cell_text(left, "POPIA | Client Privacy Notice and Consent", size=8.5, color=MUTED)
    set_cell_text(right, "[Business Name]", size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    footer = section.footer
    footer.paragraphs[0].text = ""
    ftable = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ftable, [6500, 2860], indent_dxa=0)
    lcell, rcell = ftable.rows[0].cells
    set_cell_text(lcell, "Controlled document - verify latest version before use", size=8.5, color=MUTED)
    p = rcell.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" of ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "NUMPAGES")
    set_paragraph_format(p, before=0, after=0, line=1.0)
    for cell in ftable.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)


def add_cover(doc):
    add_para(doc, "", after=48)
    add_para(
        doc,
        "Business Template",
        size=10.5,
        bold=True,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
    )
    add_para(
        doc,
        "Client Privacy Notice and Consent",
        size=29,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.08,
    )
    add_para(
        doc,
        "For a South African Cleaning Services Business",
        size=15,
        color=DARK_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=24,
    )
    add_para(
        doc,
        "Prepared for: [Business Name] | Template date: June 2026",
        size=11.5,
        bold=True,
        color=BODY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=5,
    )
    add_para(
        doc,
        "Use this controlled template to explain how client personal information is collected, used, stored, shared, retained, and protected, and to record consent for optional processing such as direct marketing and public-use photographs.",
        size=10.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=28,
    )
    add_table(
        doc,
        ["Field", "Template Entry"],
        [
            ("Document ID", "[POPIA-CLI-###]"),
            ("Version", "[0.1 / Draft / Approved]"),
            ("Effective date", "[dd/mm/yyyy]"),
            ("Review date", "[dd/mm/yyyy]"),
            ("Responsible party", "[Registered business / sole proprietor / company name]"),
            ("Information Officer", "[Name / role / email / telephone]"),
            ("Applies to", "Prospects, clients, authorised representatives, household occupants, site contacts, and emergency contacts"),
            ("Service area", "[Province / municipality / service areas]"),
        ],
        [2500, 6860],
        font_size=9.4,
    )
    add_callout(
        doc,
        "Template control",
        "Replace bracketed fields before issue. This template is a POPIA-aligned business document, not legal advice. Have the final version checked against the company's actual systems, contracts, operators, marketing channels, and retention schedule.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )
    add_page_break(doc)


def add_document_control(doc):
    add_heading(doc, "1. Document Control", 1)
    add_table(
        doc,
        ["Version", "Date", "Changed by", "Summary of change", "Approved by"],
        [
            ("0.1", "[dd/mm/yyyy]", "[Name]", "Initial template issue", "[Name]"),
            ("", "", "", "", ""),
            ("", "", "", "", ""),
        ],
        [1050, 1400, 1700, 3610, 1600],
        header_fill=LIGHT_GREY,
        font_size=8.9,
    )
    add_table(
        doc,
        ["Controlled copy holder", "Role / area", "Issue method", "Review trigger"],
        [
            ("[Name]", "Information Officer", "[Controlled master / website copy]", "POPIA change, new operator, complaint, security compromise"),
            ("[Name]", "Operations", "[Client pack / booking process]", "New service type, new site access process, photo use change"),
            ("[Name]", "Marketing / Admin", "[CRM / consent register]", "New marketing channel, objection, unsubscribe trend"),
            ("[Name]", "Supervisors", "[Crew pack / mobile app]", "New photo, key, alarm, or access-code procedure"),
        ],
        [2300, 1900, 2600, 2560],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )


def add_use_guidance(doc):
    add_heading(doc, "2. How to Use This Template", 1)
    add_table(
        doc,
        ["Component", "Use"],
        [
            ("Client Privacy Notice", "Give to clients before or at collection of personal information, for example during quote, booking, onboarding, service agreement, website enquiry, or account setup."),
            ("Consent Form", "Use only for optional processing where consent is the right basis, such as direct marketing to non-customers, public-use before/after photos, testimonials, or optional WhatsApp marketing."),
            ("Service Authorisation", "Use for practical service permissions such as access codes, keys, alarm instructions, emergency contacts, pets, and authorised site representatives."),
            ("Withdrawal / Objection Record", "Use when a client withdraws consent, objects to direct marketing, requests correction/deletion, or changes communication preferences."),
            ("Consent Register", "Keep a record of what was consented to, when, by whom, how it was captured, and when it was withdrawn or changed."),
        ],
        [2200, 7160],
        header_fill=LIGHT_BLUE,
        font_size=8.6,
    )
    add_callout(
        doc,
        "Consent discipline",
        "Do not ask for blanket consent for everything. Some processing is necessary for the quote, contract, safety, payment, records, or legal obligations. Use consent mainly where the client has a real choice and can refuse without losing the core cleaning service.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )


def add_compliance_refs(doc):
    add_heading(doc, "3. South African POPIA Reference Points", 1)
    add_para(
        doc,
        "This section is a practical reference for template users. The business should verify its final privacy notice and consent wording with a competent adviser where risk is high or systems are complex.",
        italic=True,
        color=MUTED,
        size=10.4,
    )
    add_table(
        doc,
        ["Reference point", "Operational meaning for a cleaning services business"],
        [
            (
                "Protection of Personal Information Act 4 of 2013",
                "Personal information must be processed lawfully and fairly, for a defined purpose, with appropriate security safeguards and respect for data-subject rights.",
            ),
            (
                "POPIA section 18 notification principle",
                "When collecting client information, tell clients what is collected, who collects it, why it is collected, whether it is voluntary or mandatory, consequences of not providing it, intended cross-border transfer where applicable, recipients, rights, and complaints route.",
            ),
            (
                "POPIA section 11 processing justification",
                "Processing may be based on consent, contract necessity, legal obligation, protection of a legitimate interest, or legitimate interests of the responsible party/third party, depending on the purpose.",
            ),
            (
                "Direct marketing rules and Information Regulator guidance",
                "Treat electronic marketing consent carefully. Maintain opt-in/opt-out records and stop marketing when the client objects or withdraws consent.",
            ),
            (
                "Information Officer duties",
                "The Information Officer should maintain the POPIA compliance framework, handle requests, support awareness, and work with the Information Regulator where needed.",
            ),
            (
                "Official POPIA forms",
                "Keep a route for objection, correction/deletion, direct-marketing consent, and complaints. Internal forms may be simpler but must not block legal rights.",
            ),
        ],
        [3000, 6360],
        header_fill=LIGHT_GREY,
        font_size=8.0,
    )
    add_table(
        doc,
        ["Official verification source", "Link / note"],
        [
            ("POPIA on gov.za", "https://www.gov.za/documents/protection-personal-information-act"),
            ("Information Regulator POPIA resources", "https://inforegulator.org.za/popia/"),
            ("Information Regulator direct marketing guidance", "https://inforegulator.org.za/wp-content/uploads/2020/07/GUIDANCE-NOTE-ON-DIRECT-MARKETING-IN-TERMS-OF-THE-PROTECTION-OF-PERSONAL-INFORMATION-ACT-4-OF-2013-POPIA.pdf"),
            ("Official objection / correction / consent / complaint forms", "Available from the Information Regulator POPIA resources page."),
        ],
        [3000, 6360],
        header_fill=LIGHT_GREY,
        font_size=8.0,
    )


def add_notice_identity(doc):
    add_heading(doc, "4. Client Privacy Notice", 1)
    add_heading(doc, "4.1 Who We Are", 2)
    add_para(
        doc,
        "[Business Name] is the responsible party for personal information collected for quotations, bookings, cleaning services, customer support, safety, billing, records, complaints, and lawful business administration.",
    )
    add_table(
        doc,
        ["Notice field", "Client-facing wording / entry"],
        [
            ("Responsible party", "[Business legal name, registration number if applicable, trading name]"),
            ("Physical address", "[Business address]"),
            ("Information Officer", "[Name, title, email, telephone]"),
            ("General privacy contact", "[Email / phone / WhatsApp business number]"),
            ("Client service contact", "[Bookings / complaints / service desk contact]"),
            ("Effective date", "[dd/mm/yyyy]"),
        ],
        [2500, 6860],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )
    add_heading(doc, "4.2 What Personal Information We Collect", 2)
    add_table(
        doc,
        ["Category", "Examples relevant to cleaning services"],
        [
            ("Identity and contact details", "Name, surname, company name, phone number, email address, ID/passport only where necessary, authorised representative details."),
            ("Service and location details", "Home or site address, unit number, directions, parking, access notes, room/area instructions, cleaning preferences, service history."),
            ("Access and security details", "Keys, gate codes, alarm instructions, guardhouse instructions, access cards, authorised persons, lockbox details, handover notes."),
            ("Household / site context", "Occupants, pets, children/vulnerable persons presence, allergies or sensitivities where voluntarily shared for safety or product selection."),
            ("Financial and transaction records", "Quotes, invoices, payment confirmations, account status, refund details, purchase order details."),
            ("Service evidence", "Before/after photos, damage photos, snag records, complaint notes, sign-off records, incident reports, messages, call notes."),
            ("Marketing preferences", "Opt-in/opt-out choices, preferred channels, campaign history, unsubscribe records."),
            ("Digital information", "Website enquiry data, email metadata, WhatsApp business messages, online form submissions, cookies where applicable."),
        ],
        [2200, 7160],
        header_fill=LIGHT_GREY,
        font_size=8.1,
    )


def add_notice_processing(doc):
    add_heading(doc, "4.3 Why We Use Personal Information", 2)
    add_table(
        doc,
        ["Purpose", "Information used", "Typical POPIA basis", "Required?"],
        [
            ("Quote and booking", "Contact details, address, scope, access notes, service dates.", "Contract steps / legitimate interest", "Usually required"),
            ("Deliver cleaning service", "Site instructions, access/security details, crew notes, service history.", "Contract performance / safety interest", "Required"),
            ("Protect property and safety", "Photos, incident notes, emergency contacts, hazard/sensitivity information.", "Legitimate interest / legal obligation / consent where optional", "Depends on context"),
            ("Billing and account records", "Invoices, payment references, company details, service records.", "Contract / legal obligation", "Required"),
            ("Client support and complaints", "Messages, call notes, service evidence, complaint details.", "Contract / legitimate interest", "Required for support"),
            ("Direct marketing", "Name, contact details, service interest, consent/preference record.", "Consent or permitted existing-customer marketing rules", "Optional"),
            ("Public portfolio or social media", "Before/after images, testimonial, name/logo if agreed.", "Specific consent", "Optional"),
            ("Legal, insurance, audit, dispute", "Contracts, records, photos, incident reports, correspondence.", "Legal obligation / legitimate interest", "Required when relevant"),
        ],
        [1600, 2750, 3000, 2010],
        header_fill=LIGHT_BLUE,
        font_size=7.7,
    )
    add_heading(doc, "4.4 Where We Get Information From", 2)
    add_table(
        doc,
        ["Source", "Examples"],
        [
            ("Client or authorised representative", "Quote forms, calls, WhatsApp messages, emails, booking forms, service agreements, sign-off forms."),
            ("Site contact / landlord / managing agent", "Access arrangements, building rules, parking, security, emergency contacts, work-area restrictions."),
            ("Crew and supervisor records", "Job cards, service notes, photos if authorised, damage reports, complaints, incidents, quality checks."),
            ("Payment and admin systems", "Invoice records, payment confirmations, account status, collection communications."),
            ("Public sources only where appropriate", "Business contact details, public company information, publicly listed addresses for commercial enquiries."),
        ],
        [2500, 6860],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )


def add_sharing_retention(doc):
    add_heading(doc, "4.5 Who We Share Information With", 2)
    add_table(
        doc,
        ["Recipient / operator", "Reason for sharing", "Control note"],
        [
            ("Cleaning teams and supervisors", "To deliver the service, access the site, follow instructions, report snags, and keep clients safe.", "Share only what the team needs for the job."),
            ("Admin, finance, and client support", "Bookings, invoicing, collections, complaints, and service records.", "Limit access to approved staff."),
            ("Subcontractors / temporary staff", "Service delivery where approved by the business and client contract.", "Use confidentiality and operator controls."),
            ("IT, cloud, CRM, accounting, messaging providers", "Hosting, storage, communications, scheduling, payments, backups, and support.", "Use reputable providers and restrict access."),
            ("Insurers, advisers, legal representatives", "Claims, disputes, serious incidents, debt recovery, or legal advice.", "Share relevant records only."),
            ("Authorities / regulators", "Where required by law or lawful process.", "Verify request and record disclosure."),
            ("Client-authorised contacts", "Representatives, managing agents, site contacts, emergency contacts.", "Confirm authority before sharing sensitive details."),
        ],
        [2200, 4300, 2860],
        header_fill=LIGHT_BLUE,
        font_size=8.0,
    )
    add_heading(doc, "4.6 Storage, Security, Cross-Border Transfers, and Retention", 2)
    add_table(
        doc,
        ["Control area", "Client-facing wording / business control"],
        [
            ("Security", "We use reasonable technical and organisational safeguards, including access controls, staff instructions, secure filing, device controls, approved communication channels, and confidentiality requirements."),
            ("Access codes and keys", "Gate codes, alarm instructions, keys, access cards, and lockbox details are restricted to authorised staff and used only for agreed service access."),
            ("Photos and media", "Service photos are used for quality, proof of work, damage, claims, or public portfolio only where permitted. Public-use photos require separate consent."),
            ("Cross-border / cloud", "Some providers may store or support information outside South Africa. We use providers that offer appropriate protection or contractual safeguards where applicable."),
            ("Retention", "We keep personal information only as long as needed for service delivery, records, tax/accounting, legal, insurance, dispute, safety, and legitimate business purposes."),
            ("Deletion / de-identification", "When records no longer need to identify a client, they should be securely deleted, destroyed, or de-identified according to the retention schedule."),
        ],
        [2300, 7060],
        header_fill=LIGHT_GREY,
        font_size=8.1,
    )


def add_rights_and_requests(doc):
    add_heading(doc, "4.7 Client Rights and Complaints", 2)
    add_para(
        doc,
        "Clients may ask what personal information we hold, request correction or deletion where appropriate, object to certain processing, withdraw consent for optional processing, unsubscribe from direct marketing, and complain to the Information Regulator.",
    )
    add_table(
        doc,
        ["Request type", "Business response standard"],
        [
            ("Access request", "Verify identity/authority, locate relevant records, respond through the Information Officer or approved process."),
            ("Correction / deletion", "Correct inaccurate information or consider deletion/destruction where POPIA and retention obligations allow."),
            ("Objection", "Record the objection, stop processing where required, and confirm any processing that must continue for legal/contractual reasons."),
            ("Consent withdrawal", "Update the consent register and affected systems. Withdrawal does not undo lawful processing already completed."),
            ("Marketing unsubscribe", "Stop the affected marketing channel promptly and keep a suppression record so the client is not re-added."),
            ("Complaint", "Acknowledge, investigate, respond, and advise the client of the Information Regulator route where appropriate."),
        ],
        [2600, 6760],
        header_fill=LIGHT_BLUE,
        font_size=8.4,
    )
    add_table(
        doc,
        ["Information Regulator contact", "Template entry"],
        [
            ("Website", "https://inforegulator.org.za/"),
            ("General enquiries", "enquiries@inforegulator.org.za"),
            ("Telephone", "010 023 5200"),
            ("Complaint route", "Use the latest POPIA complaint forms and instructions from the Information Regulator."),
        ],
        [2800, 6560],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )


def add_client_notice_short(doc):
    add_page_break(doc)
    add_heading(doc, "Appendix A: Client-Facing Privacy Notice", 1)
    add_para(
        doc,
        "Use this appendix as the client-facing notice in quotes, onboarding packs, booking forms, website privacy pages, or service agreements. Replace bracketed fields before issue.",
        italic=True,
        color=MUTED,
        size=10.2,
    )
    add_table(
        doc,
        ["Notice item", "Client-facing wording"],
        [
            ("Who we are", "[Business Name] is a South African cleaning services business. We collect and use personal information to quote, book, deliver, manage, improve, and support our cleaning services."),
            ("Information we collect", "We may collect your name, contact details, address/site details, access instructions, service preferences, payment records, communication records, complaints, incident details, and service photos where relevant."),
            ("Why we collect it", "We use this information to provide quotations, schedule and deliver services, manage access, protect client property and staff safety, invoice and collect payment, handle complaints, keep records, and comply with legal or insurance requirements."),
            ("Optional consent items", "Direct marketing, public-use before/after photos, testimonials, and social-media use are optional and require your separate consent where applicable."),
            ("Sharing", "We may share information with authorised staff, supervisors, subcontractors, IT/CRM/accounting providers, insurers, advisers, and authorities where necessary and lawful."),
            ("Storage and retention", "We protect information using reasonable safeguards and keep it only for as long as needed for the purposes explained in this notice or required by law."),
            ("Your rights", "You may request access, correction, deletion where appropriate, object to certain processing, withdraw optional consent, or unsubscribe from direct marketing."),
            ("Contact", "Information Officer: [Name, email, phone]. You may also complain to the Information Regulator at https://inforegulator.org.za/."),
        ],
        [2100, 7260],
        header_fill=LIGHT_BLUE,
        font_size=8.0,
    )


def add_consent_form(doc):
    add_page_break(doc)
    add_heading(doc, "Appendix B: Client Consent and Service Authorisation Form", 1)
    add_table(
        doc,
        ["Client field", "Entry"],
        [
            ("Client / representative name", "[Name and surname / company representative]"),
            ("Client company / household", "[Company / household / managing agent]"),
            ("Service address / site", "[Address / unit / area]"),
            ("Contact details", "[Phone / email / WhatsApp]"),
            ("Job / account number", "[Reference]"),
            ("Date", "[dd/mm/yyyy]"),
            ("Business representative", "[Name / role]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )
    add_heading(doc, "B1. Mandatory Service Processing Acknowledgement", 2)
    add_para(
        doc,
        "I acknowledge that [Business Name] must process certain personal information to quote, book, deliver, invoice, support, and keep records for the cleaning service. If I do not provide required service information, the business may not be able to provide the requested service safely or at all.",
        size=10.4,
    )
    add_heading(doc, "B2. Optional Consent Choices", 2)
    add_table(
        doc,
        ["Optional processing", "Consent choice", "Client initials"],
        [
            ("Send me direct marketing about cleaning services, promotions, packages, and reminders by email.", "[ ] Yes   [ ] No", ""),
            ("Send me direct marketing by SMS or WhatsApp.", "[ ] Yes   [ ] No", ""),
            ("Contact me by WhatsApp for service updates, reminders, access coordination, and support.", "[ ] Yes   [ ] No", ""),
            ("Use before/after photos of my property for internal training and quality examples with identifying details removed where practical.", "[ ] Yes   [ ] No", ""),
            ("Use before/after photos or testimonials publicly on website, proposals, social media, or marketing material.", "[ ] Yes   [ ] No", ""),
            ("Store access instructions, gate codes, alarm notes, lockbox details, and authorised-person notes for repeat services.", "[ ] Yes   [ ] No   [ ] N/A", ""),
            ("Contact the emergency or alternate contact listed below if needed for access, safety, property risk, or urgent service issues.", "[ ] Yes   [ ] No   [ ] N/A", ""),
        ],
        [5200, 2600, 1560],
        header_fill=LIGHT_GREY,
        font_size=8.1,
    )
    add_callout(
        doc,
        "Client choice",
        "Consent is voluntary for optional items. The client may refuse optional consent and still receive the core cleaning service where the service can be delivered without that optional processing.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )
    add_page_break(doc)
    add_heading(doc, "B3. Emergency Contacts and Sign-Off", 2)
    add_para(
        doc,
        "Complete this section where access, safety, property protection, or urgent service communication may require an authorised alternate contact.",
        size=10.3,
        color=BODY,
    )
    add_table(
        doc,
        ["Emergency / authorised contact", "Name", "Relationship / role", "Phone / email"],
        [
            ("Alternate contact", "", "", ""),
            ("Emergency contact", "", "", ""),
            ("Site / building contact", "", "", ""),
        ],
        [2600, 2300, 2300, 2160],
        header_fill=LIGHT_BLUE,
        font_size=8.3,
    )
    add_para(
        doc,
        "I confirm that I have read the privacy notice, understand the optional consent choices above, and have authority to provide the service instructions and contact details recorded in this form.",
        size=10.3,
    )
    add_table(
        doc,
        ["Signature field", "Entry"],
        [
            ("Client signature", "[Signature]"),
            ("Name and capacity", "[Name / role / authority to sign]"),
            ("Date and place", "[dd/mm/yyyy] / [place]"),
            ("Business witness / representative", "[Name / signature]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_GREY,
        font_size=8.7,
    )


def add_direct_marketing_register(doc):
    add_page_break(doc)
    add_heading(doc, "Appendix C: Direct Marketing Consent Register", 1)
    add_table(
        doc,
        ["Client", "Channel", "Consent source", "Date", "Status", "Withdrawn / objection date", "Updated by"],
        [
            ("", "[Email / SMS / WhatsApp / phone]", "[Form / online / call]", "", "[Active / withdrawn]", "", ""),
            ("", "[Email / SMS / WhatsApp / phone]", "[Form / online / call]", "", "[Active / withdrawn]", "", ""),
            ("", "[Email / SMS / WhatsApp / phone]", "[Form / online / call]", "", "[Active / withdrawn]", "", ""),
            ("", "[Email / SMS / WhatsApp / phone]", "[Form / online / call]", "", "[Active / withdrawn]", "", ""),
            ("", "[Email / SMS / WhatsApp / phone]", "[Form / online / call]", "", "[Active / withdrawn]", "", ""),
            ("", "[Email / SMS / WhatsApp / phone]", "[Form / online / call]", "", "[Active / withdrawn]", "", ""),
        ],
        [1500, 1500, 1700, 1050, 1400, 1400, 810],
        header_fill=LIGHT_BLUE,
        font_size=7.6,
    )
    add_heading(doc, "Appendix D: Photo / Video / Testimonial Permission", 1)
    add_table(
        doc,
        ["Permission item", "Client choice", "Notes / limits"],
        [
            ("Internal job evidence for quality, damage, snag, or proof-of-work records.", "[ ] Yes   [ ] No", ""),
            ("Internal staff training or method improvement, with identifying details removed where practical.", "[ ] Yes   [ ] No", ""),
            ("Public before/after image use on website or proposals.", "[ ] Yes   [ ] No", ""),
            ("Public social-media use.", "[ ] Yes   [ ] No", ""),
            ("Use of client name, company name, or logo with testimonial.", "[ ] Yes   [ ] No", ""),
            ("Areas that may not be photographed or published.", "[List restrictions]", ""),
        ],
        [4600, 2100, 2660],
        header_fill=LIGHT_GREY,
        font_size=8.1,
    )


def add_request_and_incident_forms(doc):
    add_page_break(doc)
    add_heading(doc, "Appendix E: Withdrawal, Objection, or Change Request", 1)
    add_table(
        doc,
        ["Request field", "Entry"],
        [
            ("Client name and contact", "[Name / phone / email]"),
            ("Request type", "[Withdraw consent / marketing objection / correction / deletion / access / complaint / change preferences]"),
            ("Details of request", "[What must change or stop? Which channel or record?]"),
            ("Identity / authority verified", "[Yes / No / method]"),
            ("Systems updated", "[CRM / WhatsApp / email platform / job file / consent register]"),
            ("Response sent to client", "[Date / method / summary]"),
            ("Handled by", "[Name / role / date]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )
    add_heading(doc, "Appendix F: Privacy Event / Security Compromise Log", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Date and time identified", "[Details]"),
            ("Reported by", "[Name / role]"),
            ("Event type", "[Lost device / wrong recipient / lost keys/access code / unauthorised photo / system issue / other]"),
            ("Information affected", "[Client name, address, access code, photos, invoice, message, etc.]"),
            ("Immediate containment", "[Access disabled, recipient contacted, device locked, code changed, photo removed]"),
            ("Client / regulator notification assessment", "[Decision, owner, date]"),
            ("Corrective action", "[Action, owner, due date]"),
            ("Closed by", "[Name / role / date]"),
        ],
        [2700, 6660],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )


def add_internal_checklists(doc):
    add_page_break(doc)
    add_heading(doc, "Appendix G: Staff Privacy Handling Checklist", 1)
    add_table(
        doc,
        ["Control", "Required practice", "Done"],
        [
            ("Collection", "Collect only the information needed for the quote, booking, access, safety, billing, and support purpose.", "[ ]"),
            ("Access details", "Do not share keys, alarm codes, gate codes, addresses, or client instructions outside the authorised job team.", "[ ]"),
            ("Photos", "Take photos only where authorised or necessary for job evidence, damage, snag, or incident records.", "[ ]"),
            ("Messaging", "Use approved business channels. Do not move client records to personal phones or social media.", "[ ]"),
            ("Marketing", "Check consent or permitted-marketing status before sending promotional messages.", "[ ]"),
            ("Subcontractors", "Confirm subcontractors receive only job-needed details and understand confidentiality expectations.", "[ ]"),
            ("Retention", "File records in the approved location. Do not keep duplicate client data in uncontrolled folders.", "[ ]"),
            ("Incidents", "Report lost devices, wrong-recipient messages, lost keys, unauthorised photos, or access-code exposure immediately.", "[ ]"),
        ],
        [1700, 6460, 1200],
        header_fill=LIGHT_BLUE,
        font_size=8.0,
    )
    add_heading(doc, "Appendix H: Operator / Supplier Privacy Check", 1)
    add_table(
        doc,
        ["Supplier / operator", "Information accessed", "Controls checked", "Review date"],
        [
            ("CRM / booking system", "[Client contacts, addresses, job notes]", "[Contract / access / backups / location / support]", ""),
            ("Accounting system", "[Invoices, payments, company details]", "[Access / backups / retention / support]", ""),
            ("WhatsApp / email / SMS provider", "[Messages, contacts, marketing preferences]", "[Business account / opt-out / access]", ""),
            ("Cloud storage", "[Photos, job cards, signed forms]", "[Access / sharing / backups / deletion]", ""),
            ("Subcontractor / labour broker", "[Job-needed client/site details]", "[Confidentiality / training / return or deletion]", ""),
            ("Other", "", "", ""),
        ],
        [2200, 2800, 3160, 1200],
        header_fill=LIGHT_GREY,
        font_size=8.0,
    )
    add_para(
        doc,
        "End of template. Adapt this document with the business owner, Information Officer, operations lead, client-service lead, marketing/admin owner, IT provider, and competent POPIA adviser.",
        italic=True,
        color=MUTED,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=8,
    )


def build_document():
    doc = Document()
    build_styles(doc)
    setup_header_footer(doc)

    props = doc.core_properties
    props.title = "Client Privacy Notice and Consent Template"
    props.subject = "POPIA-aligned client privacy notice and consent template for a South African cleaning services business"
    props.author = "OpenAI Codex"
    props.keywords = "POPIA, privacy notice, consent, direct marketing, cleaning services, South Africa"

    add_cover(doc)
    add_document_control(doc)
    add_use_guidance(doc)
    add_compliance_refs(doc)
    add_notice_identity(doc)
    add_notice_processing(doc)
    add_sharing_retention(doc)
    add_rights_and_requests(doc)
    add_client_notice_short(doc)
    add_consent_form(doc)
    add_direct_marketing_register(doc)
    add_request_and_incident_forms(doc)
    add_internal_checklists(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT.resolve())
