from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Staff_Onboarding_and_Training_Record_Template.docx")

FONT = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
CAUTION_GOLD = RGBColor(0x7A, 0x5A, 0x00)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
BORDER = "C9D3DF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn("w:ascii"), name)
    r_pr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, before=0, after=6, line=1.25, keep_next=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next is not None:
        fmt.keep_with_next = keep_next


def style_para(p, size=11, color=BODY, bold=False, italic=False):
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_para(
    doc,
    text="",
    *,
    style=None,
    size=11,
    color=BODY,
    bold=False,
    italic=False,
    align=None,
    before=0,
    after=6,
    line=1.25,
    keep_next=None,
):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=before, after=after, line=line, keep_next=keep_next)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    style_para(
        p,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **margins):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ["top", "start", "bottom", "end"]:
        if margin not in margins:
            continue
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[margin]))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, size=9.3, color=BODY, align=None, fill=None, italic=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    if text is not None:
        for idx, part in enumerate(str(text).split("\n")):
            if idx:
                p.add_run().add_break()
            run = p.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=0, after=0, line=1.15)
    if align is not None:
        p.alignment = align


def set_table_borders(table, color=BORDER, sz="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell, **CELL_MARGIN)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=8.9):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            size=font_size,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=header_fill,
        )
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if value in {"[ ]", "Yes / No / N/A"} else None
            set_cell_text(cells[idx], value, size=font_size, align=align)
        prevent_row_split(table.rows[-1])
    add_para(doc, "", after=3)
    return table


def add_label_table(doc, rows, label_width=2700, font_size=9.3):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [label_width, CONTENT_WIDTH_DXA - label_width])
    set_table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=NAVY, fill=LIGHT_GREY, size=font_size)
        set_cell_text(cells[1], value, size=font_size)
        prevent_row_split(table.rows[-1])
    add_para(doc, "", after=3)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT_FILL, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D5DCE5", sz="4")
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(title)
    set_run_font(run, size=10.3, color=title_color, bold=True)
    p.add_run().add_break()
    run = p.add_run(body)
    set_run_font(run, size=9.6, color=BODY)
    set_paragraph_format(p, before=0, after=0, line=1.2)
    add_para(doc, "", after=5)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.3, color=BODY)
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.3)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_headers_and_footers(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    left = hp.add_run("HR | Staff Onboarding and Training Record")
    set_run_font(left, size=9, color=MUTED)
    right = hp.add_run("    [Business Name]")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    left = fp.add_run("Controlled document - verify latest version before use")
    set_run_font(left, size=9, color=MUTED)


def add_cover(doc):
    add_para(doc, "Human Resources and Operations Template", size=10, color=CAUTION_GOLD, bold=True, after=0)
    add_para(
        doc,
        "Staff Onboarding and Training Record",
        size=27,
        color=NAVY,
        bold=True,
        after=6,
        line=1.1,
    )
    add_para(
        doc,
        "For a South African cleaning services business",
        size=13,
        color=MUTED,
        after=18,
    )
    add_label_table(
        doc,
        [
            ("Business name", "[Business Name]"),
            ("Employee name", "[Full name]"),
            ("Employee number", "[Employee / payroll number]"),
            ("Role / grade", "[Cleaner / Team Leader / Supervisor / Driver / Other]"),
            ("Primary site / client", "[Residential / commercial / mobile route / client site]"),
            ("Start date", "[YYYY-MM-DD]"),
            ("Supervisor / trainer", "[Name and role]"),
            ("Record status", "[Draft / active / completed / refresher due]"),
        ],
    )
    add_callout(
        doc,
        "Purpose",
        "Use this template to capture onboarding, legal induction, task training, "
        "competence checks, PPE issue, and refresher records for cleaning staff. "
        "Adapt it to the employee's duties, client site rules, product SDSs, "
        "risk assessments, employment contract, and any applicable bargaining council or sector rules.",
    )


def add_document_control(doc):
    add_heading(doc, "1. Document Control", 1)
    add_table(
        doc,
        ["Control field", "Template entry", "Notes"],
        [
            ("Document number", "HR-ONB-[000]", "Allocate one record per employee."),
            ("Version", "[1.0]", "Update when onboarding or training content changes."),
            ("Effective date", "[YYYY-MM-DD]", "Use the date the template is approved for use."),
            ("Review cycle", "At least annually or when law, client rules, products, or tasks change.", "Assign to HR / Operations."),
            ("Owner", "[HR / Operations Manager / Information Officer]", "Confirm who stores the signed record."),
            ("Retention", "[Company retention schedule]", "Keep employment and training records securely and only as long as lawful and necessary."),
        ],
        [2100, 4100, 3160],
    )
    add_heading(doc, "2. How to Use This Record", 1)
    add_bullet(doc, "Complete the employee and employment fields before the first shift or site deployment.")
    add_bullet(doc, "Use the legal induction, OHS, chemical, equipment, and task matrices as evidence of instruction and competence.")
    add_bullet(doc, "Record only necessary personal information. Keep medical details with the occupational health practitioner where applicable.")
    add_bullet(doc, "Keep this template as a business record. It is not legal advice and should be checked against current South African law and client requirements.")


def add_compliance_reference(doc):
    add_heading(doc, "3. South African Compliance Reference Points", 1)
    add_table(
        doc,
        ["Reference", "Why it matters for onboarding", "Template control"],
        [
            (
                "Occupational Health and Safety Act 85 of 1993",
                "Frames the employer's duty to provide and maintain a working environment that is safe and without risk to health, as far as reasonably practicable.",
                "OHS induction, risk controls, PPE, safe work procedures, incident escalation.",
            ),
            (
                "Hazardous Chemical Agents Regulations, 2021",
                "Cleaning staff may use detergents, disinfectants, degreasers, descalers, and other chemical agents. Training should cover SDS access, labels, PPE, storage, spills, hygiene, and disposal.",
                "Chemical training matrix, SDS sign-off, PPE issue, product authorisation.",
            ),
            (
                "Basic Conditions of Employment Act and National Minimum Wage Act",
                "Employment terms, working time, leave, wage compliance, and payslip/record expectations must be explained through HR documents.",
                "Day-one HR induction confirms the employee received contract, hours, pay, leave, overtime, and workplace rules.",
            ),
            (
                "Skills Development Act 97 of 1998",
                "Supports workplace skills development, learnerships, and training records.",
                "Training plan, attendance record, competence outcome, refresher register.",
            ),
            (
                "Compensation for Occupational Injuries and Diseases Act 130 of 1993",
                "Work injuries and occupational diseases must be reported and supported with appropriate records.",
                "Incident reporting briefing and employee acknowledgement.",
            ),
            (
                "Protection of Personal Information Act 4 of 2013",
                "Employee and client information must be processed lawfully, securely, and only for legitimate business purposes.",
                "Confidentiality, client privacy, records access, and secure storage controls.",
            ),
        ],
        [2400, 4050, 2910],
    )
    add_callout(
        doc,
        "Current-rate caution",
        "South Africa's National Minimum Wage is reviewed annually. If the template is used to confirm pay induction, verify the latest rate and any contract cleaning sector rules before issue rather than relying on an old stored copy.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )


def add_employee_record(doc):
    add_page_break(doc)
    add_heading(doc, "4. Employee and Employment Record", 1)
    add_heading(doc, "4.1 Employee Details", 2)
    add_label_table(
        doc,
        [
            ("Full legal name", "[Name and surname]"),
            ("Preferred name", "[Preferred name]"),
            ("ID / passport number", "[Record securely]"),
            ("Contact number / email", "[Phone / email]"),
            ("Residential address", "[Address]"),
            ("Emergency contact", "[Name / relationship / phone]"),
            ("Language preference", "[English / isiZulu / Afrikaans / Sesotho / other]"),
            ("Employment type", "[Permanent / fixed-term / casual / temporary / learnership]"),
            ("Work location", "[Office / depot / client site / route]"),
        ],
    )
    add_heading(doc, "4.2 Administrative Onboarding Checklist", 2)
    add_table(
        doc,
        ["Item", "Evidence or action required", "Done", "Date / owner"],
        [
            ("Identity and work eligibility", "ID/passport/right-to-work check completed and stored securely.", "[ ]", "[Date / owner]"),
            ("Employment contract", "Signed contract and role description issued before work starts.", "[ ]", "[Date / owner]"),
            ("Pay and banking", "Bank details, tax details, payroll setup, pay cycle explained.", "[ ]", "[Date / owner]"),
            ("Emergency contact", "Emergency contact captured and verified.", "[ ]", "[Date / owner]"),
            ("Uniform and PPE sizes", "Uniform sizes, shoe size, glove size, and any PPE needs recorded.", "[ ]", "[Date / owner]"),
            ("Policies issued", "Code of conduct, attendance, leave, disciplinary, grievance, anti-harassment, and substance rules explained.", "[ ]", "[Date / owner]"),
            ("Privacy and confidentiality", "Employee privacy notice, client confidentiality, photo/phone rules, and records access explained.", "[ ]", "[Date / owner]"),
            ("Medical / fit-for-duty note", "If risk-based medical surveillance applies, record certificate outcome only; avoid storing unnecessary medical details.", "[ ]", "[Date / owner]"),
            ("Driver / vehicle documents", "Driver licence, PDP if applicable, vehicle allocation, route rules, and insurance conditions checked.", "[ ]", "[Date / owner]"),
        ],
        [2400, 4550, 950, 1460],
        font_size=8.7,
    )


def add_day_one_induction(doc):
    add_heading(doc, "5. Day-One HR and Site Induction", 1)
    add_table(
        doc,
        ["Induction topic", "Employee must understand", "Initial / date", "Evidence or notes"],
        [
            ("Role and reporting line", "Job duties, supervisor, escalation route, probation expectations, and performance standards.", "[Initials / date]", "[Notes]"),
            ("Working time and attendance", "Start/finish times, timekeeping, breaks, overtime approval, absenteeism reporting, transport or route meeting point.", "[Initials / date]", "[Notes]"),
            ("Pay, leave, and records", "Pay cycle, payslip access, leave process, deductions, and where HR records are kept.", "[Initials / date]", "[Notes]"),
            ("Professional conduct", "Respectful communication, client property care, no unauthorised guests, no alcohol/drugs, no conflict at client sites.", "[Initials / date]", "[Notes]"),
            ("Uniform and presentation", "Uniform issue, name badge, personal hygiene, jewellery limits, and replacement process.", "[Initials / date]", "[Notes]"),
            ("Client access and keys", "Keys, gate codes, alarm codes, lockboxes, remote controls, and access cards must be protected and returned.", "[Initials / date]", "[Notes]"),
            ("Mobile service setup", "Vehicle loading, job card review, equipment check, arrival procedure, and pack-down responsibilities.", "[Initials / date]", "[Notes]"),
            ("Incident and near-miss reporting", "Report injury, property damage, chemical spill, security concern, client complaint, harassment, or near miss immediately.", "[Initials / date]", "[Notes]"),
        ],
        [2300, 4400, 1300, 1360],
        font_size=8.7,
    )


def add_ohs_training(doc):
    add_page_break(doc)
    add_heading(doc, "6. OHS Induction and Baseline Training Matrix", 1)
    add_callout(
        doc,
        "Risk-based rule",
        "Mark a topic as complete only when the employee has received instruction, had a chance to ask questions, and can demonstrate safe behaviour. If the employee is not yet competent, restrict the task or require supervision until signed off.",
    )
    add_table(
        doc,
        ["Topic", "Must cover", "Competence evidence", "Date / trainer"],
        [
            ("OHS duties", "Employee duties, supervisor instructions, health and safety representative route, stop-work escalation.", "Verbal Q&A / acknowledgement", "[Date / trainer]"),
            ("Site risk assessment", "Client/site hazards, wet floors, stairs, pets, public areas, security, restricted areas, and site-specific rules.", "Can identify top hazards before work", "[Date / trainer]"),
            ("PPE", "Correct PPE selection, inspection, wearing, cleaning, storage, replacement, and reporting defects.", "Demonstrates correct use", "[Date / trainer]"),
            ("Slips, trips, and falls", "Wet-floor signs, cable management, stairs, ladders/step stools, uneven surfaces, lighting, and weather.", "Sets up safe work area", "[Date / trainer]"),
            ("Manual handling", "Safe lifting, team lifts, trolley use, avoiding overreach, moving furniture only when authorised.", "Demonstrates safe lift", "[Date / trainer]"),
            ("Electrical safety", "Vacuum cables, plugs, extension leads, water near electricity, switch-off, defect reporting.", "Pre-use inspection", "[Date / trainer]"),
            ("Chemical safety", "SDS access, labels, dilution, never mixing chemicals, ventilation, PPE, storage, transport, spills, first aid.", "SDS and label check", "[Date / trainer]"),
            ("Biological contaminants", "Blood/body fluids, vomit, mould, sharps, pest waste, toilet areas, gloves, disposal, and escalation.", "Explains no-touch/escalation rules", "[Date / trainer]"),
            ("Waste handling", "General waste, chemical containers, contaminated waste, sharps, client recycling rules, bag weight, and hand hygiene.", "Correct segregation", "[Date / trainer]"),
            ("Emergency response", "Fire, evacuation, first aid, panic button/security, chemical splash, spill, vehicle accident, and emergency contacts.", "Explains response route", "[Date / trainer]"),
            ("Working alone / after hours", "Check-in rules, route plan, emergency contact, no unauthorised access, personal safety.", "Confirms check-in method", "[Date / trainer]"),
        ],
        [2100, 4370, 1750, 1140],
        font_size=8.5,
    )


def add_cleaning_skills_matrix(doc):
    add_heading(doc, "7. Cleaning Skills Training Matrix", 1)
    add_table(
        doc,
        ["Module", "Employee must demonstrate", "Initial", "Result", "Assessor"],
        [
            ("Pre-clean inspection", "Review job card, identify fragile items, access restrictions, risks, missing supplies, and client instructions.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("Colour coding", "Use correct cloths/mops by area to prevent cross-contamination; keep clean and dirty items separate.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("Kitchen cleaning", "Degreasing, surfaces, appliances, splashbacks, sinks, bins, safe food-area practices.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("Bathroom cleaning", "Toilets, basins, showers, mirrors, taps, limescale, disinfection contact time, ventilation.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("Bedroom/living areas", "Dusting, vacuuming, mopping, linen handling if applicable, respectful handling of personal items.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("High-touch points", "Door handles, switches, remotes, rails, taps, appliance handles, desks, and shared surfaces.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("Floors and carpets", "Sweep/vacuum/mop sequence, product choice, wet-floor controls, stain escalation, equipment care.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("Windows and glass", "Safe reach limits, streak-free method, no unsafe climbing, report cracks/damage.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("Quality inspection", "Final walk-through, missed-item correction, job photo rules, supervisor checklist, client sign-off.", "[Date]", "[C / NS / N/A]", "[Name]"),
            ("Pack down", "Product caps closed, equipment cleaned, waste removed, keys returned, job card updated.", "[Date]", "[C / NS / N/A]", "[Name]"),
        ],
        [2100, 4300, 950, 1000, 1010],
        font_size=8.5,
    )
    add_para(doc, "Legend: C = competent, NS = needs supervision, N/A = not applicable.", size=9.2, color=MUTED, italic=True, after=2)


def add_chemical_equipment(doc):
    add_page_break(doc)
    add_heading(doc, "8. Chemical and Equipment Authorisation", 1)
    add_heading(doc, "8.1 Chemical Product Training Record", 2)
    add_table(
        doc,
        ["Product / group", "Training evidence", "PPE / dilution / task limits", "Authorised by", "Review"],
        [
            ("All-purpose cleaner", "SDS/label explained; no mixing; correct storage.", "[PPE / dilution / tasks]", "[Name / date]", "[Date]"),
            ("Disinfectant / sanitiser", "Contact time, surfaces, ventilation, skin/eye protection.", "[PPE / dilution / tasks]", "[Name / date]", "[Date]"),
            ("Degreaser", "Dilution, gloves, eye protection, food-area rinse where required.", "[PPE / dilution / tasks]", "[Name / date]", "[Date]"),
            ("Bathroom descaler / acid product", "Acid risks, never mix with bleach/ammonia, ventilation, splash response.", "[PPE / dilution / tasks]", "[Name / date]", "[Date]"),
            ("Bleach / chlorine product", "No mixing, dilution, ventilation, splash response, storage away from acids.", "[PPE / dilution / tasks]", "[Name / date]", "[Date]"),
            ("Specialist products", "[Oven / carpet / mould / floor treatment / other]", "[PPE / dilution / tasks]", "[Name / date]", "[Date]"),
        ],
        [1900, 3100, 2200, 1080, 1080],
        font_size=8.4,
    )
    add_heading(doc, "8.2 Equipment Training and Authorisation", 2)
    add_table(
        doc,
        ["Equipment", "Pre-use check and safe use", "Authorised tasks", "Employee sign", "Supervisor sign"],
        [
            ("Vacuum cleaner", "Cable, plug, filter, bag/bin, attachments, trip control, storage.", "[Tasks]", "[Initial/date]", "[Initial/date]"),
            ("Mop system / bucket", "Correct mop head, wringer, wet-floor controls, cleaning after use.", "[Tasks]", "[Initial/date]", "[Initial/date]"),
            ("Step stool / small ladder", "Only if trained and authorised; stable surface, three points contact, no overreach.", "[Tasks]", "[Initial/date]", "[Initial/date]"),
            ("Steam cleaner", "Heat/burn risk, electrical safety, surface suitability, cooling and storage.", "[Tasks]", "[Initial/date]", "[Initial/date]"),
            ("Scrubber / polisher", "Machine-specific training, cord control, pads, chemical compatibility, client/site permission.", "[Tasks]", "[Initial/date]", "[Initial/date]"),
            ("Pressure washer", "Only trained staff; high-pressure risk, electrical/water controls, bystander exclusion.", "[Tasks]", "[Initial/date]", "[Initial/date]"),
            ("Company vehicle", "Licence check, route rules, load securing, accident reporting, fuel/card rules.", "[Tasks]", "[Initial/date]", "[Initial/date]"),
        ],
        [1900, 3320, 1700, 1220, 1220],
        font_size=8.4,
    )


def add_competence_signoff(doc):
    add_heading(doc, "9. Competence Assessment and Restrictions", 1)
    add_table(
        doc,
        ["Assessment area", "Outcome", "Restrictions / supervision required", "Review date"],
        [
            ("General induction complete", "[Competent / needs support]", "[Restrictions]", "[Date]"),
            ("OHS baseline training complete", "[Competent / needs support]", "[Restrictions]", "[Date]"),
            ("Chemical handling authorised", "[Yes / No / limited]", "[Products or tasks restricted]", "[Date]"),
            ("Equipment authorised", "[Yes / No / limited]", "[Equipment restricted]", "[Date]"),
            ("Residential cleaning process competent", "[Competent / supervised / not applicable]", "[Restrictions]", "[Date]"),
            ("Commercial / client-site process competent", "[Competent / supervised / not applicable]", "[Restrictions]", "[Date]"),
            ("Driving / mobile route duties", "[Competent / supervised / not applicable]", "[Restrictions]", "[Date]"),
        ],
        [2500, 1850, 3560, 1450],
        font_size=8.7,
    )
    add_heading(doc, "9.1 Employee and Supervisor Declaration", 2)
    add_callout(
        doc,
        "Declaration wording",
        "The employee confirms that the training marked complete was explained in a language and manner they understand, that they had an opportunity to ask questions, and that they will follow lawful instructions, site rules, safe work procedures, client confidentiality requirements, and incident reporting rules.",
    )
    add_table(
        doc,
        ["Sign-off", "Name", "Signature", "Date"],
        [
            ("Employee", "[Name]", "[Signature]", "[Date]"),
            ("Trainer / assessor", "[Name]", "[Signature]", "[Date]"),
            ("Supervisor / manager", "[Name]", "[Signature]", "[Date]"),
        ],
        [1800, 2800, 3000, 1760],
        font_size=9.0,
    )


def add_training_plan(doc):
    add_page_break(doc)
    add_heading(doc, "10. Training Plan and Refresher Register", 1)
    add_table(
        doc,
        ["Training item", "Trigger / frequency", "Due date", "Completed", "Evidence location"],
        [
            ("Initial HR induction", "Before first shift", "[Date]", "[Yes / no]", "[Record]"),
            ("OHS induction", "Before site deployment and when risks change", "[Date]", "[Yes / no]", "[Record]"),
            ("Chemical / SDS refresher", "When product changes, SDS changes, incident occurs, or annually", "[Date]", "[Yes / no]", "[Record]"),
            ("PPE refresher", "When PPE changes, incorrect use observed, or annually", "[Date]", "[Yes / no]", "[Record]"),
            ("Equipment refresher", "Before new equipment use and after unsafe use/incident", "[Date]", "[Yes / no]", "[Record]"),
            ("Client privacy / confidentiality", "On appointment and at least annually", "[Date]", "[Yes / no]", "[Record]"),
            ("Manual handling and slips/trips", "On appointment and through toolbox talks", "[Date]", "[Yes / no]", "[Record]"),
            ("Site-specific induction", "Before first shift at each new client/site", "[Date]", "[Yes / no]", "[Record]"),
            ("Incident learnings / corrective action", "After incident, near miss, complaint, audit finding, or procedure change", "[Date]", "[Yes / no]", "[Record]"),
        ],
        [2300, 3230, 1200, 1100, 1530],
        font_size=8.5,
    )


def add_appendices(doc):
    add_heading(doc, "Appendix A: Training Session Attendance Record", 1)
    add_label_table(
        doc,
        [
            ("Session title", "[Topic]"),
            ("Trainer", "[Name / role]"),
            ("Date and duration", "[Date / start / finish]"),
            ("Location", "[Depot / client site / online / other]"),
            ("Materials used", "[SOP / SDS / toolbox talk / demo / video / assessment]"),
        ],
    )
    add_table(
        doc,
        ["Attendee name", "Employee no.", "Signature", "Competence outcome", "Follow-up action"],
        [
            ("[Name]", "[No.]", "[Signature]", "[C / NS / N/A]", "[Action]"),
            ("[Name]", "[No.]", "[Signature]", "[C / NS / N/A]", "[Action]"),
            ("[Name]", "[No.]", "[Signature]", "[C / NS / N/A]", "[Action]"),
            ("[Name]", "[No.]", "[Signature]", "[C / NS / N/A]", "[Action]"),
        ],
        [2350, 1200, 2100, 1700, 2010],
        font_size=8.7,
    )
    add_page_break(doc)
    add_heading(doc, "Appendix B: On-the-Job Coaching Log", 1)
    add_table(
        doc,
        ["Date", "Task observed", "What was coached", "Outcome", "Supervisor"],
        [
            ("[Date]", "[Task/site]", "[Observation and coaching given]", "[Improved / repeat / restrict]", "[Name]"),
            ("[Date]", "[Task/site]", "[Observation and coaching given]", "[Improved / repeat / restrict]", "[Name]"),
            ("[Date]", "[Task/site]", "[Observation and coaching given]", "[Improved / repeat / restrict]", "[Name]"),
            ("[Date]", "[Task/site]", "[Observation and coaching given]", "[Improved / repeat / restrict]", "[Name]"),
            ("[Date]", "[Task/site]", "[Observation and coaching given]", "[Improved / repeat / restrict]", "[Name]"),
        ],
        [1250, 1900, 3550, 1600, 1060],
        font_size=8.7,
    )
    add_heading(doc, "Appendix C: PPE Issue and Return Register", 1)
    add_table(
        doc,
        ["PPE / uniform item", "Size / serial", "Issue date", "Return / replace", "Employee sign"],
        [
            ("Gloves", "[Size/type]", "[Date]", "[Date/reason]", "[Initials]"),
            ("Safety shoes", "[Size/type]", "[Date]", "[Date/reason]", "[Initials]"),
            ("Eye protection", "[Type]", "[Date]", "[Date/reason]", "[Initials]"),
            ("Apron / overall / uniform", "[Size/type]", "[Date]", "[Date/reason]", "[Initials]"),
            ("Mask / respirator", "[Type/filter if applicable]", "[Date]", "[Date/reason]", "[Initials]"),
            ("Other", "[Details]", "[Date]", "[Date/reason]", "[Initials]"),
        ],
        [2300, 2000, 1300, 2060, 1700],
        font_size=8.5,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix D: Site-Specific Induction Record", 1)
    add_table(
        doc,
        ["Site / client", "Site-specific risks and rules", "Access / security controls", "Employee sign", "Supervisor"],
        [
            ("[Client/site]", "[Pets, stairs, parking, public areas, restricted rooms, special surfaces]", "[Keys, codes, alarm, sign-in]", "[Initial/date]", "[Name]"),
            ("[Client/site]", "[Risks and rules]", "[Controls]", "[Initial/date]", "[Name]"),
            ("[Client/site]", "[Risks and rules]", "[Controls]", "[Initial/date]", "[Name]"),
            ("[Client/site]", "[Risks and rules]", "[Controls]", "[Initial/date]", "[Name]"),
        ],
        [1600, 3350, 1720, 1290, 1400],
        font_size=8.4,
    )
    add_heading(doc, "Appendix E: Incident / Near-Miss Briefing Record", 1)
    add_table(
        doc,
        ["Event / lesson", "What changed", "Staff briefed", "Date", "Owner"],
        [
            ("[Incident / near miss / complaint]", "[New control, training, PPE, route, product, or procedure]", "[Names/teams]", "[Date]", "[Owner]"),
            ("[Event]", "[Change]", "[Names/teams]", "[Date]", "[Owner]"),
            ("[Event]", "[Change]", "[Names/teams]", "[Date]", "[Owner]"),
        ],
        [1900, 3600, 1900, 960, 1000],
        font_size=8.5,
    )

    add_heading(doc, "Appendix F: Source and Review List", 1)
    add_table(
        doc,
        ["Source", "URL / review note"],
        [
            ("Occupational Health and Safety Act 85 of 1993", "https://www.gov.za/documents/occupational-health-and-safety-act"),
            ("Regulations for Hazardous Chemical Agents, 2021", "https://www.labour.gov.za/DocumentCenter/Publications/Occupational%20Health%20and%20Safety/Regulations%20for%20Hazardous%20Chemical%20Agents%202021.pdf"),
            ("Basic Conditions of Employment Act 75 of 1997", "https://www.gov.za/documents/basic-conditions-employment-act"),
            ("National Minimum Wage Act and latest adjustments", "https://www.gov.za/documents/acts/national-minimum-wage-act-9-2018-english-tshivenda-27-nov-2018"),
            ("Skills Development Act 97 of 1998", "https://www.gov.za/documents/skills-development-act"),
            ("Compensation for Occupational Injuries and Diseases Act 130 of 1993", "https://www.gov.za/documents/compensation-occupational-injuries-and-diseases-act"),
            ("Protection of Personal Information Act 4 of 2013", "https://www.gov.za/documents/protection-personal-information-act"),
        ],
        [3300, 6060],
        font_size=8.0,
    )


def build_doc():
    doc = Document()
    configure_styles(doc)
    set_headers_and_footers(doc)

    add_cover(doc)
    add_document_control(doc)
    add_compliance_reference(doc)
    add_employee_record(doc)
    add_day_one_induction(doc)
    add_ohs_training(doc)
    add_cleaning_skills_matrix(doc)
    add_chemical_equipment(doc)
    add_competence_signoff(doc)
    add_training_plan(doc)
    add_appendices(doc)

    doc.core_properties.title = "Staff Onboarding and Training Record"
    doc.core_properties.subject = "South African cleaning services staff onboarding and training template"
    doc.core_properties.keywords = "cleaning services, South Africa, onboarding, training, OHS, POPIA, HCA, HR"
    doc.core_properties.comments = "Template generated for business use; verify current legal and client requirements before implementation."
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUT.resolve()}")
