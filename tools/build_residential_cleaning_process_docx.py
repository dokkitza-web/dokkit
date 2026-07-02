from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Residential_Cleaning_Process_Template.docx")

FONT = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
BORDER = "C9D3DF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn("w:ascii"), name)
    r_pr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, before=0, after=6, line=1.25, keep_next=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next is not None:
        fmt.keep_with_next = keep_next


def style_para(p, size=11, color=BODY, bold=False, italic=False):
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_para(
    doc,
    text="",
    *,
    style=None,
    size=11,
    color=BODY,
    bold=False,
    italic=False,
    align=None,
    before=0,
    after=6,
    line=1.25,
    keep_next=None,
):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=before, after=after, line=line, keep_next=keep_next)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    style_para(
        p,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **margins):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ["top", "start", "bottom", "end"]:
        if m not in margins:
            continue
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[m]))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, size=9.5, color=BODY, align=None, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    if text is not None:
        for idx, part in enumerate(str(text).split("\n")):
            if idx:
                p.add_run().add_break()
            run = p.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold)
    set_paragraph_format(p, before=0, after=0, line=1.15)
    if align is not None:
        p.alignment = align


def set_table_borders(table, color=BORDER, sz="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell, **CELL_MARGIN)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            size=font_size,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=header_fill,
        )
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for idx, text in enumerate(row_data):
            short = len(str(text)) < 16
            align = WD_ALIGN_PARAGRAPH.CENTER if short and idx in (0, len(row_data) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], text, size=font_size, align=align)
    add_para(doc, "", after=3)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D7DBE2", sz="6")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.text = ""
    title_run = p.add_run(title.upper() + "\n")
    set_run_font(title_run, size=10, color=DARK_BLUE, bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=10.5, color=BODY)
    set_paragraph_format(p, before=0, after=0, line=1.2)
    add_para(doc, "", after=2)
    return table


def add_step_table(doc, steps, *, font_size=8.8):
    rows = [(str(idx), step) for idx, step in enumerate(steps, start=1)]
    return add_table(
        doc,
        ["Step", "Action"],
        rows,
        [850, 8510],
        header_fill=LIGHT_BLUE,
        font_size=font_size,
    )


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def build_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True


def setup_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = ""
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    left, right = table.rows[0].cells
    set_cell_text(left, "SOP | Residential Cleaning Process", size=8.5, color=MUTED)
    set_cell_text(right, "[Business Name]", size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    footer = section.footer
    footer.paragraphs[0].text = ""
    ftable = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ftable, [6500, 2860], indent_dxa=0)
    lcell, rcell = ftable.rows[0].cells
    set_cell_text(lcell, "Controlled document - verify latest version before use", size=8.5, color=MUTED)
    p = rcell.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" of ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "NUMPAGES")
    set_paragraph_format(p, before=0, after=0, line=1.0)
    for cell in ftable.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)


def build_document():
    doc = Document()
    build_styles(doc)
    setup_header_footer(doc)

    props = doc.core_properties
    props.title = "Residential Cleaning Process Template"
    props.subject = "Operational SOP template for a South African residential cleaning services business"
    props.author = "OpenAI Codex"
    props.keywords = "SOP, residential cleaning, cleaning services, South Africa"

    add_para(doc, "", after=78)
    add_para(
        doc,
        "Standard Operating Procedure Template",
        size=10.5,
        bold=True,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=16,
    )
    add_para(
        doc,
        "Residential Cleaning Process",
        size=30,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.08,
    )
    add_para(
        doc,
        "For a South African Cleaning Services Business",
        size=15,
        color=DARK_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=36,
    )
    add_para(
        doc,
        "Prepared for: [Business Name] | Template date: June 2026",
        size=11.5,
        bold=True,
        color=BODY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=6,
    )
    add_para(
        doc,
        "Use this controlled SOP to standardise residential bookings, arrival, room-by-room cleaning, quality control, client handover, and records.",
        size=10.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=36,
    )
    add_table(
        doc,
        ["Field", "Template Entry"],
        [
            ("Document ID", "[SOP-RES-###]"),
            ("Version", "[0.1 / Draft / Approved]"),
            ("Effective date", "[dd/mm/yyyy]"),
            ("Review date", "[dd/mm/yyyy]"),
            ("Document owner", "[Operations Manager / Residential Services Lead]"),
            ("Approved by", "[Owner / Director / Accountable Manager]"),
            ("Applies to", "Residential cleaners, supervisors, drivers, subcontractors, and temporary staff"),
            ("Service area", "[Province / Municipality / Suburbs]"),
        ],
        [2500, 6860],
        font_size=9.5,
    )
    add_callout(
        doc,
        "Template control",
        "Replace bracketed fields before issue. Keep client-specific instructions in the job card, not in uncontrolled handwritten notes.",
        fill=CAUTION_FILL,
    )
    add_page_break(doc)

    add_heading(doc, "1. Document Control", 1)
    add_table(
        doc,
        ["Version", "Date", "Changed by", "Summary of change", "Approved by"],
        [
            ("0.1", "[dd/mm/yyyy]", "[Name]", "Initial template issue", "[Name]"),
            ("", "", "", "", ""),
            ("", "", "", "", ""),
        ],
        [1050, 1400, 1700, 3610, 1600],
        header_fill=LIGHT_GREY,
        font_size=9,
    )
    add_table(
        doc,
        ["Role / Area", "Controlled copy holder", "Issue method"],
        [
            ("Operations", "[Name / Role]", "[Shared drive / printed controlled copy]"),
            ("Residential supervisors", "[Name / Role]", "[Crew pack / mobile app]"),
            ("Training / HR", "[Name / Role]", "[Onboarding file / skills matrix]"),
            ("Client service desk", "[Name / Role]", "[Booking and complaint process]"),
        ],
        [2200, 3000, 4160],
        header_fill=LIGHT_GREY,
        font_size=9.2,
    )

    add_heading(doc, "2. Purpose, Scope, and Service Types", 1)
    add_para(
        doc,
        "This SOP defines the minimum process for residential cleaning jobs, from booking confirmation to final client handover. It must be adapted to the company's services, insurance conditions, client contracts, and local site risks.",
    )
    add_table(
        doc,
        ["Purpose area", "Business requirement"],
        [
            ("Consistency", "Every client receives a predictable cleaning process regardless of crew member or suburb."),
            ("Safety", "Chemical, electrical, slip, manual handling, and client-property risks are controlled before cleaning starts."),
            ("Quality", "Room-by-room standards are checked before the team leaves the property."),
            ("Trust", "Client privacy, keys, access codes, valuables, children, pets, and personal spaces are handled with care."),
            ("Records", "Job cards, snag notes, damage reports, photos, and sign-offs are captured in one approved record trail."),
        ],
        [2100, 7260],
        header_fill=LIGHT_GREY,
        font_size=9,
    )
    add_table(
        doc,
        ["Included residential services", "Excluded unless separately authorised"],
        [
            (
                "Standard domestic cleaning, once-off cleaning, move-in/move-out cleaning, post-renovation light cleaning, spring cleaning, laundry support, ironing support, appliance exterior cleaning, cupboard exterior cleaning, and scheduled recurring home cleans.",
                "Pest control, mould remediation, biohazard cleanup, hoarding cleanup, high-access ladder work above company limit, heavy construction debris removal, hazardous waste handling, child/pet care, gardening, pool services, and work requiring specialist training or a permit.",
            )
        ],
        [4680, 4680],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )

    add_heading(doc, "3. South African Compliance Reference Points", 1)
    add_para(
        doc,
        "This section is a practical reminder for SOP users. It does not replace legal advice, site-specific risk assessment, or client contract requirements.",
        italic=True,
        color=MUTED,
        size=10.5,
    )
    add_table(
        doc,
        ["Reference point", "Operational meaning for residential cleaning"],
        [
            (
                "Occupational Health and Safety Act 85 of 1993, including chemical safety duties where applicable",
                "Keep PPE, training, safe work methods, labels, and SDS access current. Do not clean where access, chemicals, electricity, slips, or client behaviour create uncontrolled risk.",
            ),
            (
                "Compensation for Occupational Injuries and Diseases Act / Compensation Fund obligations",
                "Escalate injuries, occupational exposures, and near misses through the company incident process.",
            ),
            (
                "Protection of Personal Information Act 4 of 2013",
                "Protect client addresses, access codes, alarm details, photos, keys, contact details, complaint records, and staff/client personal information.",
            ),
            (
                "National Environmental Management: Waste Act 59 of 2008 and municipal by-laws",
                "Dispose of general waste, recyclables, sharps, broken glass, chemical residues, and unusual household waste only through approved routes.",
            ),
            (
                "Client service agreement and insurance conditions",
                "Confirm scope, exclusions, valuables, damages, access, pets, keys, alarms, and photo permissions before the team starts work.",
            ),
        ],
        [2800, 6560],
        header_fill=LIGHT_GREY,
        font_size=8.5,
    )
    add_para(
        doc,
        "Useful official sources for company verification: labour.gov.za, gov.za, justice.gov.za, and the relevant municipality.",
        size=9.5,
        color=MUTED,
        italic=True,
        after=3,
    )

    add_heading(doc, "4. Roles and Responsibilities", 1)
    add_table(
        doc,
        ["Role", "Before cleaning", "During / after cleaning"],
        [
            ("Owner / Operations Manager", "Approves SOP, pricing rules, service menus, insurance limits, and escalation routes.", "Reviews complaints, losses, incidents, trends, and corrective actions."),
            ("Residential Supervisor", "Confirms job scope, risk notes, crew, route, keys/access, client expectations, and equipment.", "Leads arrival check, quality inspection, client handover, records, and rework decisions."),
            ("Cleaning Technician", "Arrives clean, on time, trained, fit for task, and equipped with PPE and approved tools.", "Follows room sequence, product instructions, privacy rules, and reports damage or concerns immediately."),
            ("Client Service / Admin", "Confirms booking, quote, client instructions, access, pets, parking, payment terms, and exclusions.", "Files sign-off, photos if authorised, complaints, compliments, rework notes, and next-booking actions."),
            ("Stores / Laundry", "Issues clean cloths, mop heads, consumables, chemicals, labels, and equipment.", "Receives dirty textiles, updates stock, tags damaged equipment, and quarantines unknown products."),
        ],
        [2100, 3630, 3630],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )

    add_heading(doc, "5. Residential Service Readiness", 1)
    add_table(
        doc,
        ["Readiness item", "Minimum requirement", "Initial"],
        [
            ("Booking confirmed", "Client name, address, phone, service type, arrival window, estimated duration, and payment instruction confirmed.", "[ ]"),
            ("Access confirmed", "Keys, gate code, guardhouse, alarm, parking, building rules, and contact person confirmed.", "[ ]"),
            ("Scope confirmed", "Rooms, priority areas, exclusions, add-ons, pets, children, valuables, and fragile items recorded.", "[ ]"),
            ("Crew ready", "Assigned team is trained, uniformed where required, and briefed on conduct and privacy rules.", "[ ]"),
            ("Chemicals ready", "Approved chemicals, labels, dilution instructions, SDS access, and PPE are packed.", "[ ]"),
            ("Equipment ready", "Vacuum, mop system, cloth colour-coding, buckets, brushes, extension leads, and consumables checked.", "[ ]"),
            ("Client risk notes", "Previous complaints, allergies, surface warnings, security rules, or pets are visible on the job card.", "[ ]"),
            ("Stop-work triggers", "Supervisor knows when to pause and escalate unsafe, out-of-scope, or disputed work.", "[ ]"),
        ],
        [2500, 5660, 1200],
        header_fill=LIGHT_GREY,
        font_size=8.7,
    )

    add_heading(doc, "6. Standard Residential Cleaning Workflow", 1)
    workflow = [
        "Confirm arrival with the client or authorised contact and introduce the team professionally.",
        "Complete the entry walk-through before unpacking chemicals or moving client property.",
        "Confirm rooms, priority areas, exclusions, fragile items, valuables, pets, children, alarms, keys, water, electricity, and waste rules.",
        "Set a safe staging area for chemicals, equipment, clean cloths, dirty cloths, waste bags, and personal bags.",
        "Place signs or verbal warnings where floors may become wet or where clients remain in the home.",
        "Brief the crew on task allocation, room sequence, PPE, products, surface warnings, privacy expectations, and finish time.",
        "Clean from cleaner areas to dirtier areas and from high surfaces to low surfaces unless the room method states otherwise.",
        "Use colour-coded cloths or separated cloth sets to prevent toilet, kitchen, bathroom, and general-area cross-contamination.",
        "Use only approved products, correct dilution, and manufacturer/client instructions for surfaces, appliances, floors, and fixtures.",
        "Keep client property in place unless movement is necessary and safe. Return moved items to their original position where practical.",
        "Report existing damage, breakage, stains, pests, mould, leaks, biohazards, sharps, or unsafe conditions immediately.",
        "Perform room-by-room quality checks before asking for client sign-off.",
        "Dispose of waste as agreed, remove company materials, clean tools for transport, and ensure the home is safe before departure.",
        "Offer final handover, record snags or rework, obtain sign-off where available, and submit job records.",
    ]
    add_step_table(doc, workflow, font_size=8.5)
    add_callout(
        doc,
        "Residential trust rule",
        "Staff must not open private cupboards, drawers, fridges, files, safes, bedside storage, phones, computers, or personal bags unless the job card specifically authorises the task.",
        fill=CAUTION_FILL,
    )

    add_heading(doc, "7. Room-by-Room Cleaning Standards", 1)
    add_table(
        doc,
        ["Area", "Standard process", "Quality check"],
        [
            ("Entrance / hallways", "Dust reachable surfaces, wipe touch points, remove cobwebs where reachable, vacuum/sweep, mop if suitable.", "No visible loose dirt, sticky marks, or wet-slip risk."),
            ("Living areas", "Dust surfaces, wipe accessible furniture exteriors, vacuum upholstery where included, clean mirrors/glass where included, clean floors.", "Items replaced, no streaky glass, floors clean and safe."),
            ("Bedrooms", "Dust surfaces, make beds where included, clean mirrors, empty bins, vacuum/sweep/mop floors, handle personal items minimally.", "Privacy respected, personal items not rearranged unnecessarily."),
            ("Bathrooms", "Apply product safely, clean toilet, basin, taps, shower/bath, tiles where included, mirrors, counters, bins, and floors.", "Sanitary surfaces clean, product residue rinsed, floor dry/safe."),
            ("Kitchen", "Clean counters, sink, taps, appliance exteriors, cupboard exteriors where included, table, splashback, bins, and floors.", "No food contamination, no chemical residue on food-prep surfaces."),
            ("Laundry / utility", "Wipe machines externally, clean sinks, sort laundry only where included, sweep/mop floors.", "Laundry instructions followed, no client clothing mixed incorrectly."),
            ("Stairs", "Vacuum/sweep carefully, wipe handrails, spot-clean marks where included.", "No trip hazards, wet steps are never left unattended."),
            ("Patio / balcony", "Sweep, wipe furniture where included, remove loose debris according to scope.", "Doors locked/secured and outdoor waste handled as agreed."),
        ],
        [1600, 5200, 2560],
        header_fill=LIGHT_BLUE,
        font_size=8.2,
    )

    add_heading(doc, "8. Surface and Product Control", 1)
    add_table(
        doc,
        ["Surface / item", "Approved approach", "Do not do without approval"],
        [
            ("Natural stone", "Use pH-neutral approved cleaner and dry promptly.", "Do not use acidic, abrasive, or unknown products."),
            ("Wood / laminate floors", "Use minimal moisture and approved product. Dry excess water immediately.", "Do not flood mop or use harsh degreasers."),
            ("Stainless steel", "Wipe with grain where possible and use approved cloth/product.", "Do not use abrasive pads that scratch."),
            ("Glass / mirrors", "Use clean glass cloth and streak-control method.", "Do not spray near electronics, plugs, or untreated wood."),
            ("Upholstery / carpets", "Vacuum only unless spot cleaning or extraction is part of the job card.", "Do not apply stain remover without test/approval."),
            ("Appliance interiors", "Clean only if included and safe to access.", "Do not disconnect appliances or handle spoiled food without instruction."),
            ("Electronics", "Dust exterior with dry cloth only if included.", "Do not spray, unplug, move, or clean screens with chemical unless approved."),
            ("Children's items / toys", "Clean only according to client instruction with safe product and clean cloth.", "Do not apply strong chemicals or mix with bathroom cloths."),
        ],
        [1900, 4260, 3200],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )
    add_table(
        doc,
        ["Chemical control", "Required practice"],
        [
            ("Labels", "Every bottle must show product name, dilution, hazard warning if applicable, and company identifier."),
            ("SDS access", "Crew must know how to access safety data sheets or product safety instructions."),
            ("Dilution", "Use the approved dilution chart. Do not guess or mix stronger because a surface looks dirty."),
            ("Ventilation", "Ventilate where products require it and where odours may affect clients, children, pets, or staff."),
            ("Storage", "Keep chemicals upright, closed, out of client reach, and away from food, heat, and incompatible products."),
            ("Client products", "Use client-supplied products only if authorised and safe; record it on the job card."),
        ],
        [2200, 7160],
        header_fill=LIGHT_GREY,
        font_size=8.8,
    )

    add_heading(doc, "9. Client Property, Privacy, and Conduct", 1)
    add_table(
        doc,
        ["Control area", "Required conduct"],
        [
            ("Access", "Keys, remotes, tags, alarm codes, and gate codes are controlled records. Do not share them outside the authorised team."),
            ("Photos", "Take before/after photos only where authorised and only for work evidence, damage, or quality records."),
            ("Valuables", "Do not handle cash, jewellery, medication, documents, firearms, or private items unless a supervisor directs a safety step."),
            ("Children and vulnerable persons", "Staff do not supervise, transport, feed, or discipline children or vulnerable persons."),
            ("Pets", "Do not feed, walk, medicate, or release pets unless the job card specifically authorises it."),
            ("Food and drink", "Do not consume client food or drink. Staff breaks must follow company and client rules."),
            ("Communication", "Keep tone professional. Do not discuss client personal matters, property value, or household details outside the job."),
            ("Social media", "Do not post client homes, possessions, addresses, staff, or job details on personal platforms."),
        ],
        [2200, 7160],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )

    add_heading(doc, "10. Quality Control and Client Handover", 1)
    add_step_table(
        doc,
        [
            "Supervisor or senior cleaner checks each completed room against the job card and room-by-room standards.",
            "Check high-touch points, floors, bins, mirrors, taps, toilets, counters, and priority client areas.",
            "Confirm all company products, cloths, waste bags, tools, and personal items are removed from the home.",
            "Confirm moved client items are returned where practical and that doors, windows, gates, alarms, and lights are handled as instructed.",
            "Record any exclusions, locked rooms, pre-existing damage, breakages, unsafe conditions, pest signs, mould, or client requests.",
            "Offer a walk-through to the client or authorised representative where available.",
            "Correct reasonable snags before leaving if time, safety, and scope allow.",
            "Obtain client sign-off or record that the client was unavailable.",
            "Submit job card, photos if authorised, snag notes, incident reports, and next-service recommendations.",
        ],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Quality hold point",
        "The team may not depart until the supervisor confirms the property is safe, no company item is left behind, and all known snags or exclusions are recorded.",
        fill=CAUTION_FILL,
    )

    add_heading(doc, "11. Incidents, Damage, Complaints, and Stop-Work Triggers", 1)
    add_table(
        doc,
        ["Trigger", "Immediate response", "Record / escalation"],
        [
            ("Injury or exposure", "Stop work, make area safe, provide first aid, and call emergency support if needed.", "Incident report; notify Operations/SHEQ; follow company injury process."),
            ("Breakage or damage", "Stop cleaning affected item/area, preserve evidence, notify supervisor and client contact.", "Damage report with photos if authorised, witness details, and corrective action."),
            ("Chemical spill", "Keep people/pets away, ventilate where safe, use spill kit/SDS controls.", "Spill report; quarantine product if needed."),
            ("Unsafe home condition", "Pause task and isolate hazard if safe to do so.", "Risk note, client notification, and supervisor decision."),
            ("Aggressive client or security concern", "Leave the area if needed and contact supervisor immediately.", "Security/behaviour incident report."),
            ("Biohazard, sharps, pests, mould, or bodily fluids", "Do not handle unless included in approved scope and trained.", "Stop-work note and revised quotation/risk assessment if applicable."),
            ("Out-of-scope request", "Do not start extra work until supervisor and client approve scope and price.", "Variation record or rebooking note."),
            ("Complaint", "Listen, avoid argument, record facts, and notify supervisor.", "Complaint record, response owner, due date, and resolution note."),
        ],
        [1900, 4000, 3460],
        header_fill=LIGHT_GREY,
        font_size=8.4,
    )
    add_table(
        doc,
        ["Emergency / escalation contact", "Name", "Number", "After-hours instruction"],
        [
            ("Operations Manager", "[Name]", "[Number]", "[Instruction]"),
            ("Residential Supervisor", "[Name]", "[Number]", "[Instruction]"),
            ("Client service desk", "[Name]", "[Number]", "[Instruction]"),
            ("SHEQ / First Aid Lead", "[Name]", "[Number]", "[Instruction]"),
            ("Emergency services", "[Local emergency number]", "[Number]", "[Instruction]"),
        ],
        [2600, 2100, 1900, 2760],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )

    add_heading(doc, "12. Records and Retention", 1)
    add_table(
        doc,
        ["Record", "Owner", "Where filed", "Minimum retention"],
        [
            ("Approved SOP and revision history", "Operations", "[Controlled document folder]", "[Company schedule]"),
            ("Booking and quote", "Client Service", "[Client/job file]", "[Company schedule]"),
            ("Job card and client sign-off", "Supervisor", "[Client/job file]", "[Company schedule]"),
            ("Before/after photos if authorised", "Supervisor / Admin", "[Approved media folder]", "[Company/client schedule]"),
            ("Incident, damage, and complaint reports", "Operations / SHEQ", "[Incident/complaint register]", "[Company/legal schedule]"),
            ("Chemical issue and SDS records", "Stores / SHEQ", "[Chemical control file]", "[Company schedule]"),
            ("Training and competency records", "HR / Training", "[Training file]", "[Company/legal schedule]"),
        ],
        [2500, 1800, 3060, 2000],
        header_fill=LIGHT_GREY,
        font_size=8.5,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix A: Residential Job Card", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Client name", "[Name]"),
            ("Address", "[Full address and unit/access details]"),
            ("Contact details", "[Phone / email]"),
            ("Service type", "[Standard / deep clean / move-in / move-out / recurring / other]"),
            ("Date and time", "[Date / arrival window / estimated finish]"),
            ("Team", "[Supervisor and crew names]"),
            ("Rooms included", "[Bedrooms, bathrooms, kitchen, living areas, extras]"),
            ("Priority areas", "[Client priorities]"),
            ("Exclusions", "[Locked rooms, fragile areas, out-of-scope work]"),
            ("Access and alarm", "[Keys, gate, alarm, guardhouse, parking]"),
            ("Pets / children / occupants", "[Instructions]"),
            ("Surface warnings", "[Stone, wood, antiques, electronics, fragile fixtures]"),
            ("Photo permission", "[Yes / No / limited to damage evidence]"),
        ],
        [2500, 6860],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )

    add_heading(doc, "Appendix B: Arrival and Departure Checklist", 1)
    add_table(
        doc,
        ["Phase", "Checklist item", "Done", "Initial"],
        [
            ("Arrival", "Client/contact checked in and access confirmed.", "[ ]", ""),
            ("Arrival", "Rooms, priorities, exclusions, valuables, pets, and surface warnings confirmed.", "[ ]", ""),
            ("Arrival", "Safe staging area set up and chemicals controlled.", "[ ]", ""),
            ("Arrival", "Crew briefing completed.", "[ ]", ""),
            ("During", "Room sequence followed and cross-contamination controls used.", "[ ]", ""),
            ("During", "Damage, stains, pests, mould, sharps, or unsafe conditions reported.", "[ ]", ""),
            ("Departure", "Room-by-room quality check completed.", "[ ]", ""),
            ("Departure", "Company equipment, products, cloths, and waste removed or handled as agreed.", "[ ]", ""),
            ("Departure", "Doors, windows, gates, alarms, keys, and lights handled as instructed.", "[ ]", ""),
            ("Departure", "Client sign-off offered or reason unavailable recorded.", "[ ]", ""),
            ("Departure", "Job records submitted.", "[ ]", ""),
        ],
        [1300, 6040, 900, 1120],
        header_fill=LIGHT_BLUE,
        font_size=8.4,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix C: Room Quality Checklist", 1)
    add_table(
        doc,
        ["Room / area", "Standard checks", "Done", "Snags / notes"],
        [
            ("Entrance / hallway", "Floors, touch points, cobwebs, visible marks, no trip/slip risk.", "[ ]", ""),
            ("Living area", "Dusting, floors, glass/mirrors if included, items replaced.", "[ ]", ""),
            ("Bedroom 1", "Dusting, bed if included, bins, floors, privacy respected.", "[ ]", ""),
            ("Bedroom 2", "Dusting, bed if included, bins, floors, privacy respected.", "[ ]", ""),
            ("Bathroom 1", "Toilet, basin, taps, shower/bath, mirror, bins, floor, product residue removed.", "[ ]", ""),
            ("Bathroom 2", "Toilet, basin, taps, shower/bath, mirror, bins, floor, product residue removed.", "[ ]", ""),
            ("Kitchen", "Counters, sink, taps, appliance exteriors, cupboard exteriors if included, bins, floor.", "[ ]", ""),
            ("Laundry / utility", "Machine exteriors, sink, floor, laundry instructions followed if included.", "[ ]", ""),
            ("Stairs", "Steps, handrails, no wet-step hazard.", "[ ]", ""),
            ("Patio / balcony", "Scope completed, doors secured, debris/waste handled.", "[ ]", ""),
        ],
        [1800, 5060, 900, 1600],
        header_fill=LIGHT_GREY,
        font_size=8.2,
    )

    add_heading(doc, "Appendix D: Chemical and Equipment Issue Register", 1)
    add_table(
        doc,
        ["Date", "Product / equipment", "Qty out", "Condition out", "Qty back", "Condition back", "Action"],
        [
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
        ],
        [1050, 2200, 850, 1400, 850, 1450, 1560],
        header_fill=LIGHT_GREY,
        font_size=8.1,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix E: Client Completion Sign-Off", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Client / site", "[Client name and address]"),
            ("Job number", "[Job number]"),
            ("Service date and time", "[Start] to [Finish]"),
            ("Areas completed", "[Areas / rooms completed]"),
            ("Exclusions / access limits", "[Areas not completed and reason]"),
            ("Snags or rework", "[Details, owner, due date]"),
            ("Waste handled", "[Client retained / contractor removed / municipal bin / other]"),
            ("Photos authorised", "[Yes / No / Not applicable]"),
            ("Client representative", "[Name, signature, date]"),
            ("Supervisor / cleaner", "[Name, signature, date]"),
        ],
        [2500, 6860],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )

    add_heading(doc, "Appendix F: Incident, Damage, or Complaint Report", 1)
    add_table(
        doc,
        ["Field", "Entry"],
        [
            ("Report type", "[Injury / exposure / breakage / damage / complaint / security / non-conformance / other]"),
            ("Date, time, address", "[Details]"),
            ("Reported by", "[Name and role]"),
            ("People involved", "[Names / roles / contact details where appropriate]"),
            ("Description", "[What happened, where, and what was affected]"),
            ("Immediate action taken", "[Make safe, first aid, client notice, isolation, spill control]"),
            ("Photos or evidence", "[Authorised? Where filed?]"),
            ("Client notified", "[Name, time, method]"),
            ("Corrective action", "[Action, owner, due date]"),
            ("Closed by", "[Name, role, date]"),
        ],
        [2500, 6860],
        header_fill=LIGHT_GREY,
        font_size=8.8,
    )
    add_para(
        doc,
        "End of template. Adapt this SOP with the business owner, supervisor, SHEQ representative, insurer, client service process, and competent compliance adviser.",
        italic=True,
        color=MUTED,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=8,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT.resolve())
