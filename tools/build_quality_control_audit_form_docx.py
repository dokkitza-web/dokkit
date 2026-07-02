from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Quality_Control_Audit_Form.docx")

FONT = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
CAUTION_GOLD = RGBColor(0x7A, 0x5A, 0x00)
RISK_RED = RGBColor(0x9B, 0x1C, 0x1C)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"
RISK_FILL = "FBEAEA"
BORDER = "C9D3DF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn("w:ascii"), name)
    r_pr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, before=0, after=6, line=1.25, keep_next=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next is not None:
        fmt.keep_with_next = keep_next


def style_para(p, size=11, color=BODY, bold=False, italic=False):
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_para(
    doc,
    text="",
    *,
    style=None,
    size=11,
    color=BODY,
    bold=False,
    italic=False,
    align=None,
    before=0,
    after=6,
    line=1.25,
    keep_next=None,
):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=before, after=after, line=line, keep_next=keep_next)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    style_para(
        p,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **margins):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ["top", "start", "bottom", "end"]:
        if margin not in margins:
            continue
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[margin]))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, size=8.7, color=BODY, align=None, fill=None, italic=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    if text is not None:
        for idx, part in enumerate(str(text).split("\n")):
            if idx:
                p.add_run().add_break()
            run = p.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_paragraph_format(p, before=0, after=0, line=1.13)
    if align is not None:
        p.alignment = align


def set_table_borders(table, color=BORDER, sz="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell, **CELL_MARGIN)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            size=font_size,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=header_fill,
        )
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            align = None
            if str(value).strip() in {"[ ]", "[0-3 / N/A]", "[Yes / No / N/A]", "[Score]", "[%]"}:
                align = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], value, size=font_size, align=align)
        prevent_row_split(table.rows[-1])
    add_para(doc, "", after=3)
    return table


def add_label_table(doc, rows, label_width=2700, font_size=9.2):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [label_width, CONTENT_WIDTH_DXA - label_width])
    set_table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=NAVY, fill=LIGHT_GREY, size=font_size)
        set_cell_text(cells[1], value, size=font_size)
        prevent_row_split(table.rows[-1])
    add_para(doc, "", after=3)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT_FILL, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D5DCE5", sz="4")
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(title)
    set_run_font(run, size=10.2, color=title_color, bold=True)
    p.add_run().add_break()
    run = p.add_run(body)
    set_run_font(run, size=9.4, color=BODY)
    set_paragraph_format(p, before=0, after=0, line=1.18)
    add_para(doc, "", after=5)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.2, color=BODY)
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_headers_and_footers(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    left = hp.add_run("QC | Quality Control Audit Form")
    set_run_font(left, size=9, color=MUTED)
    right = hp.add_run("    [Business Name]")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    left = fp.add_run("Controlled document - verify latest version before use")
    set_run_font(left, size=9, color=MUTED)


def add_cover(doc):
    add_para(doc, "Operations Quality Template", size=10, color=CAUTION_GOLD, bold=True, after=0)
    add_para(
        doc,
        "Quality Control Audit Form",
        size=27,
        color=NAVY,
        bold=True,
        after=6,
        line=1.1,
    )
    add_para(
        doc,
        "For a South African cleaning services business",
        size=13,
        color=MUTED,
        after=18,
    )
    add_label_table(
        doc,
        [
            ("Business name", "[Business Name]"),
            ("Client / site", "[Client name / site address]"),
            ("Audit date and time", "[YYYY-MM-DD / start-finish]"),
            ("Service type", "[Residential / commercial / deep clean / move-in/out / mobile route]"),
            ("Audit type", "[Routine / re-audit / complaint follow-up / handover / supervisor check]"),
            ("Team / supervisor", "[Names]"),
            ("Auditor", "[Name and role]"),
            ("Overall result", "[Pass / conditional pass / fail / re-audit required]"),
        ],
    )
    add_callout(
        doc,
        "Purpose",
        "Use this form to check cleaning quality, safety controls, chemical handling, client property care, "
        "staff conduct, and corrective actions. Adapt the audit criteria to the service agreement, scope of work, "
        "site risk assessment, product SDSs, and client-specific requirements.",
    )


def add_document_control(doc):
    add_heading(doc, "1. Document Control", 1)
    add_table(
        doc,
        ["Control field", "Template entry", "Notes"],
        [
            ("Document number", "QC-AUD-[000]", "Allocate a unique audit number."),
            ("Version", "[1.0]", "Update when audit criteria, scope, or scoring rules change."),
            ("Effective date", "[YYYY-MM-DD]", "Use the date approved for operational use."),
            ("Owner", "[Operations / Quality Manager]", "Assign responsibility for template control."),
            ("Review cycle", "At least annually or when laws, client requirements, products, or processes change.", "Keep the form aligned with current operations."),
            ("Record location", "[Folder / CRM / shared drive / job file]", "Store completed forms securely."),
        ],
        [2100, 4100, 3160],
    )
    add_heading(doc, "2. How to Use This Audit Form", 1)
    add_bullet(doc, "Complete the audit details before starting the inspection, including the service scope and areas inspected.")
    add_bullet(doc, "Score only the items that apply to the service. Mark non-applicable items as N/A and exclude them from the percentage calculation.")
    add_bullet(doc, "Capture objective evidence: photos, job-card notes, client feedback, supervisor observations, and staff explanations where relevant.")
    add_bullet(doc, "Record every defect or non-conformance in the corrective action plan with an owner, due date, and follow-up status.")
    add_bullet(doc, "Do not store unnecessary personal information or client-private details in the audit record.")


def add_scoring(doc):
    add_heading(doc, "3. Scoring and Rating Guide", 1)
    add_table(
        doc,
        ["Score", "Meaning", "Action standard"],
        [
            ("3", "Pass: requirement fully met.", "No corrective action needed; note good practice where useful."),
            ("2", "Minor issue: service standard mostly met, but small correction needed.", "Correct during shift where possible and record trend if repeated."),
            ("1", "Major issue: standard not met and client/service/safety outcome affected.", "Immediate correction required; supervisor review and root cause check."),
            ("0", "Critical fail: unsafe, unhygienic, unauthorised, serious client-impacting, or repeat failure.", "Stop or restrict task if needed, escalate, correct, and re-audit."),
            ("N/A", "Not applicable to this service, site, or audit scope.", "Exclude from total score and note reason if unclear."),
        ],
        [1100, 4100, 4160],
        font_size=8.8,
    )
    add_table(
        doc,
        ["Rating band", "Suggested result", "Operational response"],
        [
            ("90-100%", "Pass", "Confirm close-out and share positive feedback with the team."),
            ("80-89%", "Conditional pass", "Correct minor issues and monitor for repeat findings."),
            ("65-79%", "Improvement required", "Corrective action plan, supervisor coaching, and follow-up audit."),
            ("Below 65% or any critical fail", "Fail / re-audit required", "Escalate to operations manager and re-audit before sign-off where appropriate."),
        ],
        [1800, 2200, 5360],
        font_size=8.8,
    )
    add_label_table(
        doc,
        [
            ("Score formula", "Total achieved points / total applicable points x 100 = audit percentage."),
            ("Critical fail override", "A critical safety, chemical, security, privacy, or client property issue may require a fail result even if the numeric score is high."),
        ],
        label_width=2300,
    )


def add_compliance_reference(doc):
    add_heading(doc, "4. South African Reference Points", 1)
    add_table(
        doc,
        ["Reference", "Audit relevance", "Example audit control"],
        [
            (
                "Occupational Health and Safety Act 85 of 1993",
                "Cleaning work must be managed to protect employees and others affected by work activities.",
                "Wet-floor controls, PPE use, safe equipment, incident reporting, and risk controls.",
            ),
            (
                "Hazardous Chemical Agents Regulations, 2021",
                "Cleaning products may expose staff or occupants to hazardous chemical agents.",
                "SDS access, correct dilution, labelling, storage, PPE, spill response, and no unauthorised mixing.",
            ),
            (
                "Consumer Protection Act 68 of 2008",
                "Service delivery should match agreed quality, scope, and client-facing representations.",
                "Scope adherence, defect handling, complaint capture, client sign-off, and transparent corrective actions.",
            ),
            (
                "Protection of Personal Information Act 4 of 2013",
                "Audit photos, client details, access codes, staff names, and complaint notes may contain personal information.",
                "Limit audit evidence to what is necessary, store securely, and avoid public sharing without approval.",
            ),
            (
                "National Environmental Management: Waste Act 59 of 2008",
                "Waste handling and disposal must not create health, environmental, or client-site risks.",
                "Segregation, bagging, disposal route, chemical container handling, and spill/waste escalation.",
            ),
        ],
        [2600, 4100, 2660],
        font_size=8.4,
    )
    add_callout(
        doc,
        "Important note",
        "This template is an operational quality tool, not legal advice. Check the latest legislation, client contract, scope of work, risk assessment, and product SDSs before implementation.",
        fill=CAUTION_FILL,
        title_color=CAUTION_GOLD,
    )


def add_audit_details(doc):
    add_page_break(doc)
    add_heading(doc, "5. Audit Details and Scope", 1)
    add_label_table(
        doc,
        [
            ("Audit number", "[QC-AUD-000]"),
            ("Client / site", "[Name / address / unit / area]"),
            ("Cleaning date", "[YYYY-MM-DD]"),
            ("Audit date", "[YYYY-MM-DD]"),
            ("Service scope audited", "[Routine / deep clean / post-construction / move-in/out / office / residential]"),
            ("Areas inspected", "[Kitchen / bathrooms / bedrooms / offices / floors / windows / exterior / other]"),
            ("Areas excluded", "[List and reason]"),
            ("Staff present", "[Names / roles]"),
            ("Audit evidence", "[Photos / checklist / client feedback / job card / supervisor observation]"),
            ("Client present?", "[Yes / no / representative]"),
        ],
        label_width=2500,
    )
    add_heading(doc, "5.1 Audit Summary", 2)
    add_table(
        doc,
        ["Section", "Applicable points", "Points achieved", "Percentage", "Notes / trend"],
        [
            ("Arrival, scope, and administration", "[Total]", "[Score]", "[%]", "[Notes]"),
            ("Site set-up and safety controls", "[Total]", "[Score]", "[%]", "[Notes]"),
            ("Chemicals, PPE, and equipment", "[Total]", "[Score]", "[%]", "[Notes]"),
            ("Cleaning quality by area", "[Total]", "[Score]", "[%]", "[Notes]"),
            ("Client property, privacy, and security", "[Total]", "[Score]", "[%]", "[Notes]"),
            ("Staff conduct and supervision", "[Total]", "[Score]", "[%]", "[Notes]"),
            ("Overall result", "[Total]", "[Score]", "[%]", "[Pass / conditional / fail]"),
        ],
        [2500, 1400, 1400, 1100, 2960],
        font_size=8.6,
    )


def add_audit_matrix_section(doc, heading, rows):
    add_heading(doc, heading, 1)
    add_table(
        doc,
        ["Audit item", "Check method / requirement", "Score", "Evidence / notes", "Action owner / due"],
        rows,
        [2100, 3450, 900, 1850, 1060],
        font_size=8.1,
    )


def add_audit_matrices(doc):
    add_audit_matrix_section(
        doc,
        "6. Arrival, Scope, and Administration",
        [
            ("Job card / scope", "Correct client, address, service date, access instructions, exclusions, and special requirements confirmed.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Team briefing", "Team understands scope, risks, client requests, quality focus areas, and escalation route.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Timekeeping", "Arrival, start, finish, and travel times recorded accurately.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Client communication", "Client or representative contacted according to agreed process; changes recorded.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Access control", "Keys, remotes, gate codes, alarm codes, lockboxes, and sign-in/out handled securely.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Pre-clean condition", "Existing damage, access issues, fragile items, or excluded areas noted before work begins.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
        ],
    )
    add_audit_matrix_section(
        doc,
        "7. Site Set-Up and Safety Controls",
        [
            ("Risk assessment", "Site hazards identified and controls applied: stairs, wet floors, pets, public areas, electricity, manual handling.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Wet-floor controls", "Warning signs, controlled access, dry routes, and prompt spill clean-up are in place.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Manual handling", "Loads, trolleys, team lifts, and furniture movement controlled; no unsafe overreach.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Electrical safety", "Vacuum/equipment cables, plugs, extension leads, and water/electricity separation checked.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Emergency readiness", "Emergency contacts, first aid route, incident reporting, and security/panic process understood.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Public / occupant safety", "Work area controlled so residents, visitors, pets, children, and client staff are not exposed to avoidable risk.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
        ],
    )
    add_audit_matrix_section(
        doc,
        "8. Chemicals, PPE, and Equipment",
        [
            ("Product authorisation", "Only approved products used for the surface, task, and client requirement.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Labels and SDS", "Product labels readable; SDS or product safety instructions available to staff.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Dilution and mixing", "Correct dilution used; no unauthorised mixing of products; containers not repurposed unsafely.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("PPE", "Correct gloves, eye protection, masks/respirators, footwear, aprons, or other PPE used and in good condition.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Storage and transport", "Products closed, upright, segregated where needed, and secured in vehicle or site storage.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Equipment condition", "Vacuum, mop systems, buckets, cloths, ladders/steps, machines, and tools clean, safe, and fit for task.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Colour coding", "Cloths, mops, brushes, and buckets used according to area to prevent cross-contamination.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
        ],
    )
    add_page_break(doc)
    add_audit_matrix_section(
        doc,
        "9. Cleaning Quality by Area",
        [
            ("General surfaces", "Dust, marks, smudges, cobwebs, visible debris, and missed surfaces checked.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Floors and carpets", "Sweeping/vacuuming/mopping complete; edges, corners, under movable items, stains, and wetness checked.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Bathrooms / toilets", "Toilets, basins, showers, taps, mirrors, drains, bins, odours, and hygiene touchpoints checked.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Kitchen / food areas", "Counters, sinks, splashbacks, appliances, handles, bins, grease, residue, and food-safe practices checked.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("High-touch points", "Door handles, switches, rails, remotes, appliance handles, desks, taps, and shared surfaces checked.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Windows / glass", "Glass, mirrors, frames, tracks, and safe reach limits checked; no streaking or unsafe access.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Waste and bins", "Waste removed/segregated according to client rules; bags not overloaded; liners replaced where required.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Final presentation", "Furniture returned where authorised, lights/doors/windows checked, supplies packed down, and site left neat.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
        ],
    )
    add_audit_matrix_section(
        doc,
        "10. Client Property, Privacy, and Security",
        [
            ("Client property care", "Fragile, personal, confidential, or high-value items handled respectfully and not moved without authorisation.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Damage reporting", "Breakage, stain, scratch, loss, water damage, or security concern reported immediately and recorded.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Photos / audit evidence", "Photos limited to necessary evidence; no personal/private information captured unnecessarily.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Access close-out", "Keys, cards, codes, alarms, windows, gates, and lock-up completed according to client instruction.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Confidentiality", "No client information, images, access codes, documents, or conversations shared outside authorised channels.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Client feedback", "Compliments, concerns, defects, complaints, or scope-change requests captured and routed.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
        ],
    )
    add_audit_matrix_section(
        doc,
        "11. Staff Conduct and Supervision",
        [
            ("Uniform and presentation", "Uniform/PPE clean and appropriate; identification visible where required.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Professional conduct", "Respectful communication, no unauthorised guests, no inappropriate phone use, no conflict at site.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Task method", "Staff follow SOP sequence, contact times, dwell times, surface care instructions, and safe work methods.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Supervisor oversight", "Supervisor checked quality, corrected defects during shift, and coached team where needed.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Training gap", "Any skill, product, equipment, conduct, or safety training gap identified and added to training plan.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
            ("Close-out briefing", "Findings discussed with team; corrective actions and praise shared before job closure where practical.", "[0-3 / N/A]", "[Evidence]", "[Owner / due]"),
        ],
    )


def add_corrective_action(doc):
    add_page_break(doc)
    add_heading(doc, "12. Defect and Non-Conformance Log", 1)
    add_table(
        doc,
        ["No.", "Defect / finding", "Severity", "Root cause", "Immediate correction", "Owner / due"],
        [
            ("1", "[Finding]", "[Minor / major / critical]", "[Cause]", "[Correction]", "[Owner / date]"),
            ("2", "[Finding]", "[Minor / major / critical]", "[Cause]", "[Correction]", "[Owner / date]"),
            ("3", "[Finding]", "[Minor / major / critical]", "[Cause]", "[Correction]", "[Owner / date]"),
            ("4", "[Finding]", "[Minor / major / critical]", "[Cause]", "[Correction]", "[Owner / date]"),
            ("5", "[Finding]", "[Minor / major / critical]", "[Cause]", "[Correction]", "[Owner / date]"),
        ],
        [650, 2500, 1450, 1700, 2000, 1060],
        font_size=8.2,
    )
    add_heading(doc, "13. Corrective Action Plan and Follow-Up", 1)
    add_table(
        doc,
        ["Action", "Owner", "Due date", "Verification method", "Status", "Closed by / date"],
        [
            ("[Action required]", "[Name]", "[Date]", "[Photo / re-audit / client confirmation / supervisor check]", "[Open / closed]", "[Name / date]"),
            ("[Action required]", "[Name]", "[Date]", "[Method]", "[Open / closed]", "[Name / date]"),
            ("[Action required]", "[Name]", "[Date]", "[Method]", "[Open / closed]", "[Name / date]"),
            ("[Action required]", "[Name]", "[Date]", "[Method]", "[Open / closed]", "[Name / date]"),
        ],
        [2500, 1200, 1100, 2600, 1050, 910],
        font_size=8.2,
    )
    add_heading(doc, "14. Audit Sign-Off", 1)
    add_callout(
        doc,
        "Sign-off statement",
        "The auditor confirms that the audit was completed against the applicable scope, findings were recorded honestly, and corrective actions were assigned where required. Signature does not remove the need to escalate serious safety, security, privacy, property damage, or client complaint issues.",
    )
    add_table(
        doc,
        ["Role", "Name", "Signature", "Date", "Comments"],
        [
            ("Auditor", "[Name]", "[Signature]", "[Date]", "[Comments]"),
            ("Team leader / supervisor", "[Name]", "[Signature]", "[Date]", "[Comments]"),
            ("Operations manager", "[Name]", "[Signature]", "[Date]", "[Comments]"),
            ("Client / representative", "[Optional]", "[Signature]", "[Date]", "[Comments]"),
        ],
        [1700, 1900, 1900, 1200, 2660],
        font_size=8.7,
    )


def add_appendices(doc):
    add_heading(doc, "Appendix A: Area-by-Area Room Audit Sheet", 1)
    add_table(
        doc,
        ["Area / room", "Items checked", "Result", "Defects / photos", "Corrected before leaving?"],
        [
            ("Kitchen", "[Counters, sink, appliances, handles, floor, bins, odour]", "[Pass / fail / N/A]", "[Notes / photo IDs]", "[Yes / no]"),
            ("Bathroom 1", "[Toilet, basin, shower/bath, mirror, taps, floor, bins]", "[Pass / fail / N/A]", "[Notes / photo IDs]", "[Yes / no]"),
            ("Bathroom 2", "[As applicable]", "[Pass / fail / N/A]", "[Notes / photo IDs]", "[Yes / no]"),
            ("Bedroom / office", "[Dusting, floors, surfaces, bins, touchpoints]", "[Pass / fail / N/A]", "[Notes / photo IDs]", "[Yes / no]"),
            ("Living / reception", "[Surfaces, floors, touchpoints, presentation]", "[Pass / fail / N/A]", "[Notes / photo IDs]", "[Yes / no]"),
            ("Common / exterior", "[Entrances, stairs, windows, waste, exterior area]", "[Pass / fail / N/A]", "[Notes / photo IDs]", "[Yes / no]"),
        ],
        [1700, 3500, 1500, 1660, 1000],
        font_size=8.2,
    )
    add_heading(doc, "Appendix B: Photo Evidence Register", 1)
    add_table(
        doc,
        ["Photo ID", "Area / item", "Reason for photo", "Privacy check", "Stored at"],
        [
            ("[001]", "[Area/item]", "[Before / after / defect / correction]", "[No unnecessary PI or private content]", "[Folder/link]"),
            ("[002]", "[Area/item]", "[Reason]", "[Privacy check]", "[Folder/link]"),
            ("[003]", "[Area/item]", "[Reason]", "[Privacy check]", "[Folder/link]"),
            ("[004]", "[Area/item]", "[Reason]", "[Privacy check]", "[Folder/link]"),
        ],
        [1200, 2100, 2500, 2200, 1360],
        font_size=8.4,
    )
    add_heading(doc, "Appendix C: Client Feedback / Complaint Capture", 1)
    add_table(
        doc,
        ["Feedback source", "Details", "Immediate response", "Escalated to", "Close-out date"],
        [
            ("[Client / representative / platform]", "[Compliment / complaint / request / scope change]", "[Response]", "[Name]", "[Date]"),
            ("[Source]", "[Details]", "[Response]", "[Name]", "[Date]"),
            ("[Source]", "[Details]", "[Response]", "[Name]", "[Date]"),
        ],
        [1700, 3200, 2000, 1300, 1160],
        font_size=8.4,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix D: Monthly Quality Trend Log", 1)
    add_table(
        doc,
        ["Month", "Audits completed", "Average score", "Top repeat defects", "Actions / training completed"],
        [
            ("[Month]", "[No.]", "[%]", "[Defects]", "[Actions]"),
            ("[Month]", "[No.]", "[%]", "[Defects]", "[Actions]"),
            ("[Month]", "[No.]", "[%]", "[Defects]", "[Actions]"),
            ("[Month]", "[No.]", "[%]", "[Defects]", "[Actions]"),
        ],
        [1300, 1600, 1400, 3100, 1960],
        font_size=8.5,
    )
    add_heading(doc, "Appendix E: Source and Review List", 1)
    add_table(
        doc,
        ["Source", "URL / review note"],
        [
            ("Occupational Health and Safety Act 85 of 1993", "https://www.gov.za/documents/occupational-health-and-safety-act"),
            ("Regulations for Hazardous Chemical Agents, 2021", "https://www.labour.gov.za/DocumentCenter/Publications/Occupational%20Health%20and%20Safety/Regulations%20for%20Hazardous%20Chemical%20Agents%202021.pdf"),
            ("Consumer Protection Act 68 of 2008", "https://www.gov.za/documents/consumer-protection-act"),
            ("Protection of Personal Information Act 4 of 2013", "https://www.gov.za/documents/protection-personal-information-act"),
            ("National Environmental Management: Waste Act 59 of 2008", "https://www.gov.za/documents/national-environmental-management-waste-act"),
        ],
        [3400, 5960],
        font_size=8.0,
    )


def build_doc():
    doc = Document()
    configure_styles(doc)
    set_headers_and_footers(doc)

    add_cover(doc)
    add_document_control(doc)
    add_scoring(doc)
    add_compliance_reference(doc)
    add_audit_details(doc)
    add_audit_matrices(doc)
    add_corrective_action(doc)
    add_appendices(doc)

    doc.core_properties.title = "Quality Control Audit Form"
    doc.core_properties.subject = "South African cleaning services quality control audit form"
    doc.core_properties.keywords = "cleaning services, South Africa, quality control, audit, OHS, POPIA, HCA, corrective action"
    doc.core_properties.comments = "Template generated for business use; verify current legal, client, site, and product requirements before implementation."
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUT.resolve()}")
