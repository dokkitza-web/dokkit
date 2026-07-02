from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor, Twips


SKILL_SCRIPTS = Path(
    r"C:\Users\Elzano Cox\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


OUT = Path("docs/SA_Cleaning_Services_Operations_Manual_Template.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "7A5A00"
WHITE = "FFFFFF"
INK = "202124"
RED = "9B1C1C"
GREEN = "1F5F3A"

GENERATED_DATE = "27 June 2026"


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D0D7E2", sz="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear_cell(cell):
    cell.text = ""


def add_cell_text(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table(doc, headers, rows, weights=None, header_fill=LIGHT_BLUE, font_size=9.25, total_width=None):
    if weights is None:
        weights = [1] * len(headers)
    if total_width is None:
        total_width = section_content_width_dxa(doc.sections[-1])
    widths = column_widths_from_weights(weights, total_width)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    mark_repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell, color="B7C7DA")
        add_cell_text(cell, header, bold=True, color=NAVY, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            add_cell_text(cell, value, size=font_size)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_two_col_table(doc, rows, label_width=0.28, font_size=9.5):
    return add_table(
        doc,
        ["Field", "Template entry"],
        rows,
        weights=[label_width, 1 - label_width],
        font_size=font_size,
    )


def add_para(doc, text="", style=None, bold=False, color=INK, size=None, italic=False, align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_callout(doc, title, text, fill=CALLOUT, accent=NAVY):
    total_width = section_content_width_dxa(doc.sections[-1])
    table = doc.add_table(rows=1, cols=1)
    mark_repeat_header(table.rows[0])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="D7DBE2")
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(title.upper() + ": ")
    set_run_font(r1, size=9.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5, color=INK)
    apply_table_geometry(table, [total_width], table_width_dxa=total_width, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_sep)
    field_run._r.append(fld_text)
    field_run._r.append(fld_end)
    set_run_font(field_run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    # Named South African operations-manual override: A4 portrait with 20 mm margins.
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25
        st.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = ""
        run = hp.add_run("Operations Manual Template | South African Cleaning Services")
        set_run_font(run, size=8.5, color=MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        add_page_number(fp)


def cover(doc):
    add_para(doc, "OPERATIONS MANUAL TEMPLATE", bold=True, color=GOLD, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, before=74, after=16)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("South African Cleaning Services")
    set_run_font(r, size=29, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    r = subtitle.add_run("Site operations, compliance, SOPs, quality control and ready-to-use forms")
    set_run_font(r, size=13.5, color=DARK_BLUE)

    add_callout(
        doc,
        "Template purpose",
        "Adapt this manual for commercial, industrial, retail, hospitality, education, healthcare-adjacent and public-sector cleaning contracts in South Africa. Replace bracketed fields, insert site-specific risk assessments, and attach client SLA requirements before issuing.",
        fill="EEF4FB",
        accent=BLUE,
    )

    add_two_col_table(
        doc,
        [
            ("Company", "[Insert registered company name]"),
            ("Trading name", "[Insert trading name]"),
            ("Manual owner", "[Operations Manager / SHEQ Manager]"),
            ("Effective date", "[DD Month YYYY]"),
            ("Review cycle", "Annual, or sooner after legal, client, site or incident-driven change"),
            ("Generated template date", GENERATED_DATE),
        ],
        label_width=0.24,
    )

    add_para(
        doc,
        "Important: This is a business operations template, not legal advice. Confirm the latest applicable legislation, wage determinations, municipal by-laws, client specifications and bargaining council agreements before issue.",
        italic=True,
        color=MUTED,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=20,
    )
    doc.add_page_break()


def front_matter(doc):
    add_heading(doc, "Document Control", 1)
    add_two_col_table(
        doc,
        [
            ("Document title", "Operations Manual - Cleaning Services"),
            ("Document number", "[DOC-OPS-001]"),
            ("Version", "[0.1 Draft / 1.0 Approved]"),
            ("Prepared by", "[Name and role]"),
            ("Reviewed by", "[Operations / SHEQ / HR / Finance]"),
            ("Approved by", "[Managing Director / Accountable Executive]"),
            ("Applicable sites", "[All sites / list contract sites]"),
            ("Distribution", "[Internal / Client-controlled / Tender pack]"),
        ],
    )
    add_table(
        doc,
        ["Version", "Date", "Change summary", "Approved by"],
        [
            ("0.1", "[DD/MM/YYYY]", "Initial draft for internal review", "[Name]"),
            ("1.0", "[DD/MM/YYYY]", "Approved for implementation", "[Name]"),
            ("", "", "", ""),
        ],
        weights=[0.13, 0.17, 0.50, 0.20],
    )

    add_heading(doc, "How To Use This Template", 2)
    add_bullets(
        doc,
        [
            "Replace every bracketed placeholder before issuing the manual.",
            "Insert site-specific risk assessments, method statements, emergency contacts, chemical registers and client SLA schedules for each contract.",
            "Keep signed attendance, training, incident, inspection and corrective action records with the site file.",
            "Review wage rates, conditions of employment and bargaining council requirements before pricing or renewing any contract.",
            "Archive superseded versions and keep the current approved copy available to supervisors and client representatives.",
        ]
    )

    add_heading(doc, "Contents Map", 1)
    add_table(
        doc,
        ["Part", "Section"],
        [
            ("1", "Company operating framework"),
            ("2", "South African compliance register"),
            ("3", "Roles, competencies and training"),
            ("4", "Client onboarding and site mobilisation"),
            ("5", "Site risk assessment and method statements"),
            ("6", "Cleaning standards and SOPs"),
            ("7", "Chemical, equipment and PPE management"),
            ("8", "Health, safety and environmental management"),
            ("9", "People management and labour compliance"),
            ("10", "Quality assurance and client reporting"),
            ("11", "Procurement, stock and asset control"),
            ("12", "Emergency response and incident reporting"),
            ("13", "Records, POPIA and document control"),
            ("Appendices", "Operational forms and regulatory source register"),
        ],
        weights=[0.22, 0.78],
    )


def company_framework(doc):
    add_heading(doc, "1. Company Operating Framework", 1)
    add_para(
        doc,
        "This section defines how the cleaning business is governed, how contracts are controlled, and how site teams convert client requirements into safe, measurable daily work.",
    )
    add_two_col_table(
        doc,
        [
            ("Company mission", "[Insert mission statement]"),
            ("Service promise", "[Insert measurable client promise, e.g. clean, safe and inspection-ready environments]"),
            ("Operating scope", "[Commercial cleaning / industrial cleaning / hygiene services / pest interface / waste coordination / periodic deep cleaning]"),
            ("Geographic scope", "[Province(s), municipalities and branch locations]"),
            ("Excluded work", "[High-level work, confined spaces, medical waste handling, fumigation, electrical repair, or other excluded activities unless separately authorised]"),
        ],
    )
    add_heading(doc, "Operating Principles", 2)
    add_bullets(
        doc,
        [
            "No cleaning task starts until the worker understands the method, hazards, PPE and emergency controls.",
            "Chemicals are used only according to the approved product label and current Safety Data Sheet.",
            "Client property, keys, alarms, data and confidential material are protected as controlled assets.",
            "Service quality is verified through scheduled inspections, evidence-based reporting and corrective action closure.",
            "Labour, wage, health and safety and data-protection requirements are treated as contract deliverables, not back-office admin.",
        ]
    )
    add_heading(doc, "Contract File Structure", 2)
    add_table(
        doc,
        ["File section", "Minimum contents", "Owner"],
        [
            ("Contract and SLA", "Signed agreement, scope, site rules, pricing schedule, escalation contacts", "Contracts / Operations"),
            ("SHE file", "Risk assessment, method statements, induction records, incident reports, inspection records", "SHEQ / Supervisor"),
            ("HR and payroll", "Employee contracts, attendance, leave, payslips, disciplinary records, training records", "HR / Payroll"),
            ("Quality", "Inspection schedules, audit reports, non-conformance records, client satisfaction", "Quality / Operations"),
            ("Assets and stock", "Equipment register, maintenance records, chemical register, stock counts", "Stores / Supervisor"),
            ("Data and security", "Key register, access cards, client confidentiality requirements, POPIA controls", "Supervisor / Information Officer"),
        ],
        weights=[0.24, 0.56, 0.20],
    )


def compliance_register(doc):
    add_heading(doc, "2. South African Compliance Register", 1)
    add_callout(
        doc,
        "Owner action",
        "Assign each compliance line to a named owner and review it at least quarterly. For tenders and public-sector work, attach proof documents to the tender file and site mobilisation file.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_table(
        doc,
        ["Area", "Requirement to confirm", "Evidence to keep", "Owner"],
        [
            ("Company status", "CIPC registration, SARS tax compliance status, VAT where applicable, CSD registration for public-sector work, B-BBEE affidavit/certificate where required.", "Registration docs, tax PIN, CSD report, B-BBEE proof", "[Finance / Admin]"),
            ("Labour conditions", "BCEA, LRA, National Minimum Wage Act and Sectoral Determination 1 for Contract Cleaning. Check current Area A, Area C and KwaZulu-Natal/BCCCI rates before pricing.", "Contracts, wage schedule, payslips, time records, leave records", "[HR / Payroll]"),
            ("COIDA", "Register with the Compensation Fund, submit annual Return of Earnings, keep Letter of Good Standing current, and report occupational injuries or diseases.", "COIDA registration, ROE, LOGS, claim forms", "[HR / SHEQ]"),
            ("UIF / PAYE", "Register and submit statutory payroll deductions and returns where applicable.", "UIF records, EMP returns, payroll reports", "[Payroll]"),
            ("OHS Act", "Provide and maintain a working environment that is safe and without risk as far as reasonably practicable, including risk assessments, safe work procedures, supervision and consultation.", "SHE policy, HIRA, SOPs, inspections, incident register", "[SHEQ / Operations]"),
            ("Hazardous chemicals", "Maintain a chemical inventory, current SDSs, GHS labels, training, storage compatibility, spill response and exposure controls.", "Chemical register, SDS file, training, spill checks", "[SHEQ / Stores]"),
            ("Biological exposure", "Control cleaning tasks involving blood, body fluids, sharps, sanitary waste, mould and infection risks through risk assessment, PPE and disposal controls.", "HBA risk notes, incident records, PPE issue, waste records", "[SHEQ / Supervisor]"),
            ("Waste and environment", "Follow National Environmental Management: Waste Act principles, municipal by-laws and client waste segregation/disposal rules.", "Waste transfer notes, client procedures, municipal requirements", "[Operations]"),
            ("POPIA / PAIA", "Register/confirm Information Officer arrangements, protect employee and client personal information, and control access to records.", "Privacy notice, IO registration proof, access log, retention schedule", "[Information Officer]"),
            ("Training and skills", "Maintain induction, task competency, toolbox talk and refresher training records; align learnerships with Skills Development Act where used.", "Training matrix, certificates, attendance registers", "[HR / Training]"),
        ],
        weights=[0.18, 0.39, 0.29, 0.14],
        font_size=8.5,
    )
    add_heading(doc, "Current Wage Checkpoint", 2)
    add_para(
        doc,
        "As at the 2026 Government Gazette used for this template, the national minimum wage is R30.23 per ordinary hour from 1 March 2026. Contract Cleaning Sector rates shown in the gazette include Area A at R33.27 per hour, Area C at R30.33 per hour, and KwaZulu-Natal/Area B subject to BCCCI rates. Replace this note with the latest wage schedule before issuing any priced contract or payroll instruction.",
    )


def roles_training(doc):
    add_heading(doc, "3. Roles, Competencies and Training", 1)
    add_table(
        doc,
        ["Role", "Primary responsibilities", "Minimum competency evidence"],
        [
            ("Managing Director", "Legal accountability, strategic contracts, resource approval, major client escalation.", "Appointment letter, delegation matrix"),
            ("Operations Manager", "Contract mobilisation, supervisor support, KPI review, corrective action closure.", "Operations experience, SLA management, incident escalation training"),
            ("SHEQ Manager", "Risk framework, audits, OHS compliance, chemical controls, incident investigation.", "OHS training, HIRA, incident investigation, internal auditor"),
            ("Site Supervisor", "Daily deployment, attendance, quality checks, toolbox talks, stock control, client liaison.", "Supervisor induction, site rules, task SOP competency"),
            ("Cleaner / Hygiene Operator", "Perform assigned cleaning tasks safely, report hazards, use PPE, protect client assets.", "Site induction, task training, chemical/PPE awareness"),
            ("Stores / Procurement", "Approved products, stock levels, SDS requests, equipment issue and repair tracking.", "Procurement procedure, stock control process"),
            ("Information Officer", "POPIA/PAIA compliance, privacy incidents, data access and retention controls.", "Information Officer registration/evidence, privacy training"),
        ],
        weights=[0.22, 0.50, 0.28],
        font_size=8.75,
    )
    add_heading(doc, "Training Matrix", 2)
    add_table(
        doc,
        ["Training item", "Cleaner", "Supervisor", "Manager", "Frequency / trigger"],
        [
            ("Company induction and code of conduct", "Required", "Required", "Required", "Before first shift"),
            ("Client site induction and access rules", "Required", "Required", "As applicable", "Before site access"),
            ("Chemical safety, SDS and GHS labels", "Required", "Required", "Required", "Initial and product change"),
            ("Task SOP competency", "Required", "Required", "Awareness", "Before task and annually"),
            ("PPE selection, care and issue", "Required", "Required", "Awareness", "Initial and PPE change"),
            ("Incident, near miss and hazard reporting", "Required", "Required", "Required", "Initial and annual refresher"),
            ("Emergency response and evacuation", "Required", "Required", "Required", "Site induction and drills"),
            ("POPIA, confidentiality and client security", "Required", "Required", "Required", "Initial and annual refresher"),
            ("Supervisor inspection and coaching", "N/A", "Required", "Awareness", "Before appointment"),
        ],
        weights=[0.36, 0.14, 0.14, 0.14, 0.22],
        font_size=8.25,
    )
    add_heading(doc, "Competency Sign-Off Rule", 2)
    add_bullets(
        doc,
        [
            "Workers may observe a task during training, but may not perform it unsupervised until signed competent.",
            "Competency is task-specific: wet mopping does not authorise chemical stripping, machine scrubbing or high-level window cleaning.",
            "Refresh competency after an incident, product change, equipment change, client complaint or long absence from the task.",
        ]
    )


def onboarding_risk(doc):
    add_heading(doc, "4. Client Onboarding and Site Mobilisation", 1)
    add_table(
        doc,
        ["Phase", "Key actions", "Outputs"],
        [
            ("1. Pre-award review", "Review tender/SLA, site type, hours, wage area, risks, security rules, waste streams and client reporting requirements.", "Bid risk review, pricing assumptions, legal/wage check"),
            ("2. Contract handover", "Confirm scope, exclusions, start date, access, site contact, escalation route, site rules and mobilisation budget.", "Signed handover note, mobilisation plan"),
            ("3. Site survey", "Walk the site with client; inspect floors, washrooms, kitchens, waste points, stores, electrical points, water points and hazards.", "Site survey checklist, photo record, floor plan notes"),
            ("4. Risk and method", "Complete HIRA, method statements, chemical list, PPE plan, emergency contacts and training needs.", "SHE file, SOP pack, site induction pack"),
            ("5. Mobilise", "Deploy staff, stock, PPE, equipment, registers, signage, supervision schedule and reporting templates.", "Mobilisation checklist signed off"),
            ("6. Stabilise", "Run first-week inspections, resolve snags, confirm consumable usage and adjust frequencies.", "Snag list, first-week report, revised task schedule"),
        ],
        weights=[0.20, 0.49, 0.31],
        font_size=8.75,
    )
    add_heading(doc, "Site Risk Assessment and Method Statement Control", 1)
    add_para(doc, "Each site must have a risk assessment and method statement pack that matches the actual work performed. Use the following baseline hazard register as a starting point, then add site-specific hazards.")
    add_table(
        doc,
        ["Hazard", "Typical controls", "Records"],
        [
            ("Slips, trips and wet floors", "Wet floor signs, work zoning, dry mopping before wet cleaning, cable control, spill response.", "Inspection checklist, incident reports"),
            ("Chemical exposure", "Approved chemicals, SDS access, dilution controls, gloves/eye protection, no mixing, ventilated storage.", "Chemical register, SDS file, training"),
            ("Biological exposure", "Disposable gloves, eye/face protection where needed, sharps protocol, body fluid spill kit, hand hygiene.", "Exposure report, waste disposal proof"),
            ("Manual handling", "Trolley use, team lifts, storage at safe height, training, weight limits.", "Training records, maintenance checks"),
            ("Electrical equipment", "Pre-use checks, no damaged plugs/leads, dry hands, RCD where required, isolation before cleaning near equipment.", "Equipment register, defect log"),
            ("Working at height", "Avoid unless authorised; use stable step ladders only where assessed; no improvised access.", "Permit/authorisation, inspection record"),
            ("Public/client interface", "Secure work area, respectful conduct, confidentiality, visitor awareness, complaint escalation.", "Site rules, incident/client logs"),
            ("Security/key control", "Signed key issue, access card log, no sharing, immediate loss reporting.", "Key register, access log"),
            ("Load-shedding or water interruption", "Alternate task plan, safe storage, backup water rules where approved, client communication.", "Contingency log"),
        ],
        weights=[0.22, 0.55, 0.23],
        font_size=8.5,
    )
    add_callout(
        doc,
        "Method statement rule",
        "A method statement must say who does the work, what materials are used, what PPE is required, the safe step sequence, how waste is handled, how quality is checked and when work must stop.",
    )


def cleaning_standards(doc):
    add_heading(doc, "5. Cleaning Standards and Frequencies", 1)
    add_para(doc, "Replace the frequencies below with the signed SLA. Where client instructions conflict with safe work requirements, escalate before work proceeds.")
    add_table(
        doc,
        ["Area / task", "Typical frequency", "Minimum acceptance standard"],
        [
            ("Entrances and reception", "Daily and touch-up as required", "Floors clean and dry; glass clear; bins not overflowing; no visible litter."),
            ("Offices and meeting rooms", "Daily or agreed schedule", "Work surfaces dust-free where accessible; floors clean; bins emptied; chairs aligned."),
            ("Washrooms", "High frequency", "Toilets, basins, taps, mirrors and touch points clean; consumables stocked; odour controlled."),
            ("Kitchens and canteens", "Daily and after service periods", "Food-contact-adjacent surfaces cleaned; sinks clear; bins managed; no residue or odour."),
            ("Floors - hard surfaces", "Daily / periodic", "No visible soil, streaks, standing water, gum or avoidable scuff marks."),
            ("Carpets", "Vacuum per SLA; spot clean as needed", "No visible loose debris; spot issues logged and treated within method limits."),
            ("Waste", "Daily or SLA", "Waste removed to approved points; recycling separated where required; liners replaced."),
            ("Touch points", "Daily / high-frequency sites", "Door handles, rails, switches and shared surfaces cleaned to site protocol."),
            ("Periodic deep clean", "Monthly / quarterly / annual", "Planned task completed, inspected and signed off with evidence."),
        ],
        weights=[0.28, 0.20, 0.52],
        font_size=8.75,
    )
    add_heading(doc, "Service Levels and KPIs", 2)
    add_table(
        doc,
        ["KPI", "Target", "Evidence"],
        [
            ("Attendance and shift coverage", "[Insert target, e.g. 100% planned shifts covered]", "Roster, clocking, supervisor report"),
            ("Inspection score", "[Insert target, e.g. 90% average or higher]", "Inspection checklist and monthly dashboard"),
            ("Complaint response", "[Insert target, e.g. acknowledge within 2 hours]", "Complaint log and corrective action"),
            ("Incident reporting", "Immediate escalation for serious incidents; same-shift near miss reporting", "Incident report and investigation"),
            ("Stock availability", "No critical consumable stock-outs", "Stock count and order log"),
            ("Training compliance", "100% mandatory induction and task training", "Training matrix and signed attendance"),
        ],
        weights=[0.30, 0.36, 0.34],
        font_size=8.75,
    )


def sops(doc):
    add_heading(doc, "6. Standard Operating Procedures", 1)
    add_para(doc, "Use these SOPs as baseline procedures. Each site must adapt them to the client's surfaces, products, equipment, risk assessment and SLA.")
    sop_data = [
        (
            "SOP 1: General Office Cleaning",
            "Clean offices, meeting rooms and shared work areas without disturbing client documents, IT equipment or personal belongings.",
            ["Disposable gloves where required", "Closed shoes", "Approved cloths and detergent", "Waste liners", "Vacuum or mop system"],
            [
                "Check the area for hazards, confidential material, damaged furniture, spills or client restrictions.",
                "Empty bins and replace liners; do not inspect, remove or read documents.",
                "Dust accessible surfaces without moving sensitive equipment or paperwork.",
                "Clean touch points such as door handles, switches and shared table surfaces.",
                "Clean floors according to floor type; leave wet floor signs until dry.",
                "Report damage, pests, spills, security issues or inaccessible areas to the supervisor.",
            ],
            ["No visible litter or dust build-up", "Bins managed", "Floor clean and safe", "Client property undisturbed"],
        ),
        (
            "SOP 2: Washroom Cleaning",
            "Maintain hygienic washrooms, control odour and keep consumables available.",
            ["Chemical-resistant gloves", "Eye protection if splash risk", "Colour-coded cloths", "Toilet brush", "Approved disinfectant/detergent", "Wet floor signs"],
            [
                "Place wet floor signs and ventilate where possible.",
                "Remove waste and sanitary waste only according to the approved route and container system.",
                "Apply approved product to toilets, urinals, basins and touch points for the required contact time.",
                "Clean mirrors, taps, dispensers, partitions, handles and flush controls.",
                "Mop floors from the farthest point toward the exit; prevent public access to wet areas.",
                "Restock toilet paper, soap, hand towels and other contracted consumables.",
                "Complete the washroom inspection record at the required interval.",
            ],
            ["No visible soil", "Consumables stocked", "No avoidable odour", "Inspection sheet current"],
        ),
        (
            "SOP 3: Kitchen and Canteen Cleaning",
            "Clean non-cooking kitchen/canteen areas while preventing contamination and pest attraction.",
            ["Gloves", "Apron where required", "Food-area approved detergent", "Clean cloths", "Waste liners"],
            [
                "Confirm whether any food-contact or catering equipment is excluded from the cleaning scope.",
                "Remove waste, replace liners and clean bin exteriors.",
                "Clean tables, counters, sinks, taps, splashbacks and handles using approved product.",
                "Do not mix chemicals or use washroom cloths/equipment in kitchen areas.",
                "Mop floors and leave the area dry and safe.",
                "Report blocked drains, pests, odours, leaks or food contamination concerns.",
            ],
            ["Surfaces visibly clean", "No cross-contamination", "Waste controlled", "Defects reported"],
        ),
        (
            "SOP 4: Hard Floor Wet Mopping",
            "Remove soil from hard floors while controlling slip risk.",
            ["Gloves if chemical contact", "Wet floor signs", "Mop system", "Approved detergent", "Closed shoes"],
            [
                "Dry sweep or dust mop first to remove loose debris.",
                "Place wet floor signs at all access points.",
                "Prepare chemical dilution according to label and SDS; never exceed the approved dilution.",
                "Mop in sections and keep public/client traffic away from wet areas.",
                "Change dirty water before it spreads soil.",
                "Inspect for streaks, residue and standing water before removing signs.",
            ],
            ["Floor clean", "No standing water", "Signs used", "Correct dilution followed"],
        ),
        (
            "SOP 5: Waste Handling",
            "Remove waste safely and protect staff, client personnel and the environment.",
            ["Gloves", "Closed shoes", "Trolley where needed", "Approved bags/liners", "Sharps protocol where applicable"],
            [
                "Check the waste type and confirm it is within the contracted scope.",
                "Do not compress bags by hand; watch for sharps, broken glass and liquids.",
                "Tie bags securely and use trolleys for heavy or multiple loads.",
                "Move waste only to approved holding or collection points.",
                "Clean spills immediately and report any hazardous or out-of-scope waste.",
                "Wash hands after removing gloves.",
            ],
            ["Waste removed", "No leakage", "Segregation followed", "Hazards escalated"],
        ),
        (
            "SOP 6: Body Fluid or Sharps Response",
            "Respond to blood, vomit, faeces, needles or other biological hazards without exposing workers or the public.",
            ["Disposable gloves", "Eye/face protection if splash risk", "Apron", "Biohazard/spill kit", "Tongs or sharps container where available"],
            [
                "Isolate the area and notify the supervisor immediately.",
                "Do not pick up sharps by hand; use tongs or the approved sharps container method.",
                "Apply the approved spill procedure and product contact time.",
                "Dispose of contaminated materials according to site waste rules.",
                "Remove PPE safely and wash hands thoroughly.",
                "Complete an incident/exposure report and arrange medical referral if exposure occurred.",
            ],
            ["Area isolated", "No bare-hand contact", "Waste handled correctly", "Exposure reported"],
        ),
    ]
    for title, purpose, equipment, steps, acceptance in sop_data:
        add_heading(doc, title, 2)
        add_para(doc, purpose)
        add_table(
            doc,
            ["PPE / materials", "Acceptance checks"],
            [("; ".join(equipment), "; ".join(acceptance))],
            weights=[0.50, 0.50],
            font_size=8.75,
        )
        add_numbers(doc, steps)


def resources_safety_people(doc):
    add_heading(doc, "7. Chemical, Equipment and PPE Management", 1)
    add_heading(doc, "Chemical Control Rules", 2)
    add_bullets(
        doc,
        [
            "Only approved chemicals on the site chemical register may be used.",
            "Every chemical must have a current Safety Data Sheet available to staff and supervisors.",
            "Labels must remain readable and aligned with GHS classification and supplier instructions.",
            "Decanting is allowed only into approved labelled containers; never into beverage bottles or unmarked containers.",
            "Chemicals must not be mixed unless the supplier's written instructions explicitly allow it.",
            "Store incompatible products separately and keep chemical storage locked or controlled.",
        ]
    )
    add_table(
        doc,
        ["Resource", "Control requirement", "Record"],
        [
            ("PPE", "Issue based on risk assessment, train the user, replace when damaged or contaminated.", "PPE issue register"),
            ("Electrical equipment", "Pre-use inspection, defect tagging, planned maintenance, safe storage.", "Equipment register and defect log"),
            ("Mops/cloths", "Colour-code by area, launder/dry correctly, replace when worn.", "Cleaning issue/checklist"),
            ("Floor machines", "Use only by trained operators; inspect cable, plug, brushes/pads and guards before use.", "Training and maintenance record"),
            ("Chemical dosing", "Use supplier-approved dilution system or measured method; prohibit free-pouring concentrates.", "Chemical register and supervisor checks"),
        ],
        weights=[0.22, 0.56, 0.22],
        font_size=8.75,
    )

    add_heading(doc, "8. Health, Safety and Environmental Management", 1)
    add_para(doc, "The company will maintain a practical SHE system proportionate to cleaning risks and client requirements.")
    add_table(
        doc,
        ["SHE element", "Minimum implementation"],
        [
            ("Policy and appointments", "Issue SHE policy; define legal and operational appointments where applicable."),
            ("Risk assessment", "Complete baseline and site/task HIRA; review after incidents, scope change or annual review."),
            ("Consultation", "Use toolbox talks, safety meetings and health and safety representatives/committee where thresholds apply."),
            ("Incident management", "Report, investigate and close corrective actions; escalate serious injuries immediately."),
            ("Emergency readiness", "Keep site emergency contacts, assembly points, spill kits, first aid arrangements and evacuation rules current."),
            ("Environmental control", "Prevent chemical spills, overuse of water, blocked drains and incorrect waste disposal."),
        ],
        weights=[0.30, 0.70],
        font_size=8.75,
    )
    add_heading(doc, "Environmental Practices", 2)
    add_bullets(
        doc,
        [
            "Prefer concentrated products with controlled dilution where safe and approved.",
            "Do not discharge prohibited chemicals into stormwater drains.",
            "Separate recyclable, general and hazardous/special waste according to client and municipal rules.",
            "Report leaks, blocked drains and abnormal odours immediately.",
            "Track consumption of chemicals, consumables and water where the client requires sustainability reporting.",
        ]
    )

    add_heading(doc, "9. People Management and Labour Compliance", 1)
    add_table(
        doc,
        ["Process", "Manual requirement", "Record"],
        [
            ("Recruitment", "Verify identity/right to work, role fit, references where used and job requirements.", "Recruitment checklist"),
            ("Employment contract", "Issue written particulars, wage rate, hours, overtime rules, leave and site conditions.", "Signed contract"),
            ("Time and attendance", "Record actual hours, overtime approval, absences and site coverage.", "Clocking/attendance register"),
            ("Payroll", "Pay at least the applicable legal/sectoral minimum; keep payslips and deductions records.", "Payroll reports and payslips"),
            ("Leave", "Manage annual, sick, family responsibility and other leave according to law and contract.", "Leave forms and balances"),
            ("Discipline and grievance", "Use fair procedure, evidence and consistent sanctions; escalate disputes promptly.", "Case file and outcome letter"),
            ("Transport and uniforms", "Define whether provided by employer/client/employee and how issue/recovery is managed.", "Issue register and policy acknowledgement"),
        ],
        weights=[0.25, 0.52, 0.23],
        font_size=8.5,
    )
    add_callout(
        doc,
        "Wage pricing risk",
        "Do not price a contract below the cost required to comply with wage, overtime, leave, UIF, COIDA, supervision, PPE, chemicals, equipment and statutory obligations. Section 200B liability risk under the LRA is specifically flagged in the 2026 wage gazette for contract cleaning providers and clients.",
        fill="FDECEC",
        accent=RED,
    )


def quality_procurement_emergency_records(doc):
    add_heading(doc, "10. Quality Assurance and Client Reporting", 1)
    add_table(
        doc,
        ["Control", "How it works", "Frequency"],
        [
            ("Daily supervisor walk-through", "Check critical areas, attendance, stock, incidents and client requests.", "Daily"),
            ("Formal inspection", "Score areas against SLA checklist and photograph evidence where required.", "Weekly / monthly"),
            ("Complaint log", "Record issue, root cause, action, owner and close-out confirmation.", "As raised"),
            ("Corrective action", "Use CAPA for recurring failures, safety issues or client dissatisfaction.", "As required"),
            ("Management review", "Review KPIs, incidents, staff changes, costs, audits and client feedback.", "Monthly / quarterly"),
        ],
        weights=[0.25, 0.55, 0.20],
        font_size=8.75,
    )
    add_heading(doc, "Inspection Scoring Guide", 2)
    add_table(
        doc,
        ["Score", "Meaning", "Required action"],
        [
            ("5", "Excellent: no visible defects; standard sustained.", "Recognise and maintain."),
            ("4", "Good: minor issue corrected immediately.", "Record trend if repeated."),
            ("3", "Acceptable: service meets minimum but needs improvement.", "Supervisor coaching."),
            ("2", "Poor: visible failure or missed task.", "Corrective action and re-inspection."),
            ("1", "Critical: safety, hygiene, client escalation or repeated failure.", "Immediate escalation and CAPA."),
        ],
        weights=[0.12, 0.54, 0.34],
        font_size=8.75,
    )

    add_heading(doc, "11. Procurement, Stock and Asset Control", 1)
    add_bullets(
        doc,
        [
            "Buy only approved chemicals, consumables and PPE from approved suppliers.",
            "No substitute chemical may be used until the SDS, label, suitability, dilution and risk controls are approved.",
            "Maintain minimum and maximum stock levels per site to prevent emergency purchasing and stock-outs.",
            "Tag company equipment and keep a movement/issue record for each site.",
            "Investigate abnormal stock usage as a possible quality, theft, training or dilution-control issue.",
        ]
    )
    add_table(
        doc,
        ["Stock category", "Minimum control", "Reorder trigger"],
        [
            ("Chemicals", "Approved product list, SDS, labelled storage, issue quantity control.", "[Insert minimum stock]"),
            ("Consumables", "Toilet paper, soap, towels, liners and client-specific consumables tracked by usage.", "[Insert reorder point]"),
            ("PPE", "Issue to named employees; keep spare critical PPE for replacement.", "[Insert reorder point]"),
            ("Equipment spares", "Mop heads, cloths, pads, vacuum bags/filters, extension leads where approved.", "[Insert reorder point]"),
        ],
        weights=[0.24, 0.54, 0.22],
        font_size=8.75,
    )

    add_heading(doc, "12. Emergency Response and Incident Reporting", 1)
    add_table(
        doc,
        ["Scenario", "Immediate response", "Escalation / record"],
        [
            ("Injury or illness", "Stop work, make area safe, call first aider/emergency services if needed.", "Supervisor, SHEQ, incident report, COIDA process"),
            ("Chemical spill", "Isolate area, check SDS, use spill kit/PPE, prevent drain entry if unsafe.", "Supervisor, client, spill report"),
            ("Fire / evacuation", "Raise alarm, evacuate by site route, do not re-enter until authorised.", "Roll call, client emergency lead"),
            ("Security incident", "Do not confront if unsafe; withdraw and notify security/client/supervisor.", "Security log, incident report"),
            ("Needle/sharps exposure", "Wash area, seek medical advice, preserve details, report immediately.", "Exposure report, medical referral"),
            ("Client complaint", "Acknowledge, make safe/correct immediate issue, record facts.", "Complaint log, corrective action"),
        ],
        weights=[0.24, 0.50, 0.26],
        font_size=8.75,
    )
    add_heading(doc, "Incident Classification", 2)
    add_bullets(
        doc,
        [
            "Near miss: event could have caused injury, illness, damage or client disruption.",
            "First aid: minor injury managed by first aid without lost time.",
            "Medical treatment / lost time: worker requires medical care or cannot complete normal duties.",
            "Environmental incident: spill, incorrect disposal, drain contamination or significant waste failure.",
            "Quality incident: service failure that affects client operations, hygiene or contract performance.",
        ]
    )

    add_heading(doc, "13. Records, POPIA and Document Control", 1)
    add_para(doc, "Cleaning contractors often handle employee records, client access data, key registers, CCTV-adjacent information, visitor areas and confidential client spaces. The manual must therefore include practical privacy and record controls.")
    add_table(
        doc,
        ["Record type", "Minimum retention/control", "Access owner"],
        [
            ("Employee records", "Store securely; restrict to HR/payroll/authorised management; retain according to legal and company schedule.", "HR"),
            ("Training and induction", "Keep signed records available for client and inspector audit.", "HR / SHEQ"),
            ("Incident and COIDA records", "Secure, complete and escalate according to statutory/client requirements.", "SHEQ / HR"),
            ("Client keys/access cards", "Issue only by signed register; report loss immediately.", "Supervisor"),
            ("Inspection and quality records", "Keep by site and month; use for KPI reporting and corrective actions.", "Operations"),
            ("Supplier/SDS records", "Keep current SDS and approved supplier records accessible to site teams.", "Stores / SHEQ"),
            ("Privacy incidents", "Record suspected unauthorised access, loss or disclosure; escalate to Information Officer.", "Information Officer"),
        ],
        weights=[0.30, 0.52, 0.18],
        font_size=8.5,
    )
    add_heading(doc, "Document Review Triggers", 2)
    add_bullets(
        doc,
        [
            "Change in law, wage determination, bargaining council agreement or client site rules.",
            "New chemical, equipment, task, site type or work method.",
            "Serious incident, recurring complaint, failed audit or enforcement notice.",
            "Contract renewal, scope change, price review or mobilisation of a new site.",
        ]
    )


def appendices(doc):
    doc.add_page_break()
    add_heading(doc, "Appendices: Operational Forms", 1)
    add_para(doc, "Copy these forms into site files or convert them into your digital workflow. Add site-specific fields where required.")

    add_heading(doc, "Appendix A: Site Mobilisation Checklist", 2)
    add_table(
        doc,
        ["Item", "Done", "Notes / owner"],
        [
            ("Signed contract, SLA and scope received", "[ ]", ""),
            ("Client site rules and emergency contacts received", "[ ]", ""),
            ("Site survey completed with photos/floor plan notes", "[ ]", ""),
            ("Wage area and applicable bargaining council/rate confirmed", "[ ]", ""),
            ("Risk assessment and method statements approved", "[ ]", ""),
            ("Chemical register and SDS file prepared", "[ ]", ""),
            ("PPE issued and recorded", "[ ]", ""),
            ("Staff inducted and task-trained", "[ ]", ""),
            ("Equipment delivered, tagged and inspected", "[ ]", ""),
            ("Inspection and reporting templates agreed", "[ ]", ""),
        ],
        weights=[0.58, 0.12, 0.30],
        font_size=8.75,
    )

    add_heading(doc, "Appendix B: Daily Cleaning Checklist", 2)
    add_table(
        doc,
        ["Area", "Tasks completed", "Initial", "Issues / follow-up"],
        [
            ("Reception / entrance", "Floors, glass, bins, touch points", "", ""),
            ("Offices / meeting rooms", "Bins, surfaces, floors, spot clean", "", ""),
            ("Washrooms", "Toilets, basins, mirrors, floors, consumables", "", ""),
            ("Kitchen / canteen", "Counters, sinks, tables, bins, floors", "", ""),
            ("Waste points", "Removed, liners replaced, segregation followed", "", ""),
            ("Supervisor check", "Defects logged, client requests captured", "", ""),
        ],
        weights=[0.25, 0.42, 0.12, 0.21],
        font_size=8.5,
    )

    add_heading(doc, "Appendix C: Chemical Register", 2)
    add_table(
        doc,
        ["Product", "Use", "SDS date", "Dilution", "PPE", "Storage / compatibility"],
        [
            ("[Product name]", "[Washroom / floor / glass / disinfectant]", "[Date]", "[Ratio]", "[Gloves/eye/etc.]", "[Cabinet / separate from acids/etc.]"),
            ("", "", "", "", "", ""),
            ("", "", "", "", "", ""),
            ("", "", "", "", "", ""),
        ],
        weights=[0.18, 0.19, 0.13, 0.13, 0.15, 0.22],
        font_size=8.0,
    )

    add_heading(doc, "Appendix D: Equipment Maintenance Register", 2)
    add_table(
        doc,
        ["Asset ID", "Equipment", "Site", "Pre-use issue", "Service due", "Action taken"],
        [
            ("[ID]", "[Vacuum / scrubber / trolley]", "[Site]", "[Yes/No + detail]", "[Date]", "[Repair / replace / monitor]"),
            ("", "", "", "", "", ""),
            ("", "", "", "", "", ""),
        ],
        weights=[0.13, 0.22, 0.16, 0.22, 0.13, 0.14],
        font_size=8.0,
    )

    add_heading(doc, "Appendix E: PPE Issue Register", 2)
    add_table(
        doc,
        ["Employee", "PPE item", "Size", "Date issued", "Condition", "Signature"],
        [
            ("[Name]", "[Gloves / goggles / apron / shoes]", "[Size]", "[Date]", "[New/replacement]", ""),
            ("", "", "", "", "", ""),
            ("", "", "", "", "", ""),
        ],
        weights=[0.22, 0.26, 0.10, 0.16, 0.14, 0.12],
        font_size=8.0,
    )

    add_heading(doc, "Appendix F: Toolbox Talk Attendance", 2)
    add_table(
        doc,
        ["Topic", "Date", "Presenter", "Key risk discussed"],
        [
            ("[Chemical safety / slips and trips / client security]", "[Date]", "[Name]", "[Risk and control summary]"),
        ],
        weights=[0.31, 0.16, 0.22, 0.31],
        font_size=8.5,
    )
    add_table(
        doc,
        ["Employee name", "Employee no.", "Signature", "Questions / concerns raised"],
        [
            ("", "", "", ""),
            ("", "", "", ""),
            ("", "", "", ""),
            ("", "", "", ""),
            ("", "", "", ""),
        ],
        weights=[0.28, 0.18, 0.20, 0.34],
        font_size=8.5,
    )

    add_heading(doc, "Appendix G: Incident / Near Miss Report", 2)
    add_two_col_table(
        doc,
        [
            ("Site and department", ""),
            ("Date and time", ""),
            ("Person(s) involved", ""),
            ("Incident type", "[Near miss / injury / exposure / spill / security / quality]"),
            ("What happened", ""),
            ("Immediate action taken", ""),
            ("Root cause", ""),
            ("Corrective action", ""),
            ("Owner and due date", ""),
            ("Closed by / date", ""),
        ],
        label_width=0.30,
        font_size=8.5,
    )

    add_heading(doc, "Appendix H: Non-Conformance and Corrective Action Log", 2)
    add_table(
        doc,
        ["Ref", "Issue", "Root cause", "Action", "Owner", "Due", "Closed"],
        [
            ("NCR-001", "[Describe issue]", "[Cause]", "[Action]", "[Name]", "[Date]", "[Y/N]"),
            ("", "", "", "", "", "", ""),
            ("", "", "", "", "", "", ""),
        ],
        weights=[0.10, 0.25, 0.20, 0.20, 0.11, 0.08, 0.06],
        font_size=7.75,
    )

    add_heading(doc, "Appendix I: Monthly Client Report Template", 2)
    add_two_col_table(
        doc,
        [
            ("Client / site", ""),
            ("Reporting month", ""),
            ("Attendance summary", "[Planned shifts vs actual coverage]"),
            ("Inspection score", "[Average score and trend]"),
            ("Complaints / compliments", "[Number and summary]"),
            ("Incidents / near misses", "[Number and status]"),
            ("Corrective actions", "[Open / closed / overdue]"),
            ("Stock or equipment concerns", ""),
            ("Client decisions needed", ""),
        ],
        label_width=0.30,
        font_size=8.5,
    )

    add_heading(doc, "Appendix J: Regulatory Source Register", 2)
    add_para(doc, "Use this source register during annual review. Links were checked when this template was generated; confirm latest amendments before issue.", italic=True, color=MUTED, size=9)
    sources = [
        ("National Minimum Wage Amendment 2026 and Contract Cleaning Sector rates", "Government Gazette 54075", "https://www.gov.za/sites/default/files/gcis_document/202602/54075rg11941gon7083.pdf"),
        ("Occupational Health and Safety Act 85 of 1993", "South African Government", "https://www.gov.za/documents/occupational-health-and-safety-act"),
        ("Regulations for Hazardous Chemical Agents, 2021", "Department of Employment and Labour", "https://www.labour.gov.za/DocumentCenter/Publications/Occupational%20Health%20and%20Safety/Regulations%20for%20Hazardous%20Chemical%20Agents%202021.pdf"),
        ("Basic Conditions of Employment Act 75 of 1997", "South African Government", "https://www.gov.za/documents/basic-conditions-employment-act"),
        ("Labour Relations Act 66 of 1995", "South African Government", "https://www.gov.za/documents/labour-relations-act"),
        ("Compensation for Occupational Injuries and Diseases Act 130 of 1993", "South African Government", "https://www.gov.za/documents/compensation-occupational-injuries-and-diseases-act"),
        ("Protection of Personal Information Act 4 of 2013", "South African Government", "https://www.gov.za/documents/protection-personal-information-act"),
        ("Information Regulator eServices / Information Officer guidance", "Information Regulator", "https://inforegulator.org.za/portal-posts/"),
        ("National Environmental Management: Waste Act 59 of 2008", "South African Government", "https://www.gov.za/documents/national-environmental-management-waste-act"),
        ("Skills Development Act 97 of 1998", "South African Government", "https://www.gov.za/documents/skills-development-act"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    mark_repeat_header(table.rows[0])
    for idx, h in enumerate(["Source", "Link"]):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell, color="B7C7DA")
        add_cell_text(cell, h, bold=True, color=NAVY, size=8.25)
    for label, display, url in sources:
        row = table.add_row().cells
        set_cell_border(row[0])
        add_cell_text(row[0], label, size=7.75)
        set_cell_border(row[1])
        clear_cell(row[1])
        p = row[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_hyperlink(p, display, url)
    total = section_content_width_dxa(doc.sections[-1])
    widths = column_widths_from_weights([0.38, 0.62], total)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=120)


def structural_audit(docx_path: Path):
    # Lightweight audit marker file for the builder log. Deeper table audit is run separately.
    if not docx_path.exists() or docx_path.stat().st_size < 10000:
        raise RuntimeError(f"Expected DOCX was not created correctly: {docx_path}")


def build():
    doc = Document()
    configure_document(doc)
    cp = doc.core_properties
    cp.title = "South African Cleaning Services Operations Manual Template"
    cp.subject = "Operations manual template for cleaning services in South Africa"
    cp.author = "DokKit"
    cp.keywords = "cleaning services, operations manual, South Africa, SOP, OHS, POPIA, contract cleaning"
    cp.comments = "Generated template. Review legal and site-specific requirements before issue."

    cover(doc)
    front_matter(doc)
    company_framework(doc)
    compliance_register(doc)
    roles_training(doc)
    onboarding_risk(doc)
    cleaning_standards(doc)
    sops(doc)
    resources_safety_people(doc)
    quality_procurement_emergency_records(doc)
    appendices(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    structural_audit(OUT)
    print(str(OUT.resolve()))


if __name__ == "__main__":
    build()
